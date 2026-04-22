import os
import sys
import json
import docx
import base64
import io
import re
from docx.shared import Inches
from openai import AzureOpenAI
from core.llm_utils import call_llm_with_retry
from dotenv import load_dotenv
from collections import deque


ENGLISH_VARIANT = "US English"  # Default, can be changed at runtime

def get_english_instruction(english_variant: str) -> str:
    # return f"""
    # **LANGUAGE REQUIREMENTS:**
    # - Use {english_variant} spelling, grammar, and terminology throughout
    # - Apply {english_variant} conventions for dates, numbers, and formatting
    # - Ensure all content follows {english_variant} language standards
    # - Examples:
    # - {english_variant} spelling: {'colour, organised, realise' if english_variant == 'UK English' else 'color, organized, realize'}
    # - {english_variant} terminology: {'programme' if english_variant == 'UK English' else 'program'} (for software/training)
    # """
    variant_examples = {
        "UK English": {
            "spelling": "colour, organised, realise, analyse, behaviour",
            "terminology": "programme (for training/software), practise (verb), licence (noun)",
            "dates": "DD/MM/YYYY",
            "punctuation": "Single quotation marks (‘ ’) preferred for quotes",
        },
        "US English": {
            "spelling": "color, organized, realize, analyze, behavior",
            "terminology": "program (for training/software), practice (both noun and verb), license (both noun and verb)",
            "dates": "MM/DD/YYYY",
            "punctuation": "Double quotation marks (“ ”) preferred for quotes",
        }
    }

    ex = variant_examples.get(english_variant, variant_examples["US English"])
    
    return f"""
        **LANGUAGE STYLE INSTRUCTION ({english_variant.upper()}):**

        You must write all responses in **{english_variant}**, strictly following its linguistic and stylistic norms.

        ✅ **Spelling & Grammar**
        - Use {english_variant} spellings (e.g., {ex['spelling']}).
        - Apply correct grammar, idioms, and phrasing natural to {english_variant} speakers.

        ✅ **Terminology & Vocabulary**
        - Use {english_variant} terminology (e.g., {ex['terminology']}).
        - Avoid mixing vocabulary from other English variants.

        ✅ **Formatting & Conventions**
        - Follow {english_variant} standards for:
        - Dates → {ex['dates']}
        - Numbers → use the standard numeric format for {english_variant}
        - Punctuation → {ex['punctuation']}

        ✅ **Tone & Usage**
        - Maintain a tone typical of professional {english_variant} communication.
        - Use culturally appropriate expressions and sentence structures.
        - Ensure consistency across the entire response.

        ❌ **Do NOT** use spelling, punctuation, or idioms from other English variants.
        """


# --- Load Environment Variables ---
load_dotenv()
# This dictionary defines the mapping from the original headers to your required new JSON headers.
HEADER_MAPPING = {
    "Course Title": "Course_Title",
    "Module Title": "Module_Title",
    "Topic": "Topic",
    "Screen-type": "Screen_type",
    "Bloom's Level": "Blooms_Level",
    "Duration (min)": "Duration_(min)",
    "Source Images (base64)": "Source_Images_(base64)",
    "Source Tables": "Source_Tables",
    "Learning Objectives": "Learning_Objectives",
    "On-screen text": "On_screen_text",
    "Narration": "Narration",
    "On-screen Recommendations": "On_screen_Recommendations",
    "Developer Notes": "Developer_Notes"
}


def transform_storyboard_headers(storyboards: list[dict], header_mapping: dict) -> list[dict]:
    """
    Converts the headers of the final storyboard list to a new format based on a mapping.
    This version includes special logic to:
    1. Simplify the 'Source Images (base64)' field into a simple list of strings.
    2. Convert specific text fields ('Learning Objectives', 'On-screen Recommendations',
       'Developer Notes') into a list of strings by splitting on newlines.
    """
    transformed_storyboards = []

    # Define the keys that we want to convert from a single string into a list of strings.
    keys_to_convert_to_list = {
        "Learning Objectives",
        "On-screen Recommendations",
        "Developer Notes"
    }

    for storyboard in storyboards:
        new_storyboard = {}
        for old_key, value in storyboard.items():
            # Get the new header name from our mapping (e.g., "Course Title" -> "Course_Title")
            new_key = header_mapping.get(old_key, old_key)

            # --- SPECIAL LOGIC BLOCKS ---

            # 1. Logic for Source Images (unchanged)
            if old_key == "Source Images (base64)":
                simple_image_list = [image_dict.get("content", "") for image_dict in value]
                new_storyboard[new_key] = simple_image_list

            # 2. NEW: Logic for converting specific text fields to lists
            elif old_key in keys_to_convert_to_list:
                # Ensure the value is a string before processing
                text_value = str(value)

                # Split the string by newlines, strip whitespace from each line,
                # and filter out any resulting empty lines.
                if text_value and text_value.strip():
                    string_list = [line.strip() for line in text_value.split('\n') if line.strip()]
                    new_storyboard[new_key] = string_list
                else:
                    # If the original value is empty, result in an empty list
                    new_storyboard[new_key] = []

            # 3. For all other keys, just copy the value as before.
            else:
                new_storyboard[new_key] = value

        transformed_storyboards.append(new_storyboard)
    return transformed_storyboards

BLOOM_VERBS = {
    1: [  # Level 1: Remember / Knowledge
        "Arrange", "Count", "Find", "Label", "List", "Match", "Order", "Recognize", "Relate"
    ],
    2: [  # Level 2: Understand / Comprehension
        "Check", "Choose", "Classify", "Identify", "Indicate", "Locate", "Select", "Translate"
    ],
    3: [  # Level 3: Apply / Application
        "Apply", "Assemble", "Break down", "Build", "Conduct", "Construct", "Demonstrate",
        "Employ", "Interpret", "Operate", "Perform", "Practice", "Schedule", "Solve", "Use"
    ]
}

# --- Azure OpenAI Client Setup ---
api_key = os.getenv("AZURE_OPENAI_API_KEY")
azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
api_version = os.getenv("AZURE_OPENAI_API_VERSION")
deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

_openai_client = None

def _get_openai_client():
    global _openai_client
    if _openai_client is None:
        if not all([api_key, azure_endpoint, api_version, deployment_name]):
            raise RuntimeError("Azure OpenAI credentials missing. Check .env file.")
        _openai_client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)
    return _openai_client


# ==============================================================================
# AI HELPER FUNCTIONS (Moved from previous script)
# ==============================================================================

def call_openai_for_json(system_prompt: str, user_prompt: str, temp: float = 0.7) -> dict | None:
    """Helper function to make an OpenAI call and parse JSON response."""
    try:
        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=temp,
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"ERROR: OpenAI call failed: {e}")
        print(f"System Prompt: {system_prompt}\nUser Prompt (truncated): {user_prompt[:500]}...")
        return None

def create_base_slide_dict(course_title: str, module_title: str, topic: str, screen_type: str,
                           duration: str = "01:00", bloom_level: str = "N/A") -> dict:
    """Creates a base dictionary for a new storyboard slide."""
    return {
        "Course Title": course_title,
        "Module Title": module_title,
        "Topic": topic,
        "Screen-type": screen_type,
        "Bloom's Level": bloom_level,
        "Duration (min)": duration,
        "Source Images (base64)": [],
        "Source Tables": [], # <-- NEW FIELD
        "Learning Objectives": "",
        "On-screen text": "",
        "Narration": "",
        "On-screen Recommendations": "",
        "Developer Notes": ""
    }

# DATA INGESTION AND HELPERS (Upgraded for Multimodal Content)
# ==============================================================================

def extract_rich_content_from_cell(cell) -> list[dict]:
    """
    Extracts all content from a table cell. When it encounters a nested table, it creates TWO
    representations: a plain-text version for the AI to analyze, and a formatted Markdown
    version that is stored for later but NOT sent to the AI.
    """
    content_parts = []

    # Process paragraphs for text and images (this part is unchanged)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            r_xml = run._r.xml
            if 'pic:pic' in r_xml:
                # ... (image extraction logic remains the same)
                embed_id = None
                for blip in run._r.xpath('.//a:blip'):
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')

                if embed_id:
                    try:
                        image_part = cell.part.related_parts[embed_id]
                        image_bytes = image_part.blob
                        b64_string = base64.b64encode(image_bytes).decode('utf-8')
                        mime_type = image_part.content_type
                        if not mime_type:
                            if b64_string.startswith('/9j/'):
                                mime_type = 'image/jpeg'
                            elif b64_string.startswith('iVBORw0KGgo'):
                                mime_type = 'image/png'
                            elif b64_string.startswith('R0lGODlh'):
                                mime_type = 'image/gif'
                            else:
                                mime_type = 'application/octet-stream'

                        content_parts.append({
                            "type": "image_base64",
                            "content": b64_string,
                            "mime_type": mime_type
                        })
                    except KeyError:
                        content_parts.append(
                            {"type": "text", "content": "\n[Image Found, but could not be extracted.]\n"})

        if paragraph.text.strip():
            content_parts.append({"type": "text", "content": paragraph.text})

    # --- NEW, SMARTER TABLE PROCESSING LOGIC ---
    for table in cell.tables:
        # 1. Create the plain-text representation for the AI
        plain_text_for_ai = ["[Table Data Follows]:"]
        for row in table.rows:
            row_text = " | ".join([c.text.strip() for c in row.cells])
            plain_text_for_ai.append(row_text)

        # Add this version to be sent to the AI
        content_parts.append({"type": "text", "content": "\n".join(plain_text_for_ai)})

        # 2. Create the formatted Markdown representation for the course designer
        markdown_table = ["\n[Nested Table Detected]:"]
        header = [h.text.strip() for h in table.rows[0].cells]
        markdown_table.append(f"| {' | '.join(header)} |")
        markdown_table.append(f"|{'|'.join(['---'] * len(header))}|")
        for row in table.rows[1:]:
            row_data = [c.text.strip() for c in row.cells]
            markdown_table.append(f"| {' | '.join(row_data)} |")

        # Add this version with a SPECIAL TYPE that we will filter out later
        content_parts.append({
            "type": "formatted_table_markdown",
            "content": "\n".join(markdown_table)
        })

    return content_parts


def parse_duration_string(duration_str: str) -> float:
    """
    Parses duration strings like '27 min 20 sec', '4 min 53 sec', or '3 min'
    into total minutes (float). Handles empty or non-matching strings by returning 0.0.
    """
    total_minutes = 0.0
    duration_str = duration_str.strip().lower()

    if not duration_str:
        return 0.0

    minutes_match = re.search(r'(\d+)\s*min', duration_str)
    seconds_match = re.search(r'(\d+)\s*sec', duration_str)

    if minutes_match:
        total_minutes += int(minutes_match.group(1))
    if seconds_match:
        total_minutes += int(seconds_match.group(1)) / 60.0
    return total_minutes


def read_user_context_prompt_from_word(file_path: str) -> str | None:
    """
    Reads the user_context_prompt from the Word document. It now looks for a
    heading paragraph containing 'User Prompt' and reads the content of the
    paragraph immediately following it.
    """
    if not os.path.exists(file_path):
        print(f"WARNING: File '{file_path}' not found, cannot read context prompt.")
        return None
    document = docx.Document(file_path)

    found_prompt_heading = False  # This will act as a state flag

    # Iterate through all elements to stop before the first table
    for element in document.element.body:
        # Stop searching once a table is encountered
        if element.tag.endswith('tbl'):
            break

        # We are only interested in paragraphs
        if element.tag.endswith('p'):
            paragraph = docx.text.paragraph.Paragraph(element, document)
            text = paragraph.text.strip()

            # If the flag was set in the previous iteration, this paragraph is our prompt
            if found_prompt_heading:
                user_prompt = text

                # Check if the extracted prompt is meaningful or just a placeholder
                if user_prompt and "no user prompt provided" not in user_prompt.lower():
                    print(f"    - Found User Prompt: '{user_prompt[:75]}...'")
                    return user_prompt  # Success! We found the prompt.
                else:
                    # We found the section but it was empty or had a placeholder, so we stop.
                    print("    - 'User Prompt' section found, but it contains placeholder text or is empty.")
                    return None

            # Check if the current paragraph is the heading we're looking for
            if text.lower() == 'user prompt':
                found_prompt_heading = True  # Set the flag for the next iteration

    print("    - 'User Prompt' heading not found in the document before the first table.")
    return None

def read_rows_from_word_table(file_path: str) -> list[dict]:
    """
    Reads rows from the first table of a Word document, adapting to the new column headers
    and extracting rich content.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File '{file_path}' not found.")
    document = docx.Document(file_path)
    if not document.tables:
        print("WARNING: No tables found in the document.")
        return []

    table = document.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]

    # Define expected new headers
    expected_headers = ["Chapter", "Topic", "Subtopic", "Full Page Content", "Duration (Mins)"]
    header_indices = {}
    for h in expected_headers:
        if h in headers:
            header_indices[h] = headers.index(h)
        else:
            raise ValueError(f"Expected column '{h}' not found in Word table headers. Found: {headers}")

    all_rows = []

    for row in table.rows[1:]:
        row_dict = {}

        row_dict["Chapter"] = extract_rich_content_from_cell(row.cells[header_indices["Chapter"]])
        row_dict["Topic_Outline"] = extract_rich_content_from_cell(row.cells[header_indices["Topic"]])
        row_dict["Subtopic_Outline"] = extract_rich_content_from_cell(row.cells[header_indices["Subtopic"]])
        row_dict["Full Page Content"] = extract_rich_content_from_cell(
            row.cells[header_indices["Full Page Content"]])

        duration_cell_content = extract_rich_content_from_cell(
            row.cells[header_indices["Duration (Mins)"]])
        duration_text = next((item['content'] for item in duration_cell_content if item['type'] == 'text'), "0")
        row_dict['Duration (min)'] = parse_duration_string(duration_text)

        all_rows.append(row_dict)

    return all_rows


def build_multimodal_payload(instructions: str, rich_content_parts: list[dict]) -> list[dict]:
    """
    Builds the final list for the 'content' key in a multimodal API call.
    This version is updated to filter out the 'formatted_table_markdown' type.
    """
    payload = [{"type": "text", "text": instructions}]
    for part in rich_content_parts:
        if part["type"] == "text":
            payload.append({"type": "text", "text": part["content"]})
        elif part["type"] == "image_base64":
            data_url = f"data:{part['mime_type']};base64,{part['content']}"
            payload.append({"type": "image_url", "image_url": {"url": data_url}})
        # NEW: This condition explicitly skips our special type, so it's not sent to the AI
        elif part["type"] == "formatted_table_markdown":
            continue

    return payload

def get_bloom_level_from_user() -> int:
    while True:
        user_input = input("\nEnter the desired Bloom's Taxonomy level (1=Remember, 2=Understand, 3=Apply): ")
        if user_input.isdigit() and int(user_input) in [1, 2, 3]: return int(user_input)
        print("Error: Invalid input. Please enter 1, 2, or 3.")


# NEW HELPER FUNCTION
def _get_title_from_row(row_data: dict, key: str, default_value: str) -> str:
    """Helper to extract text content from a rich content list, with fallback, cleaning 'N/A' values."""
    title = next((p['content'] for p in row_data.get(key, []) if p['type'] == 'text'), default_value)
    title = title.strip()
    return title if title and title.lower() != 'n/a' else default_value

# ==============================================================================
# NEW LOGIC FOR SEMANTIC GROUPING
# ==============================================================================

def generate_group_learning_objective(client: AzureOpenAI, deployment_name: str, group_content_text: str,
                                      bloom_level_int: int, compliance_settings: dict) -> str:
    """Generates a single LO for a group of text."""
    global ENGLISH_VARIANT
    english_variant = "UK English" if compliance_settings and getattr(compliance_settings, "ukEnglish", False) else "US English"
    english_instruction = get_english_instruction(english_variant)

    bloom_level_map = {1: "Remember", 2: "Understand", 3: "Apply"}
    bloom_level_str = bloom_level_map.get(bloom_level_int, "Understand")
    approved_verbs = BLOOM_VERBS.get(bloom_level_int, [])
    verb_instruction = ""
    if approved_verbs:
        verb_instruction = f"The objective MUST begin with one of these approved verbs for '{bloom_level_str}': `{', '.join(approved_verbs)}`."

    system_prompt = "You are an expert instructional designer."
    user_prompt = f"""
Synthesize the following course content into **ONE** comprehensive learning objective.

**Content:**
{group_content_text[:6000]}...

**Instructions:**
1. Short phrase, NOT a full sentence.
2. Start directly with an action verb.
3. {verb_instruction}

Language Style:
{english_instruction}

**Output:** Just the learning objective string.
"""
    try:
        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name, messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"      - Error generating group LO: {e}")
        return "Understand the key concepts presented in this section."


def get_semantic_groups_from_llm(client: AzureOpenAI, deployment_name: str, slide_summaries: list[dict]) -> list[list[int]]:
    """
    Sends a list of slide summaries (ID + Topic + Snippet) to LLM.
    Asks LLM to group them semantically into chunks of 3-7 slides.
    Returns a list of lists of indices: [[0,1,2,3], [4,5,6], ...]
    """
    # Prepare text for prompt
    content_block = ""
    for s in slide_summaries:
        # We limit the preview to 200 chars to save tokens, just enough for semantic check
        content_block += f"Slide_ID: {s['id']} | Topic: {s['topic']} | Content: {s['content_preview'][:200]}...\n"

    system_prompt = "You are an expert instructional designer."
    user_prompt = f"""
I have a list of sequential course slides below.
Please group these slides into logical "Learning Blocks" based on **content relevancy and similarity**.

**Rules:**
1. Group slides that discuss the same concept together.
2. Groups should ideally contain between **3 and 7 slides**.
3. **Do not reorder** the slides. Groups must be sequential (e.g., 1,2,3,4 not 1,5,9).
4. Every Slide_ID must belong to exactly one group.
5. If a single slide covers a totally different topic, it can be a group of 1 (but avoid if possible).

**Input Data:**
{content_block}

**Output Format:**
You MUST return valid JSON containing a list of groups, where each group is a list of Slide_IDs.
Example: {{ "groups": [[0, 1, 2, 3], [4, 5, 6], [7, 8, 9, 10]] }}
"""

    try:
        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        result = json.loads(response.choices[0].message.content)
        return result.get("groups", [])
    except Exception as e:
        print(f"    - Error during semantic grouping: {e}. Falling back to default chunks.")
        # Fallback: simple chunks of 5
        indices = [s['id'] for s in slide_summaries]
        return [indices[i:i + 5] for i in range(0, len(indices), 5)]


def pre_calculate_learning_objectives(all_rows: list[dict], client: AzureOpenAI, deployment_name: str,
                                      bloom_level_int: int, compliance_settings: dict, course_title: str) -> dict:
    """
    Iterates through all rows, collects them by Chapter.
    Sends chapters to LLM for semantic grouping.
    Generates LOs for those groups.
    """
    print("\n--- Pre-calculating Shared Learning Objectives (Semantic Grouping) ---")
    lo_map = {}
    
    # Buffer to hold current chapter's content rows
    chapter_buffer = [] # List of dicts: {'id': real_index, 'topic': str, 'content': str}
    current_chapter_sig = None

    def process_chapter_buffer(buffer):
        if not buffer: return
        print(f"    - Semantically grouping {len(buffer)} slides for Chapter: '{current_chapter_sig}'...")
        
        # 1. Ask LLM to define the groups
        semantic_groups = get_semantic_groups_from_llm(client, deployment_name, buffer)
        
        # 2. For each group, generate LO
        for group_indices in semantic_groups:
            # Reconstruct the full text for this group to generate the LO
            group_full_text = []
            valid_indices = []
            
            for idx in group_indices:
                # Find the matching content from our buffer
                match = next((item for item in buffer if item['id'] == idx), None)
                if match:
                    group_full_text.append(match['content'])
                    valid_indices.append(idx)
            
            if valid_indices:
                combined_text = "\n\n".join(group_full_text)
                print(f"      > Generating LO for group of {len(valid_indices)} slides...")
                shared_lo = generate_group_learning_objective(client, deployment_name, combined_text, bloom_level_int, compliance_settings)
                
                # 3. Update the map
                for idx in valid_indices:
                    lo_map[idx] = shared_lo

    # Iterate main rows
    for i, row in enumerate(all_rows):
        topic = _get_title_from_row(row, "Topic_Outline", "General")
        chapter = _get_title_from_row(row, "Chapter", course_title)
        
        # Extract text
        content_parts = row.get("Full Page Content", [])
        text_content = "\n".join([p['content'] for p in content_parts if p['type'] == 'text'])
        
        # Check if content row
        is_content_row = bool(text_content.strip()) and "knowledge check" not in topic.lower() and "overview" not in topic.lower()

        # Check for chapter change
        if current_chapter_sig is not None and chapter != current_chapter_sig:
            process_chapter_buffer(chapter_buffer)
            chapter_buffer = []

        current_chapter_sig = chapter

        if is_content_row:
            chapter_buffer.append({
                'id': i,
                'topic': topic,
                'content': text_content,
                'content_preview': text_content # We slice inside the helper function
            })

    # Flush final chapter
    process_chapter_buffer(chapter_buffer)

    print("--- Semantic LO Pre-calculation Complete ---\n")
    return lo_map


# ==============================================================================
# AI AND PROMPT LOGIC FOR CORE CONTENT SLIDES
# ==============================================================================

def create_screen_type_selection_prompt(recently_used: list[str], user_context_prompt: str | None) -> str:
    screen_type_criteria = """
### Static Screen
**Description:** For straightforward, non-interactive information.
**Use When:** Presenting introductions, objectives, facts, definitions, policies, or summaries.

### Animated Screen
**Description:** To enhance comprehension of dynamic concepts.
**Use When:** Explaining a complex process, showing change over time, visualizing abstract concepts, or simulating real-world tasks.

### Flipcard Screen
**Description:** To promote active recall and reinforce paired information.
**Use When:** For question/answer, term/definition, cause/effect, or vocabulary training.

### Tabs Screen
**Description:** To present related but distinct categories of non-sequential information.
**Use When:** Comparing items (e.g., Pros/Cons), showing different facets of a topic (e.g., Features/Specs/Pricing), or avoiding long scrolling pages.

### Accordion Screen
**Description:** To organize layered, non-sequential information vertically.
**Use When:** For FAQs, guidelines, or chunked topics where users can expand sections of interest.

### Carousel Screen
**Description:** To present a series of related, sequential items visually.
**Use When:** For steps in a process, timeline events, case studies, or a tour of features.
"""
    avoidance_instruction = ""
    if recently_used:
        avoid_list = "\n".join([f"- {item}" for item in recently_used])
        avoidance_instruction = f"""
**AVOID REPETITION:**
To create a more engaging course, please AVOID using the following recently used screen types if a suitable alternative from the criteria exists:
{avoid_list}
"""
    course_context_instruction = ""
    if user_context_prompt:
        course_context_instruction = f"""
**OVERALL COURSE CONTEXT/SCENARIO (for guiding tone and relevance):**
{user_context_prompt}
"""

    prompt = f"""
You are an expert e-learning instructional designer. Your task is to analyze the following source content (which may include text and images) and select the single best screen type to present it, based on the detailed criteria provided.
{course_context_instruction}
{avoidance_instruction}
**DETAILED SCREEN TYPE CRITERIA:**
{screen_type_criteria}
**REASONING PROCESS:**
1.  Analyze the `SOURCE CONTENT`. What is its primary purpose?
2.  Review the `AVOID REPETITION` list.
3.  Compare the content's purpose against the `DETAILED SCREEN TYPE CRITERIA`, giving preference to types NOT on the avoidance list.
4.  Select the **single most appropriate** screen type.
**OUTPUT INSTRUCTIONS:**
Your output **MUST** be a single, raw JSON object with one key, "screen_type".
"""
    return prompt


def select_screen_type(client: AzureOpenAI, deployment_name: str, rich_content_parts: list[dict],
                       recently_used: list[str], user_context_prompt: str | None) -> str | None:
    print("    - Step 1: Selecting screen type (multimodal)...")
    if not rich_content_parts:
        print("      - SKIPPING: Content is empty.")
        return None

    instructions = create_screen_type_selection_prompt(recently_used, user_context_prompt)
    payload = build_multimodal_payload(instructions, rich_content_parts)

    try:
        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name,
            messages=[{"role": "user", "content": payload}],
            temperature=0.2,
            response_format={"type": "json_object"}
        )
        result = json.loads(response.choices[0].message.content)
        screen_type = result.get("screen_type")
        if screen_type:
            print(f"      - Success. Selected screen type: '{screen_type}'")
            return screen_type
        return None
    except Exception as e:
        print(f"      - FAILED: An error occurred during screen type selection: {e}")
        return None


def create_storyboard_generation_prompt(row_data: dict, screen_type: str, bloom_level_int: int,
                                       course_title: str, recently_used_verbs: list[str],
                                       user_context_prompt: str | None,
                                       compliance_settings: dict = None,
                                       forced_learning_objective: str = None) -> str:
    """
    Creates the comprehensive, multimodal prompt for generating the full storyboard content for a single slide.
    Now includes logic to enforce a pre-calculated learning objective if provided.
    """
    print("Compliance settings received:", compliance_settings)

    # Determine English variant
    global ENGLISH_VARIANT
    ENGLISH_VARIANT = "UK English" if compliance_settings and getattr(compliance_settings, "ukEnglish", False) else "US English"

    print("Selected variant:", ENGLISH_VARIANT)

    # Build the prompt with English variant specification
    english_instruction = get_english_instruction(ENGLISH_VARIANT)

    # --- 1. Extract and Format Basic Information (Unchanged) ---
    module_title = _get_title_from_row(row_data, "Chapter", course_title)
    slide_topic = _get_title_from_row(row_data, "Subtopic_Outline",
                                     _get_title_from_row(row_data, "Topic_Outline", "General Information"))
    duration = row_data.get("Duration (min)", 0.0)
    target_word_count = int(duration * 150)
    minimum_word_count = int(target_word_count * 0.9)
    bloom_level_map = {1: "Remember", 2: "Understand", 3: "Apply"}
    bloom_level_str = bloom_level_map.get(bloom_level_int, "Understand")
    bloom_instructions = {
        "Remember": "Focus on **Remembering**: State facts, definitions, and lists directly.",
        "Understand": "Focus on **Understanding**: Explain the 'why', summarise, and describe relationships.",
        "Apply": "Focus on **Applying**: Frame narration around a practical scenario or how to use the information."
    }
    selected_bloom_instruction = bloom_instructions[bloom_level_str]

    # Build context instruction
    course_context_instruction = ""
    if user_context_prompt:
        course_context_instruction = f"""
**COURSE CONTEXT:**
{user_context_prompt}
Follow this :
{english_instruction}
Use this context to inform your content creation approach and ensure relevance to the course objectives.
"""

    # --- 2. Build Instructions for Learning Objectives (MODIFIED) ---
    
    if forced_learning_objective:
        # If we have a pre-calculated LO from grouping logic, force the AI to use it.
        learning_objective_instruction = f"""
**LEARNING OBJECTIVE (MANDATORY):**
You **MUST** use the following Learning Objective exactly as provided below. Do not create a new one.
**Objective:** "{forced_learning_objective}"
"""
    else:
        # Fallback to per-slide generation if no shared LO exists (e.g. unique slides)
        approved_verbs = BLOOM_VERBS.get(bloom_level_int, [])
        verb_instruction_parts = []
        if approved_verbs:
            verb_instruction_parts.append(
                f"The objective MUST begin with one of the following approved verbs for the '{bloom_level_str}' level: `{', '.join(approved_verbs)}`.")
            if recently_used_verbs:
                avoid_list = ", ".join(f"'{v}'" for v in recently_used_verbs)
                verb_instruction_parts.append(
                    f"To ensure variety, please AVOID starting with these recently used verbs if a suitable alternative exists: {avoid_list}.")
        verb_instruction = " ".join(verb_instruction_parts)

        learning_objective_instruction = f"""
**LEARNING OBJECTIVE INSTRUCTIONS (VERY IMPORTANT):**
1.  **Format:** The objective must be a short phrase, NOT a full sentence.
2.  **Verb First:** The objective MUST start directly with an action verb.
3.  **Verb Choice:** {verb_instruction}
4.  **Scope:** The objective should be relevant to this specific slide's content, but phrased like a high-level goal.
"""

    # --- 3. Build Instructions for Interactive Screen Types (MODIFIED) ---
    interactive_formatting_rules = {
        "Accordion Screen": {
            "on_screen_text": "The very first line of the on-screen text MUST be 'Click each item to learn more.'. Structure the rest of the text with a header for each accordion item, followed by its content.\n\nExample:\nClick each item to learn more.\n\nAccordion Header 1\nContent for the first accordion.\n\nAccordion Header 2\nContent for the second accordion.",
            "narration": "Structure the narration with an introduction, then separate narrations for each accordion click. Example:\n- [Narration for when user clicks Accordion 1]...\n- [Narration for when user clicks Accordion 2]..."
        },
        "Tabs Screen": {
            "on_screen_text": "The very first line of the on-screen text MUST be 'Click each section to explore.'. Structure the rest of the text with a title for each tab, followed by its content.\n\nExample:\nClick each section to explore.\n\nTab Title 1\nContent for the first tab.\n\nTab Title 2\nContent for the second tab.",
            "narration": "Structure the narration with an introduction, then separate narrations for each tab. Example:\n- [Narration for when user clicks Tab 1]...\n- [Narration for when user clicks Tab 2]..."
        },
        "Flipcard Screen": {
            "on_screen_text": "Structure the text clearly showing the Front and Back of each card. Use 'Front:' and 'Back:' as labels. Example:\nFront: Key Term\n\nBack: Definition of the key term.",
            "narration": "Structure narration for the front and back. Example:\n- [Narration for the front]...\n- [Narration for the back]..."
        },
        "Carousel Screen": {
            "on_screen_text": "The very first line of the on-screen text MUST be 'Click the Continue and the Previous buttons to learn more.'. Structure the rest of the text for each panel in the carousel.\n\nExample:\nClick the Continue and the Previous buttons to learn more.\n\nPanel 1: Title\nContent for the first panel.\n\nPanel 2: Title\nContent for the second panel.",
            "narration": "Structure narration with an introduction, then separate narrations for each panel. Example:\n- [Narration for Panel 1]...\n- [Narration for Panel 2]..."
        }
    }

    if screen_type in interactive_formatting_rules:
        on_screen_text_task = interactive_formatting_rules[screen_type]["on_screen_text"]
        narration_task = interactive_formatting_rules[screen_type]["narration"]
    else:
        on_screen_text_task = f"Create concise on-screen text suitable for a '{screen_type}'."
        narration_task = f"Write a detailed, single-paragraph narration."

    # --- 4. Build Contextual Information Blocks (Unchanged) ---
    course_context_instruction = ""
    if user_context_prompt:
        course_context_instruction = f"""**OVERALL COURSE CONTEXT/SCENARIO (IMPORTANT for tone, style, and scenario):**\n{user_context_prompt}"""
   
    

    # --- 5. Dynamically Generate Recommendation Instructions (Unchanged) ---
    rich_content_parts = row_data.get("Full Page Content", [])
    has_source_images = any(part['type'] == 'image_base64' for part in rich_content_parts)
    if has_source_images:
        recommendation_instruction = "Suggest relevant visuals. **Crucially, you MUST specifically reference the provided source images** and explain how they should be used."
    else:
        recommendation_instruction = "Suggest relevant non-photographic visuals like icons, charts, or diagrams. **DO NOT suggest stock photos or mention source images, as none were provided.**"

    # --- 6. Assemble the Final, High-Impact Prompt (Unchanged) ---
    prompt = f"""
You are an expert e-learning storyboard designer. Your function is to transform provided source material into a structured storyboard format without losing any information.

**THE PLAN FOR THIS SLIDE:**
- Course Title: "{course_title}"
- Module Title: "{module_title}"
- Slide Topic: "{slide_topic}"
- **Screen Type to Use:** "{screen_type}"

{course_context_instruction}


---
**CORE DIRECTIVE (NON-NEGOTIABLE):**
Your single most important responsibility is to ensure that **EVERY piece of information** from the `SOURCE CONTENT` is fully and accurately represented in the generated "Narration" and "On-screen text".
- **No piece of source information must be left behind.**
---

**SECONDARY TASKS & GUIDELINES:**
- **Slide Plan Adherence:** Generate the output using the required format for a "{screen_type}".
  - **On-screen Text Task:** {on_screen_text_task}
  - **Narration Task:** {narration_task}
- **Learning Objective:** {learning_objective_instruction}
- **Guideline Adherence:**
  - **Narration Style:** **DO NOT** use introductory phrases like "Welcome to this section...". Dive directly into the topic.
  - **Narration Length:** The word count ({minimum_word_count}-{target_word_count} words) is a loose guideline. **Covering all content is more important.**
  - **Bloom's Level ({bloom_level_str}):** Write the narration in this style: {selected_bloom_instruction}
  - **Language:** Use British English (UK).

---
**OUTPUT INSTRUCTIONS:**
Your output must be a single, raw JSON object.The value of "On-screen text" and "Narration" MUST be a SINGLE string.

**JSON OUTPUT FORMAT:**
{{
    "Learning Objectives": "The learning objective string.",
    "On-screen text": "A SINGLE string. For interactive slides, combine all headers and content into this one string, using newline characters ('\n') to separate sections exactly as shown in the examples."
    "Narration": "A SINGLE string. For interactive slides, combine the introductory narration and the narration for each interactive element into this one string, using newline characters ('\n') to separate them."
    "On-screen Recommendations": "{recommendation_instruction}",
    "Developer Notes": "Provide notes for the developer, specifically for implementing a '{screen_type}'."
}}
"""
    return prompt


def generate_storyboard_content(client: AzureOpenAI, deployment_name: str, row_data: dict, screen_type: str,
                                bloom_level: int, course_title: str, recently_used_verbs: list[str],
                                user_context_prompt: str | None , compliance_settings, forced_learning_objective: str = None ) -> dict | None:
    print("    - Step 2: Generating full storyboard content (multimodal)...")

    instructions = create_storyboard_generation_prompt(row_data, screen_type, bloom_level, course_title,
                                                       recently_used_verbs, user_context_prompt,
                                                        compliance_settings, forced_learning_objective)
    rich_content_parts = row_data.get("Full Page Content", [])
    if not rich_content_parts: return None

    # This payload will intelligently filter out the 'formatted_table_markdown' type
    payload = build_multimodal_payload(instructions, rich_content_parts)

    try:
        global ENGLISH_VARIANT
        english_instruction = get_english_instruction(ENGLISH_VARIANT)

        print(english_instruction)

        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name,
            messages=[
                {"role": "system",
                 "content": f"You are an e-learning designer who creates a single slide's storyboard content in a valid JSON format based on a pre-defined plan and multimodal source content.The english used should follow the rules mentioed in {english_instruction}."},
                {"role": "user", "content": payload}
            ],
            temperature=0.5,
            response_format={"type": "json_object"}
        )

        storyboard = json.loads(response.choices[0].message.content)

        # Populate the storyboard with metadata
        module_title = _get_title_from_row(row_data, "Chapter", course_title)
        slide_topic = _get_title_from_row(row_data, "Subtopic_Outline",
                                          _get_title_from_row(row_data, "Topic_Outline", "General Information"))

        storyboard['Course Title'] = course_title
        storyboard['Module Title'] = module_title
        storyboard['Topic'] = slide_topic
        storyboard['Screen-type'] = screen_type

        bloom_level_map = {1: "Remember", 2: "Understand", 3: "Apply"}
        storyboard['Bloom\'s Level'] = bloom_level_map.get(bloom_level, "Unknown")

        # Clean placeholders before word count calculation for accurate duration
        narration_text = storyboard.get("Narration", "")
        cleaned_narration_for_wc = re.sub(r'-\s*\[Narration for .*?\]\.\.\.\s*', '', narration_text,
                                          flags=re.IGNORECASE).strip()
        word_count = len(cleaned_narration_for_wc.split())

        calculated_duration_in_minutes = word_count / 150.0
        total_seconds = int(calculated_duration_in_minutes * 60)
        minutes, seconds = divmod(total_seconds, 60)
        formatted_duration = f"{minutes:02d}:{seconds:02d}"
        storyboard['Duration (min)'] = formatted_duration

        # Extract and add source images to the storyboard
        image_b64_list = [
            {"content": part['content'], "mime_type": part['mime_type']}
            for part in rich_content_parts if part['type'] == 'image_base64'
        ]
        storyboard['Source Images (base64)'] = image_b64_list

        # Populate the "Source Tables" field
        formatted_tables_to_add = []
        # Find all the stored Markdown tables from the original rich content
        for part in rich_content_parts:
            if part.get('type') == 'formatted_table_markdown':
                # We store just the string content of the table
                formatted_tables_to_add.append(part['content'])

        # If any tables were found, add them to the dedicated field
        if formatted_tables_to_add:
            storyboard["Source Tables"] = formatted_tables_to_add
        else:
            storyboard["Source Tables"] = []  # Ensure it's an empty list if no tables are found

        print("      - Success. Storyboard generated.")
        return storyboard

    except Exception as e:
        print(f"      - FAILED: An error occurred during storyboard generation: {e}")
        return None


def generate_knowledge_check(client: AzureOpenAI, deployment_name: str, chapter_narration_list: list[str],
                             course_title: str, module_title: str, user_context_prompt: str | None,
                             bloom_level_int: int, compliance_settings: dict = None) -> list[dict] | None:
    print(f"    - Generating Knowledge Check content for chapter: '{module_title}' with Bloom's Level {bloom_level_int}...")

    global ENGLISH_VARIANT
    english_instruction = get_english_instruction(ENGLISH_VARIANT)

    combined_narration = "\n\n".join(filter(None, chapter_narration_list)).strip()

    if not combined_narration:
        print("      - SKIPPING Knowledge Check: No narration accumulated for this chapter.")
        return None

    course_context_instruction = ""
    if user_context_prompt:
        course_context_instruction = f"""
**OVERALL COURSE CONTEXT/SCENARIO:**
{user_context_prompt}

"""

    # Define completely separate prompts for each Bloom's level
    bloom_level_prompts = {
        1: f"""
You are an expert e-learning instructional designer. Your task is to create a 1-question knowledge check for Bloom's Level 1 (Remember).

{course_context_instruction}

**BLOOM'S LEVEL 1 (REMEMBER) REQUIREMENTS:**
- Create a question that tests basic recall and recognition of facts, terms, and concepts
- The question MUST be either True/False OR Fill in the Blank
- Focus on factual recall, definitions, dates, names, and basic concepts from the narration
- The question should be straightforward and test memory of key information

**SPECIFIC INSTRUCTIONS:**
1. Create **exactly 1 question**.
2. Use either a True/False or Fill in the Blank question
3. Do NOT use multiple choice, drag and drop, or other complex question types
4. Question should be simple and test direct recall of information
5. The question should be simple and test direct recall of information

**CHAPTER NARRATION:**
{combined_narration}

**OUTPUT INSTRUCTIONS:**
Your output MUST be a single, raw JSON object with exactly 1 question in the "Questions" list.

**ALLOWED QUESTION TYPES FOR LEVEL 1:**

**True/False:**
{{
  "question_type": "True/False",
  "question": "The sky is green.",
  "instruction": "Select either True or False and click submit.",
  "options": {{ "A": "True", "B": "False" }},
  "correct_answer": "B",
  "feedback": {{ "A": "Incorrect, the sky is blue due to Rayleigh scattering.", "B": "Correct, the sky is not green." }}
}}

**Fill in the Blank:**
{{
  "question_type": "Fill in the Blank",
  "question": "The capital of France is ___.",
  "instruction": "Type your answer in the blank and click submit.",
  "correct_answer": "Paris",
  "correct_answer_feedback": "Paris is the capital and most populous city of France."
}}

**Example of Final JSON Structure:**
{{
    "Narration": "Let's test your recall of key concepts...",
    "Questions": [
        {{ ... Single True/False OR Fill in the Blank question ... }}
    ]
}}
Follow this :
{english_instruction}
""",

        2: f"""
You are an expert e-learning instructional designer. Your task is to create a 1-question knowledge check for Bloom's Level 2 (Understand).

{course_context_instruction}

**BLOOM'S LEVEL 2 (UNDERSTAND) REQUIREMENTS:**
- Create a question that tests comprehension, interpretation, and explanation of concepts
- The question MUST be Multiple Choice (either single select or multiple select)
- Focus on understanding relationships, explaining concepts, and interpreting information
- The question should require learners to demonstrate they understand the meaning and implications of concepts

**SPECIFIC INSTRUCTIONS:**
1. Create **exactly 1 question**.
2. The question must be Multiple Choice (MCSS or MCMS)
3. Use either single select or multiple-select question.
4. Do NOT use True/False, Fill in the Blank, or Drag and Drop questions
5. The question should test understanding, not just recall

**CHAPTER NARRATION:**
{combined_narration}

**OUTPUT INSTRUCTIONS:**
Your output MUST be a single, raw JSON object with exactly 1 question in the "Questions" list.

**ALLOWED QUESTION TYPES FOR LEVEL 2:**

**MCSS (Multiple Choice Single Select):**
{{
  "question_type": "MCSS",
  "question": "What is the primary purpose of project documentation?",
  "instruction": "Select the correct option and click submit.",
  "options": {{ 
    "A": "To create busy work for team members",
    "B": "To track project progress and communicate requirements",
    "C": "To satisfy management requirements only",
    "D": "To replace verbal communication entirely"
  }},
  "correct_answer": "B",
  "feedback": {{
    "A": "Incorrect, documentation serves important communication purposes.",
    "B": "Correct, documentation helps track progress and communicate requirements effectively.",
    "C": "Incorrect, documentation benefits the entire project team, not just management.",
    "D": "Incorrect, documentation complements but doesn't replace verbal communication."
  }}
}}

**MCMS (Multiple Choice Multiple Select):**
{{
  "question_type": "MCMS",
  "question": "Which of the following are benefits of effective risk management? (Choose 2)",
  "instruction": "Select all correct options and click submit.",
  "options": {{ 
    "A": "Eliminates all project risks",
    "B": "Helps anticipate potential problems",
    "C": "Reduces project uncertainty",
    "D": "Guarantees project success"
  }},
  "correct_answer": ["B", "C"],
  "feedback": {{
    "A": "Incorrect, risk management helps manage but cannot eliminate all risks.",
    "B": "Correct, risk management helps anticipate potential problems before they occur.",
    "C": "Correct, by identifying and planning for risks, uncertainty is reduced.",
    "D": "Incorrect, risk management improves success chances but doesn't guarantee it."
  }}
}}

**Example of Final JSON Structure:**
{{
    "Narration": "Let's check your understanding of the concepts...",
    "Questions": [
        {{ ... Single MCSS OR MCMS question ... }}
    ]
}}
Follow this :
{english_instruction}
""",

        3: f"""
You are an expert e-learning instructional designer. Your task is to create a 1-question knowledge check for Bloom's Level 3 (Apply).

{course_context_instruction}

**BLOOM'S LEVEL 3 (APPLY) REQUIREMENTS:**
- Create a question that tests application of knowledge in new situations and problem-solving
- The question MUST be scenario-based, presenting a real-world situation
- Question types can include MCSS, MCMS, True/False, or Fill in the Blank, but it must be scenario-based
- Focus on applying concepts to practical situations, making decisions, and solving problems

**SPECIFIC INSTRUCTIONS:**
1. Create **exactly 1 question**.
2. The question must present a realistic scenario or case study
3. The question should require learners to apply concepts to solve problems or make decisions
4. Use any question type (MCSS, MCMS, True/False, Fill in the Blank) but ensure it is scenario-driven
5. Scenario should be realistic and relevant to the course content

**CHAPTER NARRATION:**
{combined_narration}

**OUTPUT INSTRUCTIONS:**
Your output MUST be a single, raw JSON object with exactly 1 question in the "Questions" list.

**SCENARIO-BASED QUESTION EXAMPLES FOR LEVEL 3:**

**Scenario-based MCSS:**
{{
  "question_type": "MCSS",
  "question": "Scenario: You are managing a project that is behind schedule. The client is demanding an earlier delivery date. What should be your FIRST action?",
  "instruction": "Select the most appropriate first action and click submit.",
  "options": {{ 
    "A": "Immediately commit to the new deadline to please the client",
    "B": "Analyze the impact on resources and timeline before responding",
    "C": "Blame team members for the delay",
    "D": "Ignore the client's request and stick to the original plan"
  }},
  "correct_answer": "B",
  "feedback": {{
    "A": "Incorrect, committing without analysis can lead to further problems.",
    "B": "Correct, analyzing impact first allows for informed decision-making.",
    "C": "Incorrect, blaming team members is unprofessional and counterproductive.",
    "D": "Incorrect, ignoring client requests can damage the relationship."
  }}
}}

**Scenario-based True/False:**
{{
  "question_type": "True/False",
  "question": "Scenario: A team member consistently misses deadlines but produces high-quality work. As a manager, you should immediately terminate their employment. True or False?",
  "instruction": "Select either True or False and click submit.",
  "options": {{ "A": "True", "B": "False" }},
  "correct_answer": "B",
  "feedback": {{
    "A": "Incorrect, termination should be a last resort after other interventions.",
    "B": "Correct, you should first investigate reasons and provide support or coaching."
  }}
}}

**Scenario-based Fill in the Blank:**
{{
  "question_type": "Fill in the Blank",
  "question": "Scenario: During a team meeting, two members have a heated disagreement. As the team leader, you should first ______ to de-escalate the situation.",
  "instruction": "Type your answer in the blank and click submit.",
  "correct_answer": "facilitate calm communication",
  "correct_answer_feedback": "Correct, facilitating calm communication helps de-escalate conflicts and allows for constructive discussion of differences."
}}

**Example of Final JSON Structure:**
{{
    "Narration": "Let's apply these concepts to real-world scenarios...",
    "Questions": [
        {{ ... Single scenario-based question ... }}
    ]
}}
Follow this :
{english_instruction}
"""
    }

    # Select the appropriate prompt based on Bloom's level
    kc_prompt = bloom_level_prompts.get(bloom_level_int, bloom_level_prompts[2])

    try:
        response = call_llm_with_retry(
            _get_openai_client(),
            model=deployment_name, messages=[
                {"role": "system",
                 "content": "You are an e-learning designer specializing in creating Bloom's taxonomy-aligned assessment questions with detailed feedback."},
                {"role": "user", "content": kc_prompt}],
            temperature=0.8, response_format={"type": "json_object"}
        )
        kc_master_content = json.loads(response.choices[0].message.content)

        knowledge_check_slides = []
        if "Questions" in kc_master_content and kc_master_content["Questions"]:

            questions_list = kc_master_content["Questions"]
            print(f"      - Splitting Knowledge Check into {len(questions_list)} individual question slides...")
            for i, q_data in enumerate(questions_list):
                q_type = q_data.get("question_type", "MCSS")
                question_slide = create_base_slide_dict(
                    course_title, module_title, f"Knowledge Check - Q ({q_type})", "Interactive Quiz", "00:45"
                )

                on_screen_parts = [f"{q_data.get('question', '')}\n\nInstruction: {q_data.get('instruction', '')}"]

                # Logic to format based on question type
                if q_type in ["MCSS", "MCMS", "True/False"]:
                    options = q_data.get('options', {})
                    for key, value in options.items():
                        on_screen_parts.append(f"{key}. {value}")

                    correct = q_data.get('correct_answer', ' ')
                    correct_str = ", ".join(correct) if isinstance(correct, list) else correct
                    on_screen_parts.append(f"\nCorrect Answer(s): {correct_str}")

                    on_screen_parts.append("\n--- Feedback for All Options ---")
                    feedback = q_data.get('feedback', {})
                    for key in sorted(options.keys()):
                        on_screen_parts.append(f"Rationale for {key}: {feedback.get(key, 'N/A')}")

                elif q_type == "Drag and Drop":
                    on_screen_parts.append("\n--- Items to Match ---")
                    on_screen_parts.append("Drag Items: " + ", ".join(q_data.get("drag_items", [])))
                    on_screen_parts.append("Drop Targets: " + ", ".join(q_data.get("drop_targets", [])))
                    on_screen_parts.append("\n--- Correct Pairing ---")
                    answers = q_data.get('correct_answer', {})
                    for key, value in answers.items():
                        on_screen_parts.append(f"{key} -> {value}")
                    on_screen_parts.append(f"\nRationale: {q_data.get('correct_answer_feedback', 'N/A')}")

                elif q_type == "Fill in the Blank":
                    on_screen_parts.append(f"\nCorrect Answer: {q_data.get('correct_answer', 'N/A')}")
                    on_screen_parts.append(f"Rationale: {q_data.get('correct_answer_feedback', 'N/A')}")

                formatted_on_screen_text = "\n".join(on_screen_parts)

                question_slide.update({
                    "Learning Objectives": "N/A", "On-screen text": formatted_on_screen_text,
                    "Narration": f"Question. {q_data.get('instruction', 'Please complete the question.')}",
                    "Developer Notes": f"Implement as a {q_type} interaction."
                })
                knowledge_check_slides.append(question_slide)

        print(f"      - Success. Knowledge Check generated with {len(knowledge_check_slides)} slides for Bloom's Level {bloom_level_int}.")
        return knowledge_check_slides

    except Exception as e:
        print(f"      - FAILED: An error occurred during Knowledge Check generation for '{module_title}': {e}")
        return None
# ==============================================================================
# AI AND PROMPT LOGIC FOR INTRODUCTORY / CONCLUDING SLIDES (Moved from previous script)
# ==============================================================================

def generate_welcome_slide(course_title: str, course_context: str) -> dict:
    print("  - Generating Welcome Slide...")
    global ENGLISH_VARIANT
    english_instruction = get_english_instruction(ENGLISH_VARIANT)
   
    slide_dict = create_base_slide_dict(course_title, "Introduction", "Welcome", "Static Screen", "00:45")
    system_prompt = "You are an e-learning instructional designer creating a welcome slide for a new course. Your tone should be inviting and professional."
    user_prompt = f"""
   
Create a welcome slide for a course titled "{course_title}".
The course content broadly covers topics like: {course_context[:500]}... (summary of context to inform tone and brief mention of what's to come)

**Output MUST be a JSON object with these keys:**
{{
    "Learning Objectives": "A brief, inviting objective for the welcome.",
    "On-screen text": "A concise welcome message and course title, e.g., 'Welcome to [Course Title]'.",
    "Narration": "A friendly narration welcoming learners, setting the stage, and briefly mentioning the course's purpose and what they will gain. (Use British English)",
    "On-screen Recommendations": "Suggest relevant visuals for a welcome screen (e.g., course logo, welcoming abstract image, team photo if applicable).",
    "Developer Notes": "Provide notes for implementing the welcome slide, e.g., 'Ensure course title is prominent. Add a background image related to the course theme.'"
}}

Follow this :
{english_instruction}
"""
    result = call_openai_for_json(system_prompt, user_prompt, temp=0.5)
    if result:
        slide_dict.update(result)
        slide_dict["Learning Objectives"] = "N/A"
    return slide_dict


def generate_navigation_slide(course_title: str, course_context: str) -> dict:
    """
    Generates the Course Navigation slide using a predefined, hardcoded template.
    This function no longer calls the AI, ensuring consistency across all courses.
    """
    print("  - Generating Navigation Slide (using predefined template)...")

    # 1. Create the base dictionary for the slide.
    # The duration is set to a reasonable default for this amount of text.
    slide_dict = create_base_slide_dict(
        course_title,
        "Introduction",
        "Course Navigation",
        "Static Screen",
        "01:15"
    )

    # 2. Define the exact on-screen text and narration from your template.
    on_screen_text_template = """Take a look at each icon you will find featured in this course.

Menu
Lists each topic in this course. Select any visited topic in the Menu to view (or jump) directly to it.

Transcript
Provides the audio content of the screen.

Resources
Links to the various just-in-time resource documents referenced in the course.

Help
Offers assistance on how to navigate within the course.

Back
Navigates to the previous screen throughout the course.

Screen counter
Displays the current screen number along with the total number of screens in the course.

Forward
Navigates to the next screen throughout the course.

Audio
Increases or reduces the volume of the audio.

Play/Pause
Resumes or pauses progress on the current screen.

Progress bar
Displays the status of the screen.

Replay
Rewinds the current screen to the beginning/start.

Glossary
Provides access to a list of definitions for key terms that appear throughout the course.

Lesson Progress bar
Displays the status of the lessons"""

    narration_template = "Before you begin, please take a moment to get acquainted with the key functionalities you will find in this course. These will help you navigate through the course."

    # 3. Create a dictionary with the template content and standard recommendations.
    navigation_content = {
        "Learning Objectives": "N/A",  # As per previous requirement for generic slides.
        "On-screen text": on_screen_text_template.strip(),
        "Narration": narration_template.strip(),
        "On-screen Recommendations": "Use icons to visually represent each navigation feature (e.g., a menu icon, a help icon). Arrange the text and icons in a clean, easy-to-read layout, possibly a two-column grid or list.",
        "Developer Notes": "This is a standard navigation slide using a fixed template. Ensure the course player includes all functionalities listed. The layout should be static and informative."
    }

    # 4. Update the base slide dictionary with the template content.
    slide_dict.update(navigation_content)

    # 5. Return the completed slide dictionary.
    return slide_dict


# In storyboard_generator.py

def generate_course_objectives_slide(course_title: str, manual_objectives_text: str | None, course_context: str,
                                     unique_module_titles: list[str], bloom_level_int: int) -> dict:
    print("  - Generating Course Objectives Slide...")
    slide_dict = create_base_slide_dict(course_title, "Introduction", "Course Objectives", "Static Screen", "01:30")
    global ENGLISH_VARIANT
    english_instruction = get_english_instruction(ENGLISH_VARIANT)

    if manual_objectives_text:
        print("    - Using manually provided content for objectives slide.")
        # ... (manual objectives logic remains the same)
        objectives_content = {
            "Learning Objectives": "N/A",
            "On-screen text": manual_objectives_text,
            "Narration": manual_objectives_text,
            "On-screen Recommendations": "Display the provided objectives as a clean, bulleted list. Use icons next to each objective to enhance visual appeal.",
            "Developer Notes": "This slide's content was populated directly from the 'Course Overview and Objectives' row in the source outline."
        }
        slide_dict.update(objectives_content)
        return slide_dict
    else:
        print(f"    - No manual objectives found. Generating with AI for Bloom's Level {bloom_level_int}...")

        bloom_level_map = {1: "Remember", 2: "Understand", 3: "Apply"}
        bloom_level_str = bloom_level_map.get(bloom_level_int, "Understand")
        approved_verbs = BLOOM_VERBS.get(bloom_level_int, [])

        verb_instruction = ""
        if approved_verbs:
            verb_instruction = (f"The course is set to Bloom's Level {bloom_level_int} ({bloom_level_str}). "
                                f"Therefore, each objective you generate **MUST** start with one of the following "
                                f"approved action verbs: `{', '.join(approved_verbs)}`.")

        system_prompt = "You are an e-learning instructional designer creating a course objectives slide. Synthesize key objectives from the provided content. Your tone should be direct and informative."

        # --- MODIFIED PROMPT ---
        # Instructions are now outside the JSON example for clarity.
        user_prompt = f"""
Based on the following combined course narration, you will generate content for a course objectives slide.
Course Title: "{course_title}"
Course Modules: {', '.join(unique_module_titles) if unique_module_titles else 'various topics'}.

**TASKS & GUIDELINES:**
1.  **"On-screen text" Task:** Synthesize 3-5 high-level learning objectives from the `Combined Course Narration`.
2.  **Verb Requirement (CRITICAL):** {verb_instruction}
3.  **"Narration" Task:** Write a standard narration that introduces the course objectives.
4.  **Language:** Follow the language instructions provided below.

---
**Combined Course Narration:**
{course_context}
---

**OUTPUT INSTRUCTIONS:**
Your output **MUST** be a single, raw JSON object using the exact format below. Do not add explanations.

**JSON OUTPUT FORMAT:**
{{
    "Learning Objectives": "N/A",
    "On-screen text": "A bulleted list of the 3-5 course objectives you created, following all rules.",
    "Narration": "The introductory narration you wrote for this slide.",
    "On-screen Recommendations": "Suggest visuals like bullet points with icons or a graphic symbolizing achievement.",
    "Developer Notes": "Provide notes for the developer, e.g., 'Use clear bullet points. Ensure text is readable.'"
}}

Follow this :
{english_instruction}
"""
        result = call_openai_for_json(system_prompt, user_prompt, temp=0.6)
        if result:
            slide_dict.update(result)
            slide_dict["Learning Objectives"] = "N/A"
            if isinstance(slide_dict.get("On-screen text"), list):
                slide_dict["On-screen text"] = "\n".join(f"- {obj}" for obj in slide_dict["On-screen text"])
        return slide_dict

def generate_final_course_assessment_slide(course_title: str, course_context: str) -> dict:
    print("  - Generating Final Course Assessment Content (5 questions)...")
    slide_dict = create_base_slide_dict(course_title, "Final Assessment", "Final Assessment - Introduction", "Static Screen", "00:30")
    system_prompt = "You are an e-learning instructional designer specializing in creating diverse and comprehensive final assessments with detailed feedback."
    global ENGLISH_VARIANT
    english_instruction = get_english_instruction(ENGLISH_VARIANT)
    
    user_prompt = f"""

Create a varied 5-question end-of-course assessment for "{course_title}" based on the provided course content.

**Instructions:**
1.  Generate exactly **5 questions** using a variety of the question types.
2.  For each question, select the most appropriate type and provide a clear `instruction`.
3.  For multiple-choice style questions, provide a `feedback` object with a rationale for **EVERY** option. For other types, provide a single `correct_answer_feedback`.
4.  **CRITICAL:** For each individual question object you create, the values for `Narration`, `On-screen Recommendations`, and `Developer Notes` **MUST** be the string "N/A".

Course Content:
{course_context}

**OUTPUT INSTRUCTIONS:**
Your output MUST be a single, raw JSON object. The top-level object will have its own Narration/Recommendations. However, EACH of the 5 question objects inside the "Questions" list must follow one of the structures below, with "N/A" for the specified fields.

**1. MCSS (Multiple Choice Single Select):**
{{
  "question_type": "MCSS", "question": "...", "instruction": "...", "options": {{...}}, "correct_answer": "...", "feedback": {{...}},
  "Narration": "N/A", "On-screen Recommendations": "N/A", "Developer Notes": "N/A"
}}

**2. MCMS (Multiple Choice Multiple Select):**
{{
  "question_type": "MCMS", "question": "...", "instruction": "...", "options": {{...}}, "correct_answer": [...], "feedback": {{...}},
  "Narration": "N/A", "On-screen Recommendations": "N/A", "Developer Notes": "N/A"
}}

**3. True/False, Drag and Drop, Fill in the Blank, etc. must also include:**
"Narration": "N/A", "On-screen Recommendations": "N/A", "Developer Notes": "N/A"

**Example of Final JSON Structure:**
{{
    "Narration": "Welcome to the final assessment...",
    "On-screen Recommendations": "For the intro slide: a clean title. For question slides: display one per screen.",
    "Developer Notes": "The following 5 questions should be implemented as 5 sequential slides.",
    "Questions": [
        {{ ... question 1 object, including the "N/A" fields ... }},
        {{ ... question 2 object, including the "N/A" fields ... }}
    ]
}}

Follow this :
{english_instruction}
"""
    result = call_openai_for_json(system_prompt, user_prompt, temp=0.8)
    if result:
        slide_dict.update(result)
        slide_dict["Learning Objectives"] = "N/A"
    return slide_dict


def generate_summary_slide(course_title: str, course_context: str) -> dict:
    print("  - Generating Course Summary Slide...")
    global ENGLISH_VARIANT
    english_instruction = get_english_instruction(ENGLISH_VARIANT)
    slide_dict = create_base_slide_dict(course_title, "Conclusion", "Course Summary", "Static Screen", "02:30")
    system_prompt = "You are an e-learning instructional designer creating a course summary slide. Focus on key takeaways and reinforcing learning. Use a motivational and conclusive tone."
    user_prompt = f"""
Create a concise and engaging summary slide for the course "{course_title}".
Synthesize the 3-5 most important key takeaways and main points from the following entire course content.
The summary should reinforce learning and provide a sense of completion.

Course Content:
{course_context}

**Output MUST be a JSON object with these keys:**
{{
    "Learning Objectives": "An objective focused on recalling key concepts and reinforcing learning.",
    "On-screen text": "A bulleted list of 3-5 key summary points or takeaways from the entire course.",
    "Narration": "A narration that effectively reiterates the main concepts learned in the course, emphasizes their importance, and provides a concluding message. (Use British English)",
    "On-screen Recommendations": "Suggest visuals for a summary screen (e.g., key icons for each point, a graphic symbolizing completion or mastery, a 'What next?' prompt).",
    "Developer Notes": "Provide notes for implementing the summary slide, eg., 'Ensure summary points are clear and concise. Consider a final motivational message.'"
}}
Follow this :
{english_instruction}
"""
    result = call_openai_for_json(system_prompt, user_prompt, temp=0.6)
    if result:
        slide_dict.update(result)
        slide_dict["Learning Objectives"] = "N/A"

        if isinstance(slide_dict.get("On-screen text"), list):
            slide_dict["On-screen text"] = "\n".join(f"- {point}" for point in slide_dict["On-screen text"])
    return slide_dict


# ==============================================================================
# ORCHESTRATION AND OUTPUT
# ==============================================================================
def process_one_row_in_two_steps(client: AzureOpenAI, deployment_name: str, row_data: dict, bloom_level: int,
                                 recently_used: list[str], course_title: str,
                                 recently_used_verbs: list[str], user_context_prompt: str | None,
                                 compliance_settings: dict = None,
                                 forced_learning_objective: str = None) -> dict | None:
    """
    Process a single row into a storyboard slide, now with English variant support and Shared Learning Objectives.
    """
    # Step 1: Select the appropriate screen type based on content analysis
    rich_content_parts = row_data.get("Full Page Content", [])
    if not rich_content_parts:
        print("    - SKIPPING: No content for this row.")
        return None
    chosen_screen_type = select_screen_type(client, deployment_name, rich_content_parts, recently_used,
                                            user_context_prompt)
    if chosen_screen_type:
        # Step 2: Generate the complete storyboard content with the selected screen type
        # Pass compliance_settings and forced_learning_objective to the prompt generator
        
        storyboard = generate_storyboard_content(client, deployment_name, row_data, chosen_screen_type, bloom_level,
                                                 course_title, recently_used_verbs, user_context_prompt,
                                                 compliance_settings, forced_learning_objective)
        return storyboard
    else:
        print("    - ABORTING row processing because screen type could not be determined.")
        return None


def generate_all_storyboards(all_rows: list[dict], client: AzureOpenAI, deployment_name: str, bloom_level_int: int,
                             course_title: str, user_context_prompt: str | None , compliance_settings: dict = None) -> tuple[list[dict], str | None]:
    final_storyboards = []
    total_rows = len(all_rows)
    recently_used_types = deque(maxlen=3)
    recently_used_verbs = deque(maxlen=2)

    current_module_title = None
    accumulated_narration_for_kc = []
    captured_objectives_text = None

    # --- Pre-calculate Shared Learning Objectives (Grouped by content flow, not Topic) ---
    row_index_to_lo_map = pre_calculate_learning_objectives(all_rows, client, deployment_name, bloom_level_int, compliance_settings, course_title)

    print(f"\nFound {total_rows} rows to process into individual storyboards for course '{course_title}'.")
    for i, row_data in enumerate(all_rows):
        row_module_title = _get_title_from_row(row_data, "Chapter", course_title)
        log_topic_text = _get_title_from_row(row_data, "Topic_Outline", "Untitled")
        log_subtopic_text = _get_title_from_row(row_data, "Subtopic_Outline", "Untitled")

        log_title_text = log_subtopic_text if log_subtopic_text != "Untitled" else log_topic_text

        if 'course overview and objectives' in log_topic_text.lower():
            print(
                f"\n--- Found 'Course Overview and Objectives' row ({i + 1}/{total_rows}). Checking for manual content... ---")
            content_parts = row_data.get("Full Page Content", [])
            extracted_text = "\n".join([part['content'] for part in content_parts if part['type'] == 'text']).strip()
            if extracted_text:
                print(
                    "    - Manual objectives content found. Capturing it and skipping storyboard generation for this row.")
                captured_objectives_text = extracted_text
                continue
            else:
                print("    - Row is empty. Objectives will be AI-generated later.")

        if current_module_title is not None and row_module_title != current_module_title:
            print(f"\n--- Chapter Change Detected: '{current_module_title}' finished. ---")
            if accumulated_narration_for_kc:
                print(
                    "      - WARNING: Previous chapter ended without an explicit Knowledge Check. Accumulated narration will be discarded.")
            accumulated_narration_for_kc = []
            recently_used_verbs.clear()
            recently_used_types.clear()
           

        current_module_title = row_module_title

        if 'knowledge check' in log_title_text.lower():
            print(
                f"\n--- Processing Row {i + 1}/{total_rows}: '{log_title_text}' (Chapter: '{current_module_title}') ---")

            # --- NEW LOGIC START ---
            # Inspect the cell to see if a manual question was provided.
            content_parts = row_data.get("Full Page Content", [])
            manual_question_text = "\n".join([part['content'] for part in content_parts if part['type'] == 'text']).strip()

            kc_module_title = current_module_title if current_module_title else course_title

            if not manual_question_text:
                # SCENARIO 1: The cell is EMPTY. Generate a 5-question quiz with AI.
                print("    - Empty Knowledge Check row found. Generating one questions with AI...")
                knowledge_check_slides = generate_knowledge_check(
                    client, deployment_name,accumulated_narration_for_kc, course_title, kc_module_title,
                    user_context_prompt,bloom_level_int=bloom_level_int
                )
                if knowledge_check_slides:
                    final_storyboards.extend(knowledge_check_slides)
                print(f"  --- Successfully processed AI-Generated Knowledge Check for '{kc_module_title}' ---")

            else:
                # SCENARIO 2: The cell has TEXT. Create a single slide with that text.
                print("    - Manual Knowledge Check question found. Creating a single slide...")
                manual_kc_slide = create_base_slide_dict(
                    course_title=course_title,
                    module_title=kc_module_title,
                    topic="Knowledge Check - Manual Question",
                    screen_type="Interactive Quiz",
                    duration="00:45"
                )
                manual_kc_slide.update({
                    "On-screen text": manual_question_text,
                    "Narration": "Please answer the question on the screen.",
                    "Developer Notes": "Implement as a simple question/answer interaction. This question was provided manually in the source document.",
                    "Learning Objectives": "N/A"
                })
                final_storyboards.append(manual_kc_slide)
                print(f"  --- Successfully processed Manual Knowledge Check for '{kc_module_title}' ---")
            
            # --- NEW LOGIC END ---

            # Clear accumulators after generating EITHER type of KC
            accumulated_narration_for_kc = []
            recently_used_verbs.clear()
            recently_used_types.clear()
            continue # Move to the next row

        print(f"\n--- Processing Row {i + 1}/{total_rows}: '{log_title_text}' (Chapter: '{current_module_title}') ---")

        # Get the pre-calculated LO for this row, if it exists
        forced_lo = row_index_to_lo_map.get(i)
        if forced_lo:
             print(f"    - Using Pre-calculated Shared Learning Objective.")

        storyboard = process_one_row_in_two_steps(client, deployment_name, row_data, bloom_level_int,
                                                  list(recently_used_types), course_title, list(recently_used_verbs),
                                                  user_context_prompt,compliance_settings, forced_lo)
        if storyboard:
            final_storyboards.append(storyboard)
            newly_used_type = storyboard.get("Screen-type")
            if newly_used_type: recently_used_types.append(newly_used_type)

            objective = storyboard.get("Learning Objectives", "")
            if objective:
                first_word_match = re.match(r'^[A-Za-z]+', objective.strip())
                if first_word_match:
                    used_verb = first_word_match.group(0).capitalize()
                    if used_verb in BLOOM_VERBS.get(bloom_level_int, []):
                        recently_used_verbs.append(used_verb)

            narration_for_current_slide = storyboard.get("Narration", "").strip()
            narration_cleaned = re.sub(r'-\s*\[Narration for .*?\]\.\.\.\s*', '', narration_for_current_slide,
                                       flags=re.IGNORECASE).strip()

            if narration_cleaned:
                accumulated_narration_for_kc.append(narration_cleaned)
        

            print(f"  --- Successfully completed storyboard for '{log_title_text}' ---")
        else:
            print(f"  --- Failed to generate storyboard for '{log_title_text}' ---")

    if accumulated_narration_for_kc:
        print(
            f"\n--- End of Processing: There's accumulated narration for chapter '{current_module_title}' but no explicit Knowledge Check followed. ---")

    print("\n--- All core rows processed. ---")
    return final_storyboards, captured_objectives_text

def save_storyboards_to_word_document(storyboards: list[dict], filename: str):
    if not storyboards:
        print("No storyboards were generated, skipping Word document creation.")
        return
    print(f"\nCreating Word document with individual slide tables...")
    document = docx.Document()

    # Define the fields to be displayed in the table, including the new "Source Tables" field
    standard_fields = ["Course Title", "Module Title", "Topic", "Screen-type", "Bloom's Level",
                       "Learning Objectives", "On-screen text", "Narration", "On-screen Recommendations",
                       "Developer Notes", "Duration (min)", "Source Images (base64)", "Source Tables"]

    for i, storyboard in enumerate(storyboards):
        if i > 0: document.add_page_break()
        slide_title = storyboard.get("Topic", "Untitled Slide")
        module_title = storyboard.get("Module Title", "Unknown Module")
        course_title = storyboard.get("Course Title", "Unknown Course")

        document.add_heading(f"{course_title}", level=1)
        document.add_heading(f"Module: {module_title}", level=2)
        document.add_heading(f"Slide {i + 1}: {slide_title}", level=3)

        # We now use the single, updated list of fields for all slides
        current_fields_to_display = standard_fields

        table = document.add_table(rows=len(current_fields_to_display), cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Inches(1.75)
        table.columns[1].width = Inches(4.5)

        for row_index, field_name in enumerate(current_fields_to_display):
            row_cells = table.rows[row_index].cells
            p = row_cells[0].paragraphs[0]
            run = p.add_run(field_name)
            run.bold = True

            if field_name == 'Source Images (base64)':
                image_info_list = storyboard.get(field_name, [])
                images_cell = table.rows[row_index].cells[1]
                images_cell.text = ''

                if not image_info_list:
                    images_cell.add_paragraph("No images specified for this slide.")
                else:
                    for item in image_info_list:
                        b64_string = item.get("content")
                        mime_type = item.get("mime_type")

                        if b64_string and mime_type and re.match(
                                r'^(?:[A-Za-z0-p+/]{4})*(?:[A-Za-z0-p+/]{2}==|[A-Za-z0-p+/]{3}=)?$', b64_string,
                                re.IGNORECASE):
                            try:
                                decoded_image_data = base64.b64decode(b64_string)
                                image_stream = io.BytesIO(decoded_image_data)
                                p_image = images_cell.add_paragraph()
                                run_image = p_image.add_run()
                                run_image.add_picture(image_stream, width=Inches(4.0))
                                images_cell.add_paragraph(f"(Embedded image from source: {mime_type})")
                            except Exception as e:
                                images_cell.add_paragraph(f"❗️ Error embedding image from base64: {e}")
                                images_cell.add_paragraph(f"Original Data (truncated): {b64_string[:100]}...")
                        else:
                            images_cell.add_paragraph(item.get("content", str(item)))

            elif field_name == 'Source Tables':
                table_list = storyboard.get(field_name, [])
                tables_cell = table.rows[row_index].cells[1]
                tables_cell.text = ''  # Clear existing text in the cell

                if not table_list:
                    tables_cell.add_paragraph("No tables specified for this slide.")
                else:
                    # Join all found tables with a separator
                    full_text = "\n\n".join(table_list)
                    tables_cell.text = full_text

            else:
                # This generic block handles all other text-based fields
                value = storyboard.get(field_name, "")
                cell_text = str(value)
                row_cells[1].text = cell_text

    document.save(filename)
    print(f"... Successfully saved storyboards to '{filename}'")


if __name__ == "__main__":
    # --- CONFIGURATION ---
    word_file_path = 'NTC_Project Life Cycles_CO.docx'
    base_output_filename = os.path.splitext(os.path.basename(word_file_path))[0]
    json_output_filename = f"{base_output_filename}_FullCourse_Storyboard.json"
    word_output_filename = f"{base_output_filename}_FullCourse_Storyboard.docx"
    # --- END CONFIGURATION ---

    print(f"Starting full course storyboard generation based on '{word_file_path}'...")

    bloom_level_input = get_bloom_level_from_user()

    try:
        user_context_prompt = read_user_context_prompt_from_word(word_file_path)
        course_title_from_filename = os.path.splitext(os.path.basename(word_file_path))[0].replace('_', ' ').replace(
            '-', ' ').replace('new', '').strip().title()
        input_rows_from_file = read_rows_from_word_table(word_file_path)

        if not input_rows_from_file:
            raise RuntimeError("No data could be read from the input file.")

        # This call now returns two values: the list of storyboards and the optional manual objectives text.
        main_course_storyboards, manual_objectives_content = generate_all_storyboards(
            input_rows_from_file, client, deployment_name,
            bloom_level_input, course_title_from_filename,
            user_context_prompt
        )

        # Accumulate narration and module titles for intro/outro slides
        full_course_narration_list = []
        unique_module_titles = set()
        for slide in main_course_storyboards:
            if slide.get("Narration"):
                narration_cleaned = re.sub(r'-\s*\[Narration for .*?\]\.\.\.\s*', '', slide["Narration"],
                                           flags=re.IGNORECASE).strip()
                if narration_cleaned:
                    full_course_narration_list.append(narration_cleaned)
            if slide.get("Module Title"):
                unique_module_titles.add(slide["Module Title"])

        full_course_narration = "\n\n".join(filter(None, full_course_narration_list))
        max_narration_for_prompt = 100000
        if len(full_course_narration) > max_narration_for_prompt:
            print(f"WARNING: Full course narration is very long ({len(full_course_narration)} chars). Truncating.")
            full_course_narration = full_course_narration[:max_narration_for_prompt]

        if not full_course_narration.strip():
            print("WARNING: No substantial narration accumulated. Intro/outro slides may be generic.")

        print("\n--- Generating Introductory Slides ---")
        welcome_slide = generate_welcome_slide(course_title_from_filename, full_course_narration)
        navigation_slide = generate_navigation_slide(course_title_from_filename, full_course_narration)

        # This call now passes the manual_objectives_content to the function.
        course_objectives_slide = generate_course_objectives_slide(
            course_title_from_filename,
            manual_objectives_content,
            full_course_narration,
            list(unique_module_titles),
            bloom_level_input
        )

        print("\n--- Generating Concluding Slides ---")
        assessment_master_content = generate_final_course_assessment_slide(course_title_from_filename,
                                                                           full_course_narration)
        summary_slide = generate_summary_slide  (course_title_from_filename, full_course_narration)

        final_assessment_slides = []
        if assessment_master_content and "Questions" in assessment_master_content:
            assessment_intro_slide = assessment_master_content.copy()
            assessment_intro_slide.pop("Questions", None)
            final_assessment_slides.append(assessment_intro_slide)

            questions_list = assessment_master_content["Questions"]
            print(f"  - Splitting assessment into {len(questions_list)} individual question slides...")
            for i, q_data in enumerate(questions_list):
                q_type = q_data.get("question_type", "MCSS")  # Default to MCSS if type is missing
                question_slide = create_base_slide_dict(
                    course_title_from_filename, "Final Assessment", f"Question {i + 1} ({q_type})", "Interactive Quiz",
                    "00:45"
                )

                # Logic to format the On-screen text
                on_screen_parts = [
                    f"{q_data.get('question', 'Missing question text.')}\n\nInstruction: {q_data.get('instruction', 'Please answer the question.')}"]

                if q_type in ["MCSS", "MCMS", "True/False"]:
                    options = q_data.get('options', {})
                    for key, value in options.items():
                        on_screen_parts.append(f"{key}. {value}")
                    correct = q_data.get('correct_answer', 'N/A')
                    correct_str = ", ".join(correct) if isinstance(correct, list) else correct
                    on_screen_parts.append(f"\nCorrect Answer(s): {correct_str}")
                    on_screen_parts.append("\n--- Feedback for All Options ---")
                    feedback = q_data.get('feedback', {})
                    for key in sorted(options.keys()):
                        on_screen_parts.append(f"Rationale for {key}: {feedback.get(key, 'No rationale provided.')}")

                elif q_type == "Drag and Drop":
                    on_screen_parts.append("\n--- Items to Match ---")
                    on_screen_parts.append("Drag Items: " + ", ".join(q_data.get("drag_items", [])))
                    on_screen_parts.append("Drop Targets: " + ", ".join(q_data.get("drop_targets", [])))
                    on_screen_parts.append("\n--- Correct Pairing ---")
                    answers = q_data.get('correct_answer', {})
                    for key, value in answers.items():
                        on_screen_parts.append(f"{key} -> {value}")
                    on_screen_parts.append(
                        f"\nRationale: {q_data.get('correct_answer_feedback', 'No rationale provided.')}")

                elif q_type == "Fill in the Blank":
                    on_screen_parts.append(f"\nCorrect Answer: {q_data.get('correct_answer', 'N/A')}")
                    on_screen_parts.append(
                        f"Rationale: {q_data.get('correct_answer_feedback', 'No rationale provided.')}")

                formatted_on_screen_text = "\n".join(on_screen_parts)

                # Update the slide with values from the AI, defaulting to "N/A" if they are missing.
                question_slide.update({
                    "Learning Objectives": "",
                    "On-screen text": formatted_on_screen_text,
                    "Narration": q_data.get("Narration", ""),
                    "On-screen Recommendations": q_data.get("On-screen Recommendations", ""),
                    "Developer Notes": q_data.get("Developer Notes", "")
                })
                final_assessment_slides.append(question_slide)
        else:
            print("  - WARNING: Final assessment generation failed or returned no questions.")

        combined_final_storyboards = [
            welcome_slide, navigation_slide, course_objectives_slide,
            *main_course_storyboards, *final_assessment_slides, summary_slide
        ]

        if combined_final_storyboards:
            # 1. Save the Word document using the original, human-readable headers.
            save_storyboards_to_word_document(combined_final_storyboards, word_output_filename)

            # 2. Transform the headers for the JSON output.
            print(f"\nTransforming JSON headers to the required format...")
            storyboards_for_json = transform_storyboard_headers(combined_final_storyboards, HEADER_MAPPING)

            # 3. Save the JSON file using the new, transformed headers.
            with open(json_output_filename, 'w', encoding='utf-8') as f:
                json.dump(storyboards_for_json, f, indent=4, ensure_ascii=False)
            print(f"✅ JSON output with required headers saved to '{json_output_filename}'")

            print(f"\n✅✅✅ FULL COURSE STORYBOARD GENERATION COMPLETE ✅✅✅")
            print(f"Generated a total of {len(combined_final_storyboards)} slides.")
        else:
            print("\nProcessing finished, but no storyboards were generated.")

    except Exception as e:
        print(f"\nAn unexpected error occurred during the main execution: {e}")