# /core/save_storyboards_to_adls.py
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import base64
import logging
from typing import Dict
from dotenv import load_dotenv
from azure.storage.filedatalake import DataLakeServiceClient
from azure.core.exceptions import AzureError, ResourceNotFoundError
# Docx imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import AzureOpenAI

# Azure imports
from azure.storage.filedatalake import DataLakeServiceClient

# HTML parsing for tables
from bs4 import BeautifulSoup

# Load environment variables from .env file
load_dotenv()

# Azure Storage Credentials
AZURE_STORAGE_ACCOUNT_NAME = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
AZURE_STORAGE_ACCOUNT_KEY = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
AZURE_STORAGE_FILESYSTEM_NAME = os.getenv("AZURE_STORAGE_FILESYSTEM_NAME")

logger = logging.getLogger(__name__)


# ======================= HELPER FUNCTIONS =======================

def _upload_stream_to_adls(stream: io.BytesIO, filename: str, client: str, project: str, module: str):
    """
    Uploads a file stream to ADLS with a structured path.
    Returns the full file path on success.
    """
    if not all([AZURE_STORAGE_ACCOUNT_NAME, AZURE_STORAGE_ACCOUNT_KEY, AZURE_STORAGE_FILESYSTEM_NAME]):
        logger.error("Azure Storage credentials are not configured in environment variables.")
        raise ValueError("Azure Storage credentials are not configured.")

    try:
        datalake_service_client = DataLakeServiceClient(
            account_url=f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net",
            credential=AZURE_STORAGE_ACCOUNT_KEY
        )
        file_system_client = datalake_service_client.get_file_system_client(AZURE_STORAGE_FILESYSTEM_NAME)

        # Use structured path: client/project/module/Storyboard/filename
        directory_path = f"{client}/{project}/{module}/Storyboard"
        file_path = f"{directory_path}/{filename}"
        
        file_system_client.create_directory(directory_path)
        logger.info(f"Ensured directory path exists: {directory_path}")

        stream.seek(0)
        data = stream.read()

        file_client = file_system_client.get_file_client(file_path)
        file_client.upload_data(data, overwrite=True)

        full_path = f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net/{AZURE_STORAGE_FILESYSTEM_NAME}/{file_path}"
        logger.info(f"Successfully uploaded storyboard to ADLS: {full_path}")
        return full_path

    except Exception as e:
        logger.error(f"Failed to upload storyboard to ADLS: {e}", exc_info=True)
        raise

def get_versioned_filename(directory_path: str, filename: str, file_system_client) -> str:
    """
    Returns a versioned filename if the file already exists in ADLS.
    Example: If 'ÁME_Module1_SB.docx' exists, returns 'ÁME_Module1_SB(1).docx', etc.
    """
    base, ext = os.path.splitext(filename)
    version = 1
    candidate = filename
    
    # Remove trailing slash if it exists
    directory_path = directory_path.rstrip('/')
        
    while True:
        # Build the full path with a single slash
        file_path = f"{directory_path}/{candidate}"
        try:
            file_client = file_system_client.get_file_client(file_path)
            exists = file_client.exists()
        except Exception as e:
            logger.error(f"Error checking if file exists: {e}")
            return candidate  # Return original filename on error
        
        if not exists:
            return candidate
        
        candidate = f"{base}({version}){ext}"  # Version before extension
        version += 1


def objective_to_bullet(client: AzureOpenAI, deployment_name: str, objectives_text: str) -> str:
    """
    Takes a string of learning objectives, sends it to an LLM, and returns a clean,
    bulleted list.
    """
    if not objectives_text or not objectives_text.strip():
        logger.info("Input objectives text is empty. Skipping LLM call.")
        return "- Key learning outcomes for the course." # Return a default placeholder

    system_prompt = "You are a formatting expert who transforms text into a clean, bulleted list without adding any extra words or commentary."
    user_prompt = f"""
Reformat the following text into a simple, bulleted list. Each bullet point should start with an action verb.

**CRITICAL RULE:** Your response MUST contain ONLY the bulleted list itself. Do NOT include any introductory text like "Here is the list:" or any other words before or after the list. Start your response directly with the first bullet point.

**TEXT TO REFORMAT:**
---
{objectives_text}
---
"""
    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1
        )
        formatted_list = response.choices[0].message.content.strip()
        return formatted_list
    except Exception as e:
        logger.warning(f"LLM call to format objectives failed: {e}")
        return objectives_text


def save_storyboard_to_adls(storyboard_data: Dict, filename: str, llm_client: AzureOpenAI, llm_deployment_name: str) -> str:
    """
    Generates a complete storyboard DOCX file in memory with proper formatting
    and uploads it to Azure Data Lake Storage, using versioned filenames.
    """
    doc = Document()
    doc.add_heading("Storyboard Document", 0)

    storyboards = storyboard_data.get("storyboard", [])
    if not storyboards:
        raise ValueError("No storyboard items found in the provided data.")

    # ===== Course/module info (once at top) =====
    first_sb = storyboards[0]

    print(first_sb)
    
    # Add Course Title, Module Title, Bloom's Level
    doc.add_heading("Course Title", level=2)
    doc.add_paragraph(first_sb.get("Course_Title", "N/A"))
    
    # doc.add_heading("Module Title", level=2)
    # doc.add_paragraph(first_sb.get("Module_Title", "N/A"))
    
    doc.add_heading("Bloom's Level", level=2)
    blooms_level_value = first_sb.get("Blooms_Level", "N/A")
    doc.add_paragraph(blooms_level_value)

    # Calculate total duration - FIXED: Check both field name formats
    total_seconds = sum(parse_duration(sb.get("Duration_(min)", sb.get("Duration_min", ""))) for sb in storyboards)
    formatted_duration = format_hhmmss(total_seconds)
    
    doc.add_heading("Total Duration (HH:MM:SS)", level=2)
    doc.add_paragraph(formatted_duration)
    doc.add_paragraph()

    # ====== Add Course Objectives (On_screen_text of every third storyboard) ======
    doc.add_heading("Course objectives:", level=2)

    raw_objectives_text = "No objectives found."
    if len(storyboards) >= 3:
        raw_objectives_text = storyboards[2].get("On_screen_text", "No objectives found.")

    # 2. Pre-process the objectives to get a clean bulleted list using the LLM
    cleaned_objectives = objective_to_bullet(
        client=llm_client,
        deployment_name=llm_deployment_name,
        objectives_text=raw_objectives_text
    )

    # 3. Add each objective as a separate bullet point to the document
    objective_lines = [line.strip() for line in cleaned_objectives.split('\n') if line.strip()]
    if not objective_lines:
        objective_lines.append("No objectives found.")

    for line in objective_lines:
        # Remove any leading bullet characters the LLM might have added
        if line.startswith(('-', '*', '•')):
            line = line[1:].lstrip()

        # Add the cleaned line as a bullet point with black text
        p = doc.add_paragraph(line, style='List Bullet')
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()

    # ====== Add Glossary Table ======
    doc.add_heading("Glossary:", level=2)

    glossary_table = doc.add_table(rows=1, cols=2)
    glossary_table.style = 'Table Grid'
    hdr_cells = glossary_table.rows[0].cells
    hdr_cells[0].text = "Term"
    hdr_cells[1].text = "Description"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    glossary_data = [
        ("Term 1", "Description 1"),
        ("Term 2", "Description 2"),
        ("Term 3", "Description 3"),
        ("Term X", "Description X"),
    ]
    for term, desc in glossary_data:
        row_cells = glossary_table.add_row().cells
        row_cells[0].text = term
        row_cells[1].text = desc
        for i, cell in enumerate(row_cells):
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    keep_table_on_one_page(glossary_table)
    doc.add_paragraph()

    # ====== Add Resources Table ======
    doc.add_heading("Resources:", level=2)

    resources_table = doc.add_table(rows=1, cols=2)
    resources_table.style = 'Table Grid'
    hdr_cells = resources_table.rows[0].cells
    hdr_cells[0].text = "Resource"
    hdr_cells[1].text = "URL/File name and path"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    resources_data = [
        ("Resource 1", "URL 1"),
        ("Resource 2", "URL 2"),
        ("Resource 3", "URL 3"),
        ("Resource X", "URL X"),
    ]
    for res, url in resources_data:
        row_cells = resources_table.add_row().cells
        row_cells[0].text = res
        row_cells[1].text = url
        for i, cell in enumerate(row_cells):
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    keep_table_on_one_page(resources_table)
    doc.add_paragraph()

    # ===== Loop through each storyboard screen =====
    for i, sb in enumerate(storyboards):
        # Add page break before each screen (except the first one)
        if i > 0:
            doc.add_page_break()
        
        doc.add_heading(f"Screen {i + 1}", level=2)

        fields = [
            ("Topic", sb.get("Topic", "")),
            ("Screen Type", sb.get("Screen_type", "")),
            ("Learning Objectives", "\n".join(sb.get("Learning_Objectives", []))),
            ("On-screen Text", sb.get("On_screen_text", "")),
            ("Narration", sb.get("Narration", "")),
            ("On-screen Recommendations", "\n".join(sb.get("On_screen_Recommendations", []))),
            ("Developer Notes", "\n".join(sb.get("Developer_Notes", []))),
            # FIXED: Check both field name formats for Duration
            ("Duration (min)", sb.get("Duration_(min)", sb.get("Duration_min", ""))),
            ("Source Images", ""),  # Placeholder
            ("Source Tables", ""),  # Placeholder
        ]

        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"

        for row_idx, (label, value) in enumerate(fields):
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            if value:
                row_cells[1].text = value

        # --- Handle Images ---
        image_cell = table.rows[-2].cells[1]
        image_cell.text = ""  # Clear placeholder
        # Use the correct field name that matches frontend
        image_list = sb.get("Source_Images_base64", []) or sb.get("Source_Images_(base64)", [])
        if image_list and image_list != ["N/A"]:
            p = image_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for img_base64 in image_list:
                try:
                    if img_base64 and img_base64 != "N/A":
                        image_data = base64.b64decode(img_base64)
                        image_stream = io.BytesIO(image_data)
                        p.add_run().add_picture(image_stream, width=Inches(2.0))
                        p.add_run("  ")
                except Exception as e:
                    logger.warning(f"Could not process image on Screen {i+1}: {e}")
        else:
            image_cell.text = "N/A"

        # --- Handle HTML Tables ---
        tables_cell = table.rows[-1].cells[1]
        tables_cell.text = ""  # Clear placeholder
        table_html_list = sb.get("Source_Tables", [])
        if table_html_list and table_html_list != ["N/A"]:
            for table_html in table_html_list:
                soup = BeautifulSoup(table_html, "html.parser")
                for html_table in soup.find_all("table"):
                    rows = html_table.find_all("tr")
                    if not rows:
                        continue
                    cols_count = max(len(r.find_all(["th", "td"])) for r in rows)
                    docx_table = tables_cell.add_table(rows=len(rows), cols=cols_count)
                    docx_table.style = "Table Grid"
                    for r_idx, row in enumerate(rows):
                        cells = row.find_all(["th", "td"])
                        for c_idx, cell in enumerate(cells):
                            docx_table.cell(r_idx, c_idx).text = cell.get_text(strip=True)
                    tables_cell.add_paragraph()
        else:
            tables_cell.text = "N/A"

        keep_table_on_one_page(table)
        doc.add_paragraph()

    # ===== Save document to memory stream and upload =====
    stream = io.BytesIO()
    doc.save(stream)
    
    # Extract metadata from storyboard data
    client = storyboard_data.get("client", "UnknownClient")
    project = storyboard_data.get("project", "UnknownProject")
    module = storyboard_data.get("module", "UnknownModule") 

    # Prepare ADLS client and directory path
    datalake_service_client = DataLakeServiceClient(
        account_url=f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net",
        credential=AZURE_STORAGE_ACCOUNT_KEY
    )
    file_system_client = datalake_service_client.get_file_system_client(AZURE_STORAGE_FILESYSTEM_NAME)
    directory_path = f"{client}/{project}/{module}/Storyboard"

    # Ensure directory exists
    try:
        file_system_client.create_directory(directory_path)
        logger.info(f"Ensured directory path exists: {directory_path}")
    except Exception as e:
        # Directory might already exist - this is fine
        logger.debug(f"Directory path check: {e}")

    # Use versioned filename logic 
    versioned_filename = get_versioned_filename(directory_path, filename, file_system_client)
    logger.info(f"Using versioned filename: {versioned_filename}")

    # Upload using the versioned filename
    return _upload_stream_to_adls(stream, versioned_filename, client, project, module)
##getoutline metadata


def download_file_from_adls(client: str, project: str, module: str, filename: str) -> bytes | None:
    """
    Download a file from Azure Data Lake Storage (ADLS).

    Args:
        client (str): Client name (top-level folder).
        project (str): Project name (sub-folder).
        module (str): Module name (sub-folder).
        filename (str): Name of the file to download.

    Returns:
        bytes | None: File content in bytes if successful, else None.
    """
    file_path = f"{client}/{project}/{module}/Outline/{filename}"
    try:
        # Validate required configurations
        if not (AZURE_STORAGE_ACCOUNT_NAME and AZURE_STORAGE_ACCOUNT_KEY and AZURE_STORAGE_FILESYSTEM_NAME):
            logger.error("Missing Azure Storage configuration.")
            return None

        account_url = f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
        service_client = DataLakeServiceClient(account_url=account_url, credential=AZURE_STORAGE_ACCOUNT_KEY)
        file_system_client = service_client.get_file_system_client(file_system=AZURE_STORAGE_FILESYSTEM_NAME)

        # Get file client
        file_client = file_system_client.get_file_client(file_path)

        # Check if file exists
        if not file_client.exists():
            logger.warning(f"File not found: {file_path}")
            return None

        # Download file
        download = file_client.download_file()
        downloaded_bytes = download.readall()

        logger.info(f"Successfully downloaded file: {file_path}")
        return downloaded_bytes

    except ResourceNotFoundError:
        logger.warning(f"File does not exist in ADLS: {file_path}")
        return None
    except AzureError as azure_err:
        logger.error(f"Azure error while downloading file '{file_path}': {azure_err}")
        return None
    except Exception as e:
        logger.exception(f"Unexpected error while downloading file '{file_path}': {e}")
        return None



def parse_duration(value):
    """
    Parses Duration_min values like:
    Returns total seconds.
    """
    if not value:
        return 0
    value = str(value).strip()

    # Case: MM:SS format
    if ":" in value:
        parts = value.split(":")
        if len(parts) == 2:
            try:
                minutes, seconds = int(parts[0]), int(parts[1])
                return minutes * 60 + seconds
            except:
                return 0

    # Case: plain number → minutes
    if value.isdigit():
        return int(value) * 60

    return 0


def format_hhmmss(total_seconds):
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}"


def keep_table_on_one_page(table):
    # For each row except the last, set "keep with next" for all paragraphs
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


