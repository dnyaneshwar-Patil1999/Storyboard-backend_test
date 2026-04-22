# """
# Style guide routes for handling client style guide operations.
# """
# from azure.core.exceptions import ResourceNotFoundError, AzureError
# from core.adls import FILESYSTEM_NAME, STORAGE_ACCOUNT_KEY, STORAGE_ACCOUNT_NAME, get_file_structure
# from azure.storage.filedatalake import DataLakeServiceClient
# # Fix time import - don't use datetime.time
# import time
# import os
# import sys
# import docx
# from dotenv import load_dotenv
# from fastapi import APIRouter, HTTPException, Form, File, UploadFile, Path
# from fastapi.responses import JSONResponse
# import fitz
# from openai import AzureOpenAI
# from app.models.project_models import get_style_guide_info, upload_style_guide
# import logging
# import json
# import traceback
# import tempfile
# # Add this import to fix the JobStore undefined error
# from app.services.job_store import JobStore

# logger = logging.getLogger(__name__)

# router = APIRouter()
# @router.get("/style-guide/{client}")
# async def get_style_guide_info_route(client: str = Path(...)):
#     """Get style guide information for a client"""
#     try:
#         # Get style guide info from project_models
#         result = get_style_guide_info(client)
        
#         if not result.get("exists"):
#             raise HTTPException(status_code=404, detail="No style guide found")
        
#         # Get rules from the corresponding TXT file
#         pdf_filename = result.get("file_name", "")
#         if pdf_filename:
#             rules_text = get_rules_from_adls(client, pdf_filename)
#             result["rules_text"] = rules_text
#             result["rules_found"] = bool(rules_text)
        
#         return JSONResponse(content=result)
        
#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Error in get_style_guide_info_route: {e}")
#         raise HTTPException(status_code=500, detail="Internal server error")

# # @router.post("/style-guide/upload")
# # async def upload_style_guide_route(
# #     client: str = Form(...),
# #     file: UploadFile = File(...),
# #     replace: bool = Form(True),  # Default to replace existing files
# # ):
# #     """Upload style guide file for a client and extract rules"""
# #     try:
# #         if not file.filename or not file.filename.lower().endswith(".pdf"):
# #             raise HTTPException(status_code=400, detail="Only PDF files are allowed")

# #         file_content = await file.read()
# #         if len(file_content) > 20 * 1024 * 1024:
# #             raise HTTPException(
# #                 status_code=400,
# #                 detail="File size too large. Maximum 20MB allowed",
# #             )

# #         # Save temporarily for processing
# #         temp_path = f"/tmp/{file.filename}"
# #         with open(temp_path, "wb") as f:
# #             f.write(file_content)

# #         # Upload to ADLS with checks for existing style guides
# #         result = upload_style_guide(file.filename, file_content, client, replace)
# #         if not result.get("success"):
# #             raise HTTPException(status_code=400, detail=result.get("message"))

# #         # Process PDF and extract rules
# #         try:
# #             rules_text = process_pdf_and_extract_rules(temp_path, file.filename)
# #             logger.info(f"Extracted rules: {rules_text[:200]}...")  # Log first 200 chars
            
# #             # Save rules to ADLS as TXT file
# #             if rules_text:
# #                 if save_rules_to_adls(rules_text, client, file.filename):
# #                     result["rules_saved"] = True
# #                     result["rules_text"] = rules_text
# #                 else:
# #                     result["rules_saved"] = False
# #                     result["rules_text"] = None
# #                     logger.error("Failed to save rules to ADLS")
# #         except Exception as e:
# #             logger.error(f"Error extracting rules: {e}")
# #             result["rules_saved"] = False
# #             result["rules_text"] = None

# #         # Cleanup temp PDF
# #         os.remove(temp_path)
# #         return JSONResponse(content=result)
# #     except HTTPException:
# #         raise
# #     except Exception as e:
# #         logger.error(f"Error in upload_style_guide_route: {e}")
# #         raise HTTPException(status_code=500, detail="Internal server error")



# @router.post("/style-guide/upload")
# async def upload_style_guide_route(
#     client: str = Form(...),
#     file: UploadFile = File(...),
#     replace: bool = Form(True),  # Default to replace existing files
# ):
#     """Upload style guide file for a client and extract rules"""
#     temp_path = None

#     try:
#         if not file.filename or not file.filename.lower().endswith(".pdf"):
#             raise HTTPException(status_code=400, detail="Only PDF files are allowed")

#         file_content = await file.read()
#         if len(file_content) > 20 * 1024 * 1024:  # 20MB limit
#             raise HTTPException(
#                 status_code=400,
#                 detail="File size too large. Maximum 20MB allowed.",
#             )

#         # Use a safe temporary directory compatible with Azure
#         temp_dir = tempfile.gettempdir()
#         temp_path = os.path.join(temp_dir, file.filename)

#         # Write file to temp path
#         with open(temp_path, "wb") as f:
#             f.write(file_content)

#         # Upload to ADLS with checks for existing style guides
#         result = upload_style_guide(file.filename, file_content, client, replace)
#         if not result.get("success"):
#             raise HTTPException(status_code=400, detail=result.get("message"))

#         # Process PDF and extract rules
#         try:
#             rules_text = process_pdf_and_extract_rules(temp_path, file.filename)
#             logger.info(f"Extracted rules (first 200 chars): {rules_text[:200]}")

#             if rules_text:
#                 saved = save_rules_to_adls(rules_text, client, file.filename)
#                 result["rules_saved"] = saved
#                 result["rules_text"] = rules_text if saved else None

#                 if not saved:
#                     logger.error("Failed to save extracted rules to ADLS")
#             else:
#                 result["rules_saved"] = False
#                 result["rules_text"] = None
#                 logger.warning("No rules extracted from PDF")

#         except Exception as e:
#             logger.error(f"Error extracting rules: {e}\n{traceback.format_exc()}")
#             result["rules_saved"] = False
#             result["rules_text"] = None

#         return JSONResponse(content=result)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Unexpected error in upload_style_guide_route: {e}\n{traceback.format_exc()}")
#         raise HTTPException(status_code=500, detail="Internal server error")
#     finally:
#         # Always clean up temp file
#         if temp_path and os.path.exists(temp_path):
#             try:
#                 os.remove(temp_path)
#             except Exception as e:
#                 logger.warning(f"Failed to delete temp file {temp_path}: {e}")


# def extract_text_from_pdf(pdf_path: str) -> list[str]:
#     """
#     Extracts text page by page from a PDF document.

#     Args:
#         pdf_path (str): The path to the PDF file.

#     Returns:
#         list[str]: A list where each element is the text content of a page.
#                    Returns an empty list if the file cannot be opened or has no pages.
#     """
#     if not os.path.exists(pdf_path):
#         print(f"Error: PDF file not found at '{pdf_path}'")
#         return []

#     page_texts = []
#     try:
#         document = fitz.open(pdf_path)
#         for page_num in range(len(document)):
#             page = document.load_page(page_num)
#             text = page.get_text("text")
#             page_texts.append(text)
#         document.close()
#     except Exception as e:
#         print(f"Error reading PDF '{pdf_path}': {e}")
#     return page_texts


# def generate_grammatical_rules_from_text(
#         client: AzureOpenAI,
#         deployment_name: str,
#         page_text: str,
#         page_number: int  # For logging/context
# ) -> str:
#     """
#     Uses an LLM to extract only grammatical and writing style rules from a given text chunk.

#     Args:
#         client (AzureOpenAI): The Azure OpenAI client instance.
#         deployment_name (str): The name of the Azure OpenAI deployment.
#         page_text (str): The text content from a page of the style guide.
#         page_number (int): The page number for logging.

#     Returns:
#         str: A string containing the extracted rules, or an empty string if no relevant rules
#              were found or an error occurred.
#     """
#     if not page_text.strip():
#         return ""  # No text to process

#     print(f"  - Analyzing page {page_number} with LLM...")

#     # The prompt explicitly tells the LLM what to focus on and what to ignore.
#     prompt_instruction = f"""
# You are an expert editor and linguist, specializing in analyzing detailed style guides.
# Your task is to review the following text from a client's style guide and **extract ONLY the grammatical and writing style rules**.

# **Specifically, focus on rules related to:**
# -   **Punctuation:** (commas, periods, semicolons, dashes, apostrophes, quotation marks, parentheses, etc.)
# -   **Capitalization:** (e.g., proper nouns, titles, headings, sentence beginnings)
# -   **Spelling:** (e.g., preference for UK vs. US English, specific terms, compound words)
# -   **Sentence Structure and Syntax:** (e.g., clear and concise sentences, avoiding jargon, parallel structure)
# -   **Word Choice and Terminology:** (e.g., preferred terms, forbidden words, consistent vocabulary)
# -   **Active/Passive Voice:** (guidelines on when to use each)
# -   **Tone of Voice:** (e.g., formal, informal, empathetic, direct, authoritative)
# -   **General Grammar:** (verb tenses, subject-verb agreement, modifiers, pronouns)
# -   **Abbreviations and Acronyms:** (first use expansion, consistent usage)
# -   **Number Formatting:** (when to use figures vs. words, percentages, dates)
# -   **Formatting for lists:** (bullet points, numbered lists, consistency)
# -   **Accessibility considerations for text:** (e.g., clear language, avoidance of complex sentences)

# **Crucially, you MUST IGNORE and filter out ANY rules related to:**
# -   **Visual Design:** (layout, spacing, white space, margins)
# -   **Font Choices:** (typefaces, sizes, weights, leading)
# -   **Color Palettes:** (brand colors, usage in UI)
# -   **Logo Usage:** (placement, sizing, clear space)
# -   **Image or Photography Guidelines:** (types of images, resolutions, content)
# -   **Iconography:** (style, usage)
# -   **Display-Specific Instructions:** (e.g., UI elements, button styles, interactive component behavior)
# -   **Any technical or programming-specific guidelines.**

# **Output Format:**
# -   If you find any relevant grammatical or writing style rules, present them as a clear, concise list or paragraph. Preserve the original phrasing as much as possible, but reformat for readability if necessary.
# -   If **no** grammatical or writing style rules are found in the provided text, respond with the exact string "NO_GRAMMAR_RULES_FOUND".
# -   Do NOT include any introductory or conversational text, just the extracted rules or the "NO_GRAMMAR_RULES_FOUND" string.

# ---
# **TEXT FROM STYLE GUIDE PAGE (Page {page_number}):**
# {page_text}
# ---

# **Extracted Grammatical and Writing Style Rules:**
# """

#     messages = [
#         {"role": "system",
#          "content": "You are a highly analytical AI specialized in distilling specific types of information from complex documents."},
#         {"role": "user", "content": prompt_instruction}
#     ]

#     try:
#         response = client.chat.completions.create(
#             model=deployment_name,
#             messages=messages,
#             temperature=0.1,  # Keep it low for factual extraction
#             max_tokens=2000  # Allow enough tokens for detailed rules
#         )
#         extracted_content = response.choices[0].message.content.strip()

#         if extracted_content == "NO_GRAMMAR_RULES_FOUND":
#             print(f"    - No grammatical rules found on page {page_number}.")
#             return ""
#         else:
#             print(f"    - Grammatical rules extracted from page {page_number}.")
#             return extracted_content

#     except Exception as e:
#         print(f"    - Error processing page {page_number} with LLM: {e}")
#         return f"\n--- Error processing Page {page_number} ---\n{e}\n"
    

# def save_rules_to_txt(rules_text: str, output_file_path: str):
#     """
#     Saves the extracted style guide rules to a plain text file.

#     Args:
#         rules_text (str): The content to be saved.
#         output_file_path (str): The path to the output .txt file.
#     """
#     try:
#         with open(output_file_path, 'w', encoding='utf-8') as f:
#             f.write(rules_text)
#         print(f"\n✅ Successfully saved extracted rules to '{output_file_path}'")
#     except Exception as e:
#         print(f"Error saving rules to '{output_file_path}': {e}")

# def save_rules_to_adls(rules_text: str, client: str, pdf_filename: str) -> bool:
#     """
#     Save extracted rules as TXT file in ADLS with name matching the PDF.
    
#     Args:
#         rules_text: Extracted rules content
#         client: Client name
#         pdf_filename: Original PDF filename
        
#     Returns:
#         True if successful, False otherwise
#     """
#     try:
#         if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
#             logger.error("Missing Azure Storage configuration.")
#             return False

#         # Generate TXT filename from PDF name
#         txt_filename = os.path.splitext(pdf_filename)[0] + ".txt"
#         file_path = f"{client}/{txt_filename}"
        
#         # Setup ADLS client
#         account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
#         service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
#         file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
#         # Upload TXT content
#         file_client = file_system_client.get_file_client(file_path)
#         file_client.upload_data(rules_text.encode('utf-8'), overwrite=True)
#         logger.info(f"Successfully saved rules to: {file_path}")
#         return True
#     except Exception as e:
#         logger.exception(f"Error saving rules to ADLS for client '{client}': {e}")
#         return False


# def get_rules_from_adls(client: str, pdf_filename: str) -> str:
#     """
#     Retrieve rules from the corresponding TXT file in ADLS.
    
#     Args:
#         client: Client name
#         pdf_filename: PDF filename whose rules to retrieve
        
#     Returns:
#         The rules text content or empty string if not found
#     """
#     try:
#         if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
#             logger.error("Missing Azure Storage configuration.")
#             return ""

#         # Generate TXT filename from PDF name
#         txt_filename = os.path.splitext(pdf_filename)[0] + ".txt"
#         file_path = f"{client}/{txt_filename}"
        
#         # Setup ADLS client
#         account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
#         service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
#         file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
#         # Download TXT content
#         file_client = file_system_client.get_file_client(file_path)
#         download = file_client.download_file()
#         content = download.readall().decode('utf-8')
        
#         logger.info(f"Successfully retrieved rules from: {file_path}")
#         return content
#     except ResourceNotFoundError:
#         logger.warning(f"Rules file not found: {file_path}")
#         return ""
#     except Exception as e:
#         logger.exception(f"Error getting rules from ADLS for client '{client}': {e}")
#         return ""

# def process_pdf_and_extract_rules(temp_pdf_path: str, original_filename: str) -> str:
#     """
#     Main pipeline: Extract text from PDF, call Azure OpenAI for rules,
#     and return the rules as a string.
#     """
#     # Load environment
#     load_dotenv()
#     api_key = os.getenv("AZURE_OPENAI_API_KEY")
#     azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
#     api_version = os.getenv("AZURE_OPENAI_API_VERSION")
#     deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

#     if not all([api_key, azure_endpoint, api_version, deployment_name]):
#         raise RuntimeError("Azure OpenAI credentials missing in environment variables")

#     client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)

#     # Extract text
#     page_texts = extract_text_from_pdf(temp_pdf_path)
#     all_rules = []
#     for i, text in enumerate(page_texts, start=1):
#         rules = generate_grammatical_rules_from_text(client, deployment_name, text, i)
#         if rules:
#             all_rules.append(f"--- Rules from Page {i} ---\n{rules}")

#     final_rules = "\n\n".join(all_rules).strip()
#     if not final_rules:
#         final_rules = "No grammatical or writing style rules were identified."

#     return final_rules


# def load_text_file(file_path: str) -> str | None:
#     """Loads a plain text file."""
#     if not os.path.exists(file_path):
#         print(f"Error: File not found at '{file_path}'")
#         return None
#     try:
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     except Exception as e:
#         print(f"Error loading '{file_path}': {e}")
#         return None


# def apply_rules_to_text(
#     client: AzureOpenAI,
#     deployment_name: str,
#     text: str,
#     grammatical_rules: str,
#     item_type: str,
#     slide_identifier: str
# ) -> str:
#     """Uses Azure OpenAI to apply grammatical rules to text with minimal changes."""
#     if not text.strip():
#         return text
#     if not grammatical_rules.strip():
#         print(f"    - Skipping {item_type} for {slide_identifier}: No rules provided.")
#         return text

#     print(f"    - Applying rules to {item_type} for {slide_identifier}...")

#     prompt_instruction = f"""
# You are an expert editor specializing in strict adherence to style guides.
# Review the following text and **strictly apply ONLY the provided grammatical and writing style rules**.

# RULES:
# {grammatical_rules}

# TEXT ({item_type} for {slide_identifier}):
# {text}

# DIRECTIVES:
# 1. Make **only changes strictly required** by the rules.
# 2. Do NOT rewrite or improve unnecessarily.
# 3. Preserve meaning and structure.
# 4. If fully compliant, return text exactly as-is.
# 5. Output ONLY the corrected text, nothing else.
# """

#     try:
#         response = client.chat.completions.create(
#             model=deployment_name,
#             messages=[
#                 {"role": "system", "content": "You are a meticulous style-guide editor."},
#                 {"role": "user", "content": prompt_instruction}
#             ],
#             temperature=0.1,
#             max_tokens=2000
#         )
#         corrected_text = response.choices[0].message.content.strip()

#         if corrected_text != text:
#             print(f"      -> Changes applied to {item_type} for {slide_identifier}")
#         else:
#             print(f"      -> No changes needed for {item_type} on {slide_identifier}")

#         return corrected_text

#     except Exception as e:
#         print(f"    - Error applying rules to {item_type} for {slide_identifier}: {e}")
#         return text


# def read_storyboard_from_word_document(docx_path: str) -> list[dict] | None:
#     """Reads storyboard slides from a Word DOCX file (each table = one slide)."""
#     if not os.path.exists(docx_path):
#         print(f"Error: Word document not found at '{docx_path}'")
#         return None

#     document = docx.Document(docx_path)
#     storyboards_data = []
#     current_module_title = "Unknown Chapter"

#     fields_to_read = [
#         "Course Title", "Module Title", "Learning Objectives", "Topic", "Screen-type",
#         "On-screen text", "Narration", "Developer Notes", "Duration (min)",
#         "Source Images (base64)", "Questions"
#     ]

#     for element in document.element.body:
#         if element.tag.endswith('p'):  # Paragraph
#             paragraph = docx.text.paragraph.Paragraph(element, document)
#             if "Chapter:" in paragraph.text:
#                 current_module_title = paragraph.text.replace("Chapter:", "").strip()

#         if element.tag.endswith('tbl'):  # Table
#             table = docx.table.Table(element, document)
#             slide_dict = {"Module Title": current_module_title}

#             for row in table.rows:
#                 if len(row.cells) < 2:
#                     continue
#                 field_name = row.cells[0].text.strip()
#                 if field_name in fields_to_read:
#                     slide_dict[field_name] = row.cells[1].text.strip()

#             if slide_dict:
#                 storyboards_data.append(slide_dict)

#     print(f"Read {len(storyboards_data)} slides from '{docx_path}'.")
#     return storyboards_data


# def save_storyboards_to_word_document(storyboards: list[dict], filename: str):
#     """Saves updated storyboard slides to a Word DOCX file."""
#     if not storyboards:
#         print("No storyboards to save.")
#         return

#     document = docx.Document()
#     fields_to_display = [
#         "Course Title", "Module Title", "Learning Objectives", "Topic", "Screen-type",
#         "On-screen text", "Narration", "Developer Notes", "Duration (min)", "Source Images (base64)"
#     ]

#     last_module_title = None
#     for i, storyboard_data in enumerate(storyboards):
#         if i > 0:
#             document.add_page_break()

#         slide_title = storyboard_data.get("Topic", "Untitled Slide")
#         module_title = storyboard_data.get("Module Title", "Unknown Chapter")

#         if module_title != last_module_title:
#             document.add_heading(f"Chapter: {module_title}", level=1)
#             last_module_title = module_title

#         document.add_heading(f"Slide {i + 1}: {slide_title}", level=2)

#         table = document.add_table(rows=len(fields_to_display), cols=2)
#         table.style = 'Table Grid'
#         for row_index, field_name in enumerate(fields_to_display):
#             row_cells = table.rows[row_index].cells
#             row_cells[0].text = field_name
#             row_cells[1].text = str(storyboard_data.get(field_name, ""))

#     document.save(filename)
#     print(f"... Saved updated storyboards to '{filename}'")


# def apply_grammatical_rules_to_storyboard_docx(
#     input_docx: str,
#     rules_txt: str,
#     output_docx: str,
#     num_slides_to_process: int = None
# ):
#     """
#     Main function to process storyboard DOCX with grammatical rules applied.
#     - input_docx: Path to storyboard DOCX
#     - rules_txt: Path to client grammatical rules TXT
#     - output_docx: Path to save updated DOCX
#     - num_slides_to_process: Limit number of slides (None = all)
#     """
#     api_key = os.getenv("AZURE_OPENAI_API_KEY")
#     azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
#     api_version = os.getenv("AZURE_OPENAI_API_VERSION")
#     deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

#     if not all([api_key, azure_endpoint, api_version, deployment_name]):
#         sys.exit("FATAL ERROR: Azure OpenAI credentials missing in .env file.")

#     storyboard_list = read_storyboard_from_word_document(input_docx)
#     if storyboard_list is None:
#         sys.exit(1)

#     grammatical_rules = load_text_file(rules_txt) or ""
#     client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)

#     print(f"\n--- Applying grammatical rules ---")
#     slides_processed = 0
#     for i, slide_data in enumerate(storyboard_list):
#         if num_slides_to_process and slides_processed >= num_slides_to_process:
#             break

#         slide_identifier = f"Slide {i + 1}: {slide_data.get('Topic', 'Untitled')}"
#         print(f"\nProcessing {slide_identifier}...")

#         if "Narration" in slide_data:
#             slide_data["Narration"] = apply_rules_to_text(client, deployment_name, slide_data["Narration"],
#                                                          grammatical_rules, "Narration", slide_identifier)
#         if "On-screen text" in slide_data:
#             slide_data["On-screen text"] = apply_rules_to_text(client, deployment_name, slide_data["On-screen text"],
#                                                               grammatical_rules, "On-screen text", slide_identifier)

#         slides_processed += 1
#         time.sleep(0.5)

#     save_storyboards_to_word_document(storyboard_list, output_docx)
#     print("\n--- Rule application complete ---")



# """
# from my_module import apply_grammatical_rules_to_storyboard_docx

# apply_grammatical_rules_to_storyboard_docx(
#     input_docx="Sales2.docx",
#     rules_txt="client_grammatical_style_rules.txt",
#     output_docx="Sales2_StyleApplied.docx",
#     num_slides_to_process=10  # or None for all slides
# )
# """

# def apply_rules_to_batch(
#         client: AzureOpenAI,
#         deployment_name: str,
#         batch_of_texts: list[dict],
#         grammatical_rules: str
# ) -> list[dict]:
#     """
#     Uses an LLM to apply grammatical rules to a batch of texts in a single API call.
    
#     Args:
#         client: AzureOpenAI client
#         deployment_name: Model deployment name
#         batch_of_texts: List of {"id": "...", "text": "..."} dicts
#         grammatical_rules: Style guide rules to apply
        
#     Returns:
#         Updated list of dicts with rules applied
#     """
#     if not batch_of_texts:
#         return []
#     if not grammatical_rules or not grammatical_rules.strip():
#         logger.warning("Skipping rule application for batch: No grammatical rules provided.")
#         return batch_of_texts

#     logger.info(f"Applying rules to a batch of {len(batch_of_texts)} text items...")
#     input_json_str = json.dumps(batch_of_texts, indent=2)

#     prompt_instruction = f"""
# You are an expert editor specializing in strict adherence to style guides.
# Your task is to review an array of JSON objects, where each object contains an 'id' and a 'text'.
# For each object in the array, you must **strictly apply ONLY the provided grammatical and writing style rules** to its 'text' field.

# **KEY DIRECTIVES FOR EACH TEXT ITEM:**
# 1.  **EXTREMELY MINIMAL CHANGES:** Only make changes that are **ABSOLUTELY NECESSARY** to correct a clear violation of the rules.
# 2.  **DO NOT IMPROVE:** Do NOT rewrite for clarity, conciseness, or flow if it is not a direct requirement of a given rule.
# 3.  **PRESERVE ORIGINAL MEANING:** The core message must be perfectly maintained.
# 4.  **RETURN ORIGINAL IF NO VIOLATION:** If a text completely adheres to the rules, its 'text' field in the output must be **EXACTLY** the same as the input.
# 5.  **OUTPUT FORMAT:** Your response MUST be a valid JSON array of objects, with the exact same structure and 'id' values as the input.

# ---
# **CLIENT GRAMMATICAL AND WRITING STYLE RULES:**
# {grammatical_rules}
# ---

# **INPUT JSON ARRAY TO REVIEW:**
# {input_json_str}
# ---

# **OUTPUT JSON ARRAY (Corrected):**
# """

#     messages = [
#         {"role": "system",
#          "content": "You are a meticulous JSON-processing editor that only makes changes strictly dictated by style rules and always returns valid JSON."},
#         {"role": "user", "content": prompt_instruction}
#     ]

#     try:
#         response = client.chat.completions.create(
#             model=deployment_name,
#             messages=messages,
#             temperature=0.0,
#             max_tokens=4096
#         )
#         response_content = response.choices[0].message.content.strip()

#         # Attempt to find the JSON array in the response
#         json_start = response_content.find('[')
#         json_end = response_content.rfind(']')
#         if json_start == -1 or json_end == -1:
#             raise json.JSONDecodeError("Could not find a JSON array in the LLM response.", response_content, 0)

#         json_str = response_content[json_start:json_end + 1]
#         corrected_batch = json.loads(json_str)

#         logger.info(f"Successfully processed batch.")
#         return corrected_batch

#     except json.JSONDecodeError as e:
#         logger.error(f"Failed to decode JSON from LLM response: {e}")
#         logger.debug(f"Raw LLM Response: {response_content}")
#         return batch_of_texts  # Return original batch on error
#     except Exception as e:
#         logger.error(f"Error applying rules to batch with LLM: {e}")
#         return batch_of_texts  # Return original batch on error


# def find_client_style_guide(client: str) -> tuple[str, str]:
#     """
#     Find any style guide file for a given client.
    
#     Args:
#         client: Client name
        
#     Returns:
#         Tuple of (pdf_filename, rules_text)
#     """
#     try:
#         if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
#             logger.error("Missing Azure Storage configuration.")
#             return "", ""
        
#         # Setup ADLS client
#         account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
#         service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
#         file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
#         # Look for any PDF files in client directory
#         client_dir_path = f"{client}"
#         pdf_files = []
#         txt_files = []
        
#         try:
#             paths = file_system_client.get_paths(path=client_dir_path)
#             for path in paths:
#                 if path.is_directory:
#                     continue
                
#                 file_name = os.path.basename(path.name)
#                 if file_name.lower().endswith('.pdf'):
#                     pdf_files.append(file_name)
#                 elif file_name.lower().endswith('.txt'):
#                     txt_files.append(file_name)
#         except Exception as e:
#             logger.warning(f"Error listing files in client directory: {e}")
#             return "", ""
        
#         if not pdf_files:
#             logger.warning(f"No style guide PDF found for client '{client}'")
#             return "", ""
            
#         # Get the first PDF file found
#         pdf_filename = pdf_files[0]
#         logger.info(f"Found style guide PDF for client '{client}': {pdf_filename}")
        
#         # First try to find a matching TXT file
#         txt_filename = os.path.splitext(pdf_filename)[0] + '.txt'
#         if txt_filename in txt_files:
#             # Get the rules content
#             try:
#                 file_path = f"{client}/{txt_filename}"
#                 file_client = file_system_client.get_file_client(file_path)
#                 download = file_client.download_file()
#                 rules_text = download.readall().decode('utf-8')
#                 logger.info(f"Found and loaded rules from TXT file: {txt_filename}")
#                 return pdf_filename, rules_text
#             except Exception as e:
#                 logger.error(f"Error reading TXT file {txt_filename}: {e}")
                
#         # If no matching TXT found, try any TXT file
#         if txt_files:
#             txt_filename = txt_files[0]
#             try:
#                 file_path = f"{client}/{txt_filename}"
#                 file_client = file_system_client.get_file_client(file_path)
#                 download = file_client.download_file()
#                 rules_text = download.readall().decode('utf-8')
#                 logger.info(f"Found and loaded rules from TXT file: {txt_filename}")
#                 return pdf_filename, rules_text
#             except Exception as e:
#                 logger.error(f"Error reading TXT file {txt_filename}: {e}")
                
#         logger.warning(f"No rules TXT file found for client '{client}'")
#         return pdf_filename, ""
        
#     except Exception as e:
#         logger.exception(f"Error finding style guide for client '{client}': {e}")
#         return "", ""

# def apply_style_guide_to_storyboard(storyboard_list: list, client: str, batch_size: int = 10, job_id: str = None) -> list:
#     """
#     Applies client style guide rules to a generated storyboard.
    
#     Args:
#         storyboard_list: List of storyboard items
#         client: Client name to fetch style guide for
#         batch_size: Number of items to process in a single batch
#         job_id: Optional job ID to update progress status
        
#     Returns:
#         Updated storyboard list with rules applied
#     """
#     try:
#         total_items = len(storyboard_list)
#         if job_id:
#             JobStore.update_job(job_id, {
#                 "status": "processing", 
#                 "message": f"Looking for style guide for client '{client}'",
#                 "progress": 0
#             })
        
#         # Find any style guide and rules for this client
#         pdf_filename, grammatical_rules = find_client_style_guide(client)
        
#         if not pdf_filename:
#             logger.info(f"No style guide found for client '{client}'. Returning original storyboard.")
#             if job_id:
#                 JobStore.update_job(job_id, {
#                     "status": "completed", 
#                     "message": f"No style guide found for client '{client}'",
#                     "progress": 100
#                 })
#             return storyboard_list
            
#         if not grammatical_rules:
#             logger.warning(f"Style guide found ({pdf_filename}) but no rules text available. Returning original storyboard.")
#             if job_id:
#                 JobStore.update_job(job_id, {
#                     "status": "completed", 
#                     "message": f"Style guide found but no rules available",
#                     "progress": 100
#                 })
#             return storyboard_list
            
#         # Set up Azure OpenAI client
#         api_key = os.getenv("AZURE_OPENAI_API_KEY")
#         azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
#         api_version = os.getenv("AZURE_OPENAI_API_VERSION")
#         deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

#         if not all([api_key, azure_endpoint, api_version, deployment_name]):
#             logger.error("Azure OpenAI credentials missing in environment variables")
#             if job_id:
#                 JobStore.update_job(job_id, {
#                     "status": "error", 
#                     "message": "Azure OpenAI credentials missing",
#                     "progress": 100
#                 })
#             return storyboard_list

#         openai_client = AzureOpenAI(
#             api_key=api_key, 
#             azure_endpoint=azure_endpoint, 
#             api_version=api_version
#         )

#         logger.info(f"Applying style guide rules in batches to {len(storyboard_list)} storyboard items")
#         logger.info(f"Using style guide: {pdf_filename}")
#         logger.info(f"Batch size: {batch_size}")
        
#         if job_id:
#             JobStore.update_job(job_id, {
#                 "status": "processing", 
#                 "message": f"Starting style guide application: {pdf_filename}",
#                 "progress": 5
#             })

#         # Process the storyboard list in chunks of batch_size
#         total_batches = (total_items + batch_size - 1) // batch_size
        
#         for i in range(0, len(storyboard_list), batch_size):
#             batch_start_index = i
#             batch_end_index = min(i + batch_size, len(storyboard_list))
#             batch_number = i // batch_size + 1
            
#             # Calculate progress percentage
#             progress = min(5 + (batch_number * 95) // total_batches, 95)
            
#             if job_id:
#                 JobStore.update_job(job_id, {
#                     "status": "processing", 
#                     "message": f"Processing batch {batch_number}/{total_batches} (Items {batch_start_index+1}-{batch_end_index})",
#                     "progress": progress
#                 })

#             logger.info(f"Processing Batch {batch_number} (Items {batch_start_index + 1}-{batch_end_index})")

#             current_batch_slides = storyboard_list[batch_start_index:batch_end_index]
#             texts_to_process_in_batch = []

#             # 1. Gather all texts from the current batch of slides
#             for j, slide_data in enumerate(current_batch_slides):
#                 original_slide_index = batch_start_index + j

#                 narration_text = slide_data.get("Narration")
#                 if narration_text and isinstance(narration_text, str) and narration_text.strip():
#                     texts_to_process_in_batch.append({
#                         "id": f"slide_{original_slide_index}_narration",
#                         "text": narration_text
#                     })

#                 on_screen_text = slide_data.get("On_screen_text")
#                 if on_screen_text and isinstance(on_screen_text, str) and on_screen_text.strip():
#                     texts_to_process_in_batch.append({
#                         "id": f"slide_{original_slide_index}_onscreen",
#                         "text": on_screen_text
#                     })

#             if not texts_to_process_in_batch:
#                 logger.info(f"No text found in batch {batch_number} to process. Skipping.")
#                 continue

#             # 2. Send the entire batch for processing in one API call
#             corrected_batch = apply_rules_to_batch(
#                 openai_client, 
#                 deployment_name, 
#                 texts_to_process_in_batch, 
#                 grammatical_rules
#             )

#             # 3. Map the corrected texts back to the main storyboard list
#             if corrected_batch:
#                 corrected_map = {item['id']: item['text'] for item in corrected_batch}

#                 for original_id in corrected_map.keys():
#                     parts = original_id.split('_')
#                     try:
#                         slide_index = int(parts[1])
#                         field_key = "Narration" if parts[2] == "narration" else "On_screen_text"

#                         original_text = storyboard_list[slide_index].get(field_key)
#                         corrected_text = corrected_map[original_id]

#                         storyboard_list[slide_index][field_key] = corrected_text

#                         if original_text != corrected_text:
#                             logger.info(f"Changes applied to {field_key} for item {slide_index + 1}")
#                     except (IndexError, ValueError) as e:
#                         logger.warning(f"Could not parse ID '{original_id}'. Error: {e}")
#             else:
#                 logger.warning(f"Batch {batch_number} failed. Original text for this batch will be kept.")

#             # Add a small delay between API calls - using the correct time module
#             time.sleep(1)

#         logger.info("Style guide application complete")
        
#         # Final status update
#         if job_id:
            
#             JobStore.update_job(job_id, {
#                 "status": "completed", 
#                 "message": f"Style guide application complete",
#                 "progress": 100
#             })
            
#         return storyboard_list
        
#     except Exception as e:
#         logger.exception(f"Error applying style guide to storyboard: {e}")
#         # Update job status on error
#         if job_id:
#             JobStore.update_job(job_id, {
#                 "status": "error", 
#                 "message": f"Error: {str(e)}",
#                 "progress": 100
#             })
#         # Return original storyboard if there's an error
#         return storyboard_list


"""
Style guide routes for handling client style guide operations.
"""
from azure.core.exceptions import ResourceNotFoundError, AzureError
from core.adls import FILESYSTEM_NAME, STORAGE_ACCOUNT_KEY, STORAGE_ACCOUNT_NAME, get_file_structure
from azure.storage.filedatalake import DataLakeServiceClient
# Fix time import - don't use datetime.time
import time
import os
import sys
import docx
from dotenv import load_dotenv
from fastapi import APIRouter, HTTPException, Form, File, UploadFile, Path
from fastapi.responses import JSONResponse
import fitz
from openai import AzureOpenAI
from app.models.project_models import get_style_guide_info, upload_style_guide
import logging
import json
import traceback
import tempfile
# Add this import to fix the JobStore undefined error
from app.services.job_store import JobStore

logger = logging.getLogger(__name__)

router = APIRouter()
@router.get("/style-guide/{client}")
async def get_style_guide_info_route(client: str = Path(...)):
    """Get style guide information for a client"""
    try:
        # Get style guide info from project_models
        result = get_style_guide_info(client)
        
        if not result.get("exists"):
            raise HTTPException(status_code=404, detail="No style guide found")
        
        # Get rules from the corresponding TXT file
        pdf_filename = result.get("file_name", "")
        if pdf_filename:
            rules_text = get_rules_from_adls(client, pdf_filename)
            result["rules_text"] = rules_text
            result["rules_found"] = bool(rules_text)
        
        return JSONResponse(content=result)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in get_style_guide_info_route: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

# @router.post("/style-guide/upload")
# async def upload_style_guide_route(
#     client: str = Form(...),
#     file: UploadFile = File(...),
#     replace: bool = Form(True),  # Default to replace existing files
# ):
#     """Upload style guide file for a client and extract rules"""
#     try:
#         if not file.filename or not file.filename.lower().endswith(".pdf"):
#             raise HTTPException(status_code=400, detail="Only PDF files are allowed")

#         file_content = await file.read()
#         if len(file_content) > 20 * 1024 * 1024:
#             raise HTTPException(
#                 status_code=400,
#                 detail="File size too large. Maximum 20MB allowed",
#             )

#         # Save temporarily for processing
#         temp_path = f"/tmp/{file.filename}"
#         with open(temp_path, "wb") as f:
#             f.write(file_content)

#         # Upload to ADLS with checks for existing style guides
#         result = upload_style_guide(file.filename, file_content, client, replace)
#         if not result.get("success"):
#             raise HTTPException(status_code=400, detail=result.get("message"))

#         # Process PDF and extract rules
#         try:
#             rules_text = process_pdf_and_extract_rules(temp_path, file.filename)
#             logger.info(f"Extracted rules: {rules_text[:200]}...")  # Log first 200 chars
            
#             # Save rules to ADLS as TXT file
#             if rules_text:
#                 if save_rules_to_adls(rules_text, client, file.filename):
#                     result["rules_saved"] = True
#                     result["rules_text"] = rules_text
#                 else:
#                     result["rules_saved"] = False
#                     result["rules_text"] = None
#                     logger.error("Failed to save rules to ADLS")
#         except Exception as e:
#             logger.error(f"Error extracting rules: {e}")
#             result["rules_saved"] = False
#             result["rules_text"] = None

#         # Cleanup temp PDF
#         os.remove(temp_path)
#         return JSONResponse(content=result)
#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Error in upload_style_guide_route: {e}")
#         raise HTTPException(status_code=500, detail="Internal server error")



# @router.post("/style-guide/upload")
# async def upload_style_guide_route(
#     client: str = Form(...),
#     file: UploadFile = File(...),
#     replace: bool = Form(True),  # Default to replace existing files
# ):
#     """Upload style guide file for a client and extract rules"""
#     temp_path = None

#     try:
#         if not file.filename or not file.filename.lower().endswith(".pdf"):
#             raise HTTPException(status_code=400, detail="Only PDF files are allowed")

#         file_content = await file.read()
#         if len(file_content) > 20 * 1024 * 1024:  # 20MB limit
#             raise HTTPException(
#                 status_code=400,
#                 detail="File size too large. Maximum 20MB allowed.",
#             )

#         # Use a safe temporary directory compatible with Azure
#         temp_dir = tempfile.gettempdir()
#         temp_path = os.path.join(temp_dir, file.filename)

#         # Write file to temp path
#         with open(temp_path, "wb") as f:
#             f.write(file_content)

#         # Upload to ADLS with checks for existing style guides
#         result = upload_style_guide(file.filename, file_content, client, replace)
#         if not result.get("success"):
#             raise HTTPException(status_code=400, detail=result.get("message"))

#         # Process PDF and extract rules
#         try:
#             rules_text = process_pdf_and_extract_rules(temp_path, file.filename)
#             logger.info(f"Extracted rules (first 200 chars): {rules_text[:200]}")

#             if rules_text:
#                 saved = save_rules_to_adls(rules_text, client, file.filename)
#                 result["rules_saved"] = saved
#                 result["rules_text"] = rules_text if saved else None

#                 if not saved:
#                     logger.error("Failed to save extracted rules to ADLS")
#             else:
#                 result["rules_saved"] = False
#                 result["rules_text"] = None
#                 logger.warning("No rules extracted from PDF")

#         except Exception as e:
#             logger.error(f"Error extracting rules: {e}\n{traceback.format_exc()}")
#             result["rules_saved"] = False
#             result["rules_text"] = None

#         return JSONResponse(content=result)

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Unexpected error in upload_style_guide_route: {e}\n{traceback.format_exc()}")
#         raise HTTPException(status_code=500, detail="Internal server error")
#     finally:
#         # Always clean up temp file
#         if temp_path and os.path.exists(temp_path):
#             try:
#                 os.remove(temp_path)
#             except Exception as e:
#                 logger.warning(f"Failed to delete temp file {temp_path}: {e}")

import os
import tempfile
import traceback
import fitz

import asyncio
from fastapi import APIRouter, File, Form, UploadFile, HTTPException
from fastapi.responses import JSONResponse
from azure.storage.filedatalake import DataLakeServiceClient
from azure.core.exceptions import ResourceNotFoundError
from dotenv import load_dotenv
from openai import AzureOpenAI  # or your own wrapper
from typing import List

router = APIRouter()

# Azure Storage configs (should be loaded from env)
STORAGE_ACCOUNT_NAME = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
STORAGE_ACCOUNT_KEY = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
FILESYSTEM_NAME = os.getenv("AZURE_STORAGE_FILESYSTEM_NAME")

# ---------------------------------------------------------------------
# Utility: memory logging
# ---------------------------------------------------------------------


# ---------------------------------------------------------------------
# Route: upload + process PDF
# ---------------------------------------------------------------------
@router.post("/style-guide/upload")
async def upload_style_guide_route(
    client: str = Form(...),
    file: UploadFile = File(...),
    replace: bool = Form(True),
):
    """Upload style guide PDF for a client, extract rules using OpenAI, and store result."""
    temp_path = None

    try:
        # ------------------ Validate file ------------------
        if not file.filename or not file.filename.lower().endswith(".pdf"):
            raise HTTPException(status_code=400, detail="Only PDF files are allowed")

        # Save file in chunks to avoid large memory use
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, file.filename)

        with open(temp_path, "wb") as f:
            while chunk := await file.read(1024 * 1024):  # 1 MB chunks
                f.write(chunk)

        # ------------------ Upload to ADLS ------------------
        result = upload_style_guide(file.filename, temp_path, client, replace)
        if not result.get("success"):
            raise HTTPException(status_code=400, detail=result.get("message"))

        # ------------------ Process PDF ------------------
        rules_text = await process_pdf_and_extract_rules(temp_path, file.filename)
        print(f"Extracted rules (first 200 chars): {rules_text[:200]}")

        # ------------------ Save extracted rules ------------------
        try:
            saved = save_rules_to_adls(rules_text, client, file.filename)
        except Exception as e:
            saved = False
            print(f"Error saving rules to ADLS: {e}")

        result["rules_saved"] = saved
        result["rules_text"] = rules_text if saved else None
        return JSONResponse(content=result)

    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error in upload_style_guide_route: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail="Internal server error")
    finally:
        # Clean up temp file
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception as e:
                print(f"Failed to delete temp file {temp_path}: {e}")

                
# def extract_text_from_pdf(pdf_path: str) -> list[str]:
#     """
#     Extracts text page by page from a PDF document.

#     Args:
#         pdf_path (str): The path to the PDF file.

#     Returns:
#         list[str]: A list where each element is the text content of a page.
#                    Returns an empty list if the file cannot be opened or has no pages.
#     """
#     if not os.path.exists(pdf_path):
#         print(f"Error: PDF file not found at '{pdf_path}'")
#         return []

#     page_texts = []
#     try:
#         document = fitz.open(pdf_path)
#         for page_num in range(len(document)):
#             page = document.load_page(page_num)
#             text = page.get_text("text")
#             page_texts.append(text)
#         document.close()
#     except Exception as e:
#         print(f"Error reading PDF '{pdf_path}': {e}")
#     return page_texts


def extract_text_from_pdf(pdf_path: str) -> List[str]:
    if not os.path.exists(pdf_path):
        print(f"Error: PDF not found at {pdf_path}")
        return []

    page_texts = []
    try:
        doc = fitz.open(pdf_path)
        for i in range(len(doc)):
            text = doc.load_page(i).get_text("text")
            page_texts.append(text)
        doc.close()
    except Exception as e:
        print(f"Error reading PDF '{pdf_path}': {e}")
    return page_texts


def generate_grammatical_rules_from_text(
        client: AzureOpenAI,
        deployment_name: str,
        page_text: str,
        page_number: int  # For logging/context
) -> str:
    """
    Uses an LLM to extract only grammatical and writing style rules from a given text chunk.

    Args:
        client (AzureOpenAI): The Azure OpenAI client instance.
        deployment_name (str): The name of the Azure OpenAI deployment.
        page_text (str): The text content from a page of the style guide.
        page_number (int): The page number for logging.

    Returns:
        str: A string containing the extracted rules, or an empty string if no relevant rules
             were found or an error occurred.
    """
    if not page_text.strip():
        return ""  # No text to process

    print(f"  - Analyzing page {page_number} with LLM...")

    # The prompt explicitly tells the LLM what to focus on and what to ignore.
    prompt_instruction = f"""
You are an expert editor and linguist, specializing in analyzing detailed style guides.
Your task is to review the following text from a client's style guide and **extract ONLY the grammatical and writing style rules**.

**Specifically, focus on rules related to:**
-   **Punctuation:** (commas, periods, semicolons, dashes, apostrophes, quotation marks, parentheses, etc.)
-   **Capitalization:** (e.g., proper nouns, titles, headings, sentence beginnings)
-   **Spelling:** (e.g., preference for UK vs. US English, specific terms, compound words)
-   **Sentence Structure and Syntax:** (e.g., clear and concise sentences, avoiding jargon, parallel structure)
-   **Word Choice and Terminology:** (e.g., preferred terms, forbidden words, consistent vocabulary)
-   **Active/Passive Voice:** (guidelines on when to use each)
-   **Tone of Voice:** (e.g., formal, informal, empathetic, direct, authoritative)
-   **General Grammar:** (verb tenses, subject-verb agreement, modifiers, pronouns)
-   **Abbreviations and Acronyms:** (first use expansion, consistent usage)
-   **Number Formatting:** (when to use figures vs. words, percentages, dates)
-   **Formatting for lists:** (bullet points, numbered lists, consistency)
-   **Accessibility considerations for text:** (e.g., clear language, avoidance of complex sentences)

**Crucially, you MUST IGNORE and filter out ANY rules related to:**
-   **Visual Design:** (layout, spacing, white space, margins)
-   **Font Choices:** (typefaces, sizes, weights, leading)
-   **Color Palettes:** (brand colors, usage in UI)
-   **Logo Usage:** (placement, sizing, clear space)
-   **Image or Photography Guidelines:** (types of images, resolutions, content)
-   **Iconography:** (style, usage)
-   **Display-Specific Instructions:** (e.g., UI elements, button styles, interactive component behavior)
-   **Any technical or programming-specific guidelines.**

**Output Format:**
-   If you find any relevant grammatical or writing style rules, present them as a clear, concise list or paragraph. Preserve the original phrasing as much as possible, but reformat for readability if necessary.
-   If **no** grammatical or writing style rules are found in the provided text, respond with the exact string "NO_GRAMMAR_RULES_FOUND".
-   Do NOT include any introductory or conversational text, just the extracted rules or the "NO_GRAMMAR_RULES_FOUND" string.

---
**TEXT FROM STYLE GUIDE PAGE (Page {page_number}):**
{page_text}
---

**Extracted Grammatical and Writing Style Rules:**
"""

    messages = [
        {"role": "system",
         "content": "You are a highly analytical AI specialized in distilling specific types of information from complex documents."},
        {"role": "user", "content": prompt_instruction}
    ]

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=messages,
            temperature=0.1,  # Keep it low for factual extraction
            max_tokens=2000  # Allow enough tokens for detailed rules
        )
        extracted_content = response.choices[0].message.content.strip()

        if extracted_content == "NO_GRAMMAR_RULES_FOUND":
            print(f"    - No grammatical rules found on page {page_number}.")
            return ""
        else:
            print(f"    - Grammatical rules extracted from page {page_number}.")
            return extracted_content

    except Exception as e:
        print(f"    - Error processing page {page_number} with LLM: {e}")
        return f"\n--- Error processing Page {page_number} ---\n{e}\n"
    

def save_rules_to_txt(rules_text: str, output_file_path: str):
    """
    Saves the extracted style guide rules to a plain text file.

    Args:
        rules_text (str): The content to be saved.
        output_file_path (str): The path to the output .txt file.
    """
    try:
        with open(output_file_path, 'w', encoding='utf-8') as f:
            f.write(rules_text)
        print(f"\n✅ Successfully saved extracted rules to '{output_file_path}'")
    except Exception as e:
        print(f"Error saving rules to '{output_file_path}': {e}")

def save_rules_to_adls(rules_text: str, client: str, pdf_filename: str) -> bool:
    """
    Save extracted rules as TXT file in ADLS with name matching the PDF.
    
    Args:
        rules_text: Extracted rules content
        client: Client name
        pdf_filename: Original PDF filename
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
            logger.error("Missing Azure Storage configuration.")
            return False

        # Generate TXT filename from PDF name
        txt_filename = os.path.splitext(pdf_filename)[0] + ".txt"
        file_path = f"{client}/{txt_filename}"
        
        # Setup ADLS client
        account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
        service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
        file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
        # Upload TXT content
        file_client = file_system_client.get_file_client(file_path)
        file_client.upload_data(rules_text.encode('utf-8'), overwrite=True)
        logger.info(f"Successfully saved rules to: {file_path}")
        return True
    except Exception as e:
        logger.exception(f"Error saving rules to ADLS for client '{client}': {e}")
        return False


def get_rules_from_adls(client: str, pdf_filename: str) -> str:
    """
    Retrieve rules from the corresponding TXT file in ADLS.
    
    Args:
        client: Client name
        pdf_filename: PDF filename whose rules to retrieve
        
    Returns:
        The rules text content or empty string if not found
    """
    try:
        if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
            logger.error("Missing Azure Storage configuration.")
            return ""

        # Generate TXT filename from PDF name
        txt_filename = os.path.splitext(pdf_filename)[0] + ".txt"
        file_path = f"{client}/{txt_filename}"
        
        # Setup ADLS client
        account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
        service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
        file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
        # Download TXT content
        file_client = file_system_client.get_file_client(file_path)
        download = file_client.download_file()
        content = download.readall().decode('utf-8')
        
        logger.info(f"Successfully retrieved rules from: {file_path}")
        return content
    except ResourceNotFoundError:
        logger.warning(f"Rules file not found: {file_path}")
        return ""
    except Exception as e:
        logger.exception(f"Error getting rules from ADLS for client '{client}': {e}")
        return ""

# def process_pdf_and_extract_rules(temp_pdf_path: str, original_filename: str) -> str:
#     """
#     Main pipeline: Extract text from PDF, call Azure OpenAI for rules,
#     and return the rules as a string.
#     """
#     # Load environment
#     load_dotenv()
#     api_key = os.getenv("AZURE_OPENAI_API_KEY")
#     azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
#     api_version = os.getenv("AZURE_OPENAI_API_VERSION")
#     deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

#     if not all([api_key, azure_endpoint, api_version, deployment_name]):
#         raise RuntimeError("Azure OpenAI credentials missing in environment variables")

#     client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)

#     # Extract text
#     page_texts = extract_text_from_pdf(temp_pdf_path)
#     all_rules = []
#     for i, text in enumerate(page_texts, start=1):
#         rules = generate_grammatical_rules_from_text(client, deployment_name, text, i)
#         if rules:
#             all_rules.append(f"--- Rules from Page {i} ---\n{rules}")

#     final_rules = "\n\n".join(all_rules).strip()
#     if not final_rules:
#         final_rules = "No grammatical or writing style rules were identified."

#     return final_rules


async def process_pdf_and_extract_rules(temp_pdf_path: str, original_filename: str) -> str:
    load_dotenv()
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION")
    deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

    if not all([api_key, azure_endpoint, api_version, deployment_name]):
        raise RuntimeError("Azure OpenAI credentials missing in environment variables")

    client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)

   
    page_texts = extract_text_from_pdf(temp_pdf_path)
    

    if not page_texts:
        return "No text extracted from PDF."

    # Combine pages into 1-3 chunks to reduce OpenAI calls
    total_pages = len(page_texts)
    chunk_size = max(1, total_pages // 3)
    chunks = ["\n\n".join(page_texts[i:i + chunk_size]) for i in range(0, total_pages, chunk_size)]

  
  

    async def extract_rules_from_chunk(chunk_text: str, idx: int):
        try:
            prompt = (
                "You are an editor extracting grammatical and writing style rules from the following text.\n"
                "Return concise bullet points of rules found.\n\n"
                f"--- Text chunk {idx} ---\n{chunk_text}"
            )
            response = await client.chat.completions.create(
                model=deployment_name,
                messages=[
                    {"role": "system", "content": "You are an expert style guide extractor."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=1500,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"OpenAI call failed: {e}")
            return ""

    tasks = [extract_rules_from_chunk(text, i + 1) for i, text in enumerate(chunks)]
    results = await asyncio.gather(*tasks)

  
  

    final_rules = "\n\n".join([r for r in results if r]).strip()
    return final_rules or "No grammatical or writing style rules were identified."


def load_text_file(file_path: str) -> str | None:
    """Loads a plain text file."""
    if not os.path.exists(file_path):
        print(f"Error: File not found at '{file_path}'")
        return None
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error loading '{file_path}': {e}")
        return None


def apply_rules_to_text(
    client: AzureOpenAI,
    deployment_name: str,
    text: str,
    grammatical_rules: str,
    item_type: str,
    slide_identifier: str
) -> str:
    """Uses Azure OpenAI to apply grammatical rules to text with minimal changes."""
    if not text.strip():
        return text
    if not grammatical_rules.strip():
        print(f"    - Skipping {item_type} for {slide_identifier}: No rules provided.")
        return text

    print(f"    - Applying rules to {item_type} for {slide_identifier}...")

    prompt_instruction = f"""
You are an expert editor specializing in strict adherence to style guides.
Review the following text and **strictly apply ONLY the provided grammatical and writing style rules**.

RULES:
{grammatical_rules}

TEXT ({item_type} for {slide_identifier}):
{text}

DIRECTIVES:
1. Make **only changes strictly required** by the rules.
2. Do NOT rewrite or improve unnecessarily.
3. Preserve meaning and structure.
4. If fully compliant, return text exactly as-is.
5. Output ONLY the corrected text, nothing else.
"""

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "You are a meticulous style-guide editor."},
                {"role": "user", "content": prompt_instruction}
            ],
            temperature=0.1,
            max_tokens=2000
        )
        corrected_text = response.choices[0].message.content.strip()

        if corrected_text != text:
            print(f"      -> Changes applied to {item_type} for {slide_identifier}")
        else:
            print(f"      -> No changes needed for {item_type} on {slide_identifier}")

        return corrected_text

    except Exception as e:
        print(f"    - Error applying rules to {item_type} for {slide_identifier}: {e}")
        return text


def read_storyboard_from_word_document(docx_path: str) -> list[dict] | None:
    """Reads storyboard slides from a Word DOCX file (each table = one slide)."""
    if not os.path.exists(docx_path):
        print(f"Error: Word document not found at '{docx_path}'")
        return None

    document = docx.Document(docx_path)
    storyboards_data = []
    current_module_title = "Unknown Chapter"

    fields_to_read = [
        "Course Title", "Module Title", "Learning Objectives", "Topic", "Screen-type",
        "On-screen text", "Narration", "Developer Notes", "Duration (min)",
        "Source Images (base64)", "Questions"
    ]

    for element in document.element.body:
        if element.tag.endswith('p'):  # Paragraph
            paragraph = docx.text.paragraph.Paragraph(element, document)
            if "Chapter:" in paragraph.text:
                current_module_title = paragraph.text.replace("Chapter:", "").strip()

        if element.tag.endswith('tbl'):  # Table
            table = docx.table.Table(element, document)
            slide_dict = {"Module Title": current_module_title}

            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                field_name = row.cells[0].text.strip()
                if field_name in fields_to_read:
                    slide_dict[field_name] = row.cells[1].text.strip()

            if slide_dict:
                storyboards_data.append(slide_dict)

    print(f"Read {len(storyboards_data)} slides from '{docx_path}'.")
    return storyboards_data


def save_storyboards_to_word_document(storyboards: list[dict], filename: str):
    """Saves updated storyboard slides to a Word DOCX file."""
    if not storyboards:
        print("No storyboards to save.")
        return

    document = docx.Document()
    fields_to_display = [
        "Course Title", "Module Title", "Learning Objectives", "Topic", "Screen-type",
        "On-screen text", "Narration", "Developer Notes", "Duration (min)", "Source Images (base64)"
    ]

    last_module_title = None
    for i, storyboard_data in enumerate(storyboards):
        if i > 0:
            document.add_page_break()

        slide_title = storyboard_data.get("Topic", "Untitled Slide")
        module_title = storyboard_data.get("Module Title", "Unknown Chapter")

        if module_title != last_module_title:
            document.add_heading(f"Chapter: {module_title}", level=1)
            last_module_title = module_title

        document.add_heading(f"Slide {i + 1}: {slide_title}", level=2)

        table = document.add_table(rows=len(fields_to_display), cols=2)
        table.style = 'Table Grid'
        for row_index, field_name in enumerate(fields_to_display):
            row_cells = table.rows[row_index].cells
            row_cells[0].text = field_name
            row_cells[1].text = str(storyboard_data.get(field_name, ""))

    document.save(filename)
    print(f"... Saved updated storyboards to '{filename}'")


def apply_grammatical_rules_to_storyboard_docx(
    input_docx: str,
    rules_txt: str,
    output_docx: str,
    num_slides_to_process: int = None
):
    """
    Main function to process storyboard DOCX with grammatical rules applied.
    - input_docx: Path to storyboard DOCX
    - rules_txt: Path to client grammatical rules TXT
    - output_docx: Path to save updated DOCX
    - num_slides_to_process: Limit number of slides (None = all)
    """
    api_key = os.getenv("AZURE_OPENAI_API_KEY")
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    api_version = os.getenv("AZURE_OPENAI_API_VERSION")
    deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

    if not all([api_key, azure_endpoint, api_version, deployment_name]):
        sys.exit("FATAL ERROR: Azure OpenAI credentials missing in .env file.")

    storyboard_list = read_storyboard_from_word_document(input_docx)
    if storyboard_list is None:
        sys.exit(1)

    grammatical_rules = load_text_file(rules_txt) or ""
    client = AzureOpenAI(api_key=api_key, azure_endpoint=azure_endpoint, api_version=api_version)

    print(f"\n--- Applying grammatical rules ---")
    slides_processed = 0
    for i, slide_data in enumerate(storyboard_list):
        if num_slides_to_process and slides_processed >= num_slides_to_process:
            break

        slide_identifier = f"Slide {i + 1}: {slide_data.get('Topic', 'Untitled')}"
        print(f"\nProcessing {slide_identifier}...")

        if "Narration" in slide_data:
            slide_data["Narration"] = apply_rules_to_text(client, deployment_name, slide_data["Narration"],
                                                         grammatical_rules, "Narration", slide_identifier)
        if "On-screen text" in slide_data:
            slide_data["On-screen text"] = apply_rules_to_text(client, deployment_name, slide_data["On-screen text"],
                                                              grammatical_rules, "On-screen text", slide_identifier)

        slides_processed += 1
        time.sleep(0.5)

    save_storyboards_to_word_document(storyboard_list, output_docx)
    print("\n--- Rule application complete ---")



"""
from my_module import apply_grammatical_rules_to_storyboard_docx

apply_grammatical_rules_to_storyboard_docx(
    input_docx="Sales2.docx",
    rules_txt="client_grammatical_style_rules.txt",
    output_docx="Sales2_StyleApplied.docx",
    num_slides_to_process=10  # or None for all slides
)
"""

def apply_rules_to_batch(
        client: AzureOpenAI,
        deployment_name: str,
        batch_of_texts: list[dict],
        grammatical_rules: str
) -> list[dict]:
    """
    Uses an LLM to apply grammatical rules to a batch of texts in a single API call.
    
    Args:
        client: AzureOpenAI client
        deployment_name: Model deployment name
        batch_of_texts: List of {"id": "...", "text": "..."} dicts
        grammatical_rules: Style guide rules to apply
        
    Returns:
        Updated list of dicts with rules applied
    """
    if not batch_of_texts:
        return []
    if not grammatical_rules or not grammatical_rules.strip():
        logger.warning("Skipping rule application for batch: No grammatical rules provided.")
        return batch_of_texts

    logger.info(f"Applying rules to a batch of {len(batch_of_texts)} text items...")
    input_json_str = json.dumps(batch_of_texts, indent=2)

    prompt_instruction = f"""
You are an expert editor specializing in strict adherence to style guides.
Your task is to review an array of JSON objects, where each object contains an 'id' and a 'text'.
For each object in the array, you must **strictly apply ONLY the provided grammatical and writing style rules** to its 'text' field.

**KEY DIRECTIVES FOR EACH TEXT ITEM:**
1.  **EXTREMELY MINIMAL CHANGES:** Only make changes that are **ABSOLUTELY NECESSARY** to correct a clear violation of the rules.
2.  **DO NOT IMPROVE:** Do NOT rewrite for clarity, conciseness, or flow if it is not a direct requirement of a given rule.
3.  **PRESERVE ORIGINAL MEANING:** The core message must be perfectly maintained.
4.  **RETURN ORIGINAL IF NO VIOLATION:** If a text completely adheres to the rules, its 'text' field in the output must be **EXACTLY** the same as the input.
5.  **OUTPUT FORMAT:** Your response MUST be a valid JSON array of objects, with the exact same structure and 'id' values as the input.

---
**CLIENT GRAMMATICAL AND WRITING STYLE RULES:**
{grammatical_rules}
---

**INPUT JSON ARRAY TO REVIEW:**
{input_json_str}
---

**OUTPUT JSON ARRAY (Corrected):**
"""

    messages = [
        {"role": "system",
         "content": "You are a meticulous JSON-processing editor that only makes changes strictly dictated by style rules and always returns valid JSON."},
        {"role": "user", "content": prompt_instruction}
    ]

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=messages,
            temperature=0.0,
            max_tokens=4096
        )
        response_content = response.choices[0].message.content.strip()

        # Attempt to find the JSON array in the response
        json_start = response_content.find('[')
        json_end = response_content.rfind(']')
        if json_start == -1 or json_end == -1:
            raise json.JSONDecodeError("Could not find a JSON array in the LLM response.", response_content, 0)

        json_str = response_content[json_start:json_end + 1]
        corrected_batch = json.loads(json_str)

        logger.info(f"Successfully processed batch.")
        return corrected_batch

    except json.JSONDecodeError as e:
        logger.error(f"Failed to decode JSON from LLM response: {e}")
        logger.debug(f"Raw LLM Response: {response_content}")
        return batch_of_texts  # Return original batch on error
    except Exception as e:
        logger.error(f"Error applying rules to batch with LLM: {e}")
        return batch_of_texts  # Return original batch on error


def find_client_style_guide(client: str) -> tuple[str, str]:
    """
    Find any style guide file for a given client.
    
    Args:
        client: Client name
        
    Returns:
        Tuple of (pdf_filename, rules_text)
    """
    try:
        if not all([STORAGE_ACCOUNT_NAME, STORAGE_ACCOUNT_KEY, FILESYSTEM_NAME]):
            logger.error("Missing Azure Storage configuration.")
            return "", ""
        
        # Setup ADLS client
        account_url = f"https://{STORAGE_ACCOUNT_NAME}.dfs.core.windows.net"
        service_client = DataLakeServiceClient(account_url=account_url, credential=STORAGE_ACCOUNT_KEY)
        file_system_client = service_client.get_file_system_client(file_system=FILESYSTEM_NAME)
        
        # Look for any PDF files in client directory
        client_dir_path = f"{client}"
        pdf_files = []
        txt_files = []
        
        try:
            paths = file_system_client.get_paths(path=client_dir_path)
            for path in paths:
                if path.is_directory:
                    continue
                
                file_name = os.path.basename(path.name)
                if file_name.lower().endswith('.pdf'):
                    pdf_files.append(file_name)
                elif file_name.lower().endswith('.txt'):
                    txt_files.append(file_name)
        except Exception as e:
            logger.warning(f"Error listing files in client directory: {e}")
            return "", ""
        
        if not pdf_files:
            logger.warning(f"No style guide PDF found for client '{client}'")
            return "", ""
            
        # Get the first PDF file found
        pdf_filename = pdf_files[0]
        logger.info(f"Found style guide PDF for client '{client}': {pdf_filename}")
        
        # First try to find a matching TXT file
        txt_filename = os.path.splitext(pdf_filename)[0] + '.txt'
        if txt_filename in txt_files:
            # Get the rules content
            try:
                file_path = f"{client}/{txt_filename}"
                file_client = file_system_client.get_file_client(file_path)
                download = file_client.download_file()
                rules_text = download.readall().decode('utf-8')
                logger.info(f"Found and loaded rules from TXT file: {txt_filename}")
                return pdf_filename, rules_text
            except Exception as e:
                logger.error(f"Error reading TXT file {txt_filename}: {e}")
                
        # If no matching TXT found, try any TXT file
        if txt_files:
            txt_filename = txt_files[0]
            try:
                file_path = f"{client}/{txt_filename}"
                file_client = file_system_client.get_file_client(file_path)
                download = file_client.download_file()
                rules_text = download.readall().decode('utf-8')
                logger.info(f"Found and loaded rules from TXT file: {txt_filename}")
                return pdf_filename, rules_text
            except Exception as e:
                logger.error(f"Error reading TXT file {txt_filename}: {e}")
                
        logger.warning(f"No rules TXT file found for client '{client}'")
        return pdf_filename, ""
        
    except Exception as e:
        logger.exception(f"Error finding style guide for client '{client}': {e}")
        return "", ""

def apply_style_guide_to_storyboard(storyboard_list: list, client: str, batch_size: int = 10, job_id: str = None) -> list:
    """
    Applies client style guide rules to a generated storyboard.
    
    Args:
        storyboard_list: List of storyboard items
        client: Client name to fetch style guide for
        batch_size: Number of items to process in a single batch
        job_id: Optional job ID to update progress status
        
    Returns:
        Updated storyboard list with rules applied
    """
    try:
        total_items = len(storyboard_list)
        if job_id:
            JobStore.update_job(job_id, {
                "status": "processing", 
                "message": f"Looking for style guide for client '{client}'",
                "progress": 0
            })
        
        # Find any style guide and rules for this client
        pdf_filename, grammatical_rules = find_client_style_guide(client)
        
        if not pdf_filename:
            logger.info(f"No style guide found for client '{client}'. Returning original storyboard.")
            if job_id:
                JobStore.update_job(job_id, {
                    "status": "completed", 
                    "message": f"No style guide found for client '{client}'",
                    "progress": 100
                })
            return storyboard_list
            
        if not grammatical_rules:
            logger.warning(f"Style guide found ({pdf_filename}) but no rules text available. Returning original storyboard.")
            if job_id:
                JobStore.update_job(job_id, {
                    "status": "completed", 
                    "message": f"Style guide found but no rules available",
                    "progress": 100
                })
            return storyboard_list
            
        # Set up Azure OpenAI client
        api_key = os.getenv("AZURE_OPENAI_API_KEY")
        azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        api_version = os.getenv("AZURE_OPENAI_API_VERSION")
        deployment_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

        if not all([api_key, azure_endpoint, api_version, deployment_name]):
            logger.error("Azure OpenAI credentials missing in environment variables")
            if job_id:
                JobStore.update_job(job_id, {
                    "status": "error", 
                    "message": "Azure OpenAI credentials missing",
                    "progress": 100
                })
            return storyboard_list

        openai_client = AzureOpenAI(
            api_key=api_key, 
            azure_endpoint=azure_endpoint, 
            api_version=api_version
        )

        logger.info(f"Applying style guide rules in batches to {len(storyboard_list)} storyboard items")
        logger.info(f"Using style guide: {pdf_filename}")
        logger.info(f"Batch size: {batch_size}")
        
        if job_id:
            JobStore.update_job(job_id, {
                "status": "processing", 
                "message": f"Starting style guide application: {pdf_filename}",
                "progress": 5
            })

        # Process the storyboard list in chunks of batch_size
        total_batches = (total_items + batch_size - 1) // batch_size
        
        for i in range(0, len(storyboard_list), batch_size):
            batch_start_index = i
            batch_end_index = min(i + batch_size, len(storyboard_list))
            batch_number = i // batch_size + 1
            
            # Calculate progress percentage
            progress = min(5 + (batch_number * 95) // total_batches, 95)
            
            if job_id:
                JobStore.update_job(job_id, {
                    "status": "processing", 
                    "message": f"Processing batch {batch_number}/{total_batches} (Items {batch_start_index+1}-{batch_end_index})",
                    "progress": progress
                })

            logger.info(f"Processing Batch {batch_number} (Items {batch_start_index + 1}-{batch_end_index})")

            current_batch_slides = storyboard_list[batch_start_index:batch_end_index]
            texts_to_process_in_batch = []

            # 1. Gather all texts from the current batch of slides
            for j, slide_data in enumerate(current_batch_slides):
                original_slide_index = batch_start_index + j

                narration_text = slide_data.get("Narration")
                if narration_text and isinstance(narration_text, str) and narration_text.strip():
                    texts_to_process_in_batch.append({
                        "id": f"slide_{original_slide_index}_narration",
                        "text": narration_text
                    })

                on_screen_text = slide_data.get("On_screen_text")
                if on_screen_text and isinstance(on_screen_text, str) and on_screen_text.strip():
                    texts_to_process_in_batch.append({
                        "id": f"slide_{original_slide_index}_onscreen",
                        "text": on_screen_text
                    })

            if not texts_to_process_in_batch:
                logger.info(f"No text found in batch {batch_number} to process. Skipping.")
                continue

            # 2. Send the entire batch for processing in one API call
            corrected_batch = apply_rules_to_batch(
                openai_client, 
                deployment_name, 
                texts_to_process_in_batch, 
                grammatical_rules
            )

            # 3. Map the corrected texts back to the main storyboard list
            if corrected_batch:
                corrected_map = {item['id']: item['text'] for item in corrected_batch}

                for original_id in corrected_map.keys():
                    parts = original_id.split('_')
                    try:
                        slide_index = int(parts[1])
                        field_key = "Narration" if parts[2] == "narration" else "On_screen_text"

                        original_text = storyboard_list[slide_index].get(field_key)
                        corrected_text = corrected_map[original_id]

                        storyboard_list[slide_index][field_key] = corrected_text

                        if original_text != corrected_text:
                            logger.info(f"Changes applied to {field_key} for item {slide_index + 1}")
                    except (IndexError, ValueError) as e:
                        logger.warning(f"Could not parse ID '{original_id}'. Error: {e}")
            else:
                logger.warning(f"Batch {batch_number} failed. Original text for this batch will be kept.")

            # Add a small delay between API calls - using the correct time module
            time.sleep(1)

        logger.info("Style guide application complete")
        
        # Final status update
        if job_id:
            
            JobStore.update_job(job_id, {
                "status": "completed", 
                "message": f"Style guide application complete",
                "progress": 100
            })
            
        return storyboard_list
        
    except Exception as e:
        logger.exception(f"Error applying style guide to storyboard: {e}")
        # Update job status on error
        if job_id:
            JobStore.update_job(job_id, {
                "status": "error", 
                "message": f"Error: {str(e)}",
                "progress": 100
            })
        # Return original storyboard if there's an error
        return storyboard_list