"""
co_content_renderer.py — Markdown → Word formatting renderer
==============================================================
Patch for: Storyboard-backend-main/core/save_adls.py

The PROBLEM this fixes
======================
The extractors output content with markdown formatting markers:
  **bold text**       → should render as actual Word bold
  *italic text*       → should render as actual Word italic
  ***bold italic***   → should render as Word bold+italic
  <RED>text</RED>     → should render as red Word font
  <BLUE>text</BLUE>   → should render as blue Word font
  [link](url)         → should render as Word hyperlink
  [TABLE]\n|...|...   → should render as actual Word table

The CURRENT save_adls.py just dumps the raw markdown text into a paragraph,
so reviewers see literal "**bold**" with the asterisks visible. They flagged
this 10+ times: "These stars are not present in the source document."

This module replaces the cell-content rendering with a proper markdown-aware
renderer that:
  1. Parses [TABLE] markdown blocks and creates Word tables
  2. Parses **bold**, *italic*, color markers and applies real Word formatting
  3. Strips ALL marker syntax — nothing leaks as visible text
  4. Filters footer artifacts (Section Break, Sample Footer Text, dates)
  5. Suppresses [IMAGE-ONLY SLIDE] markers (they're internal flags, not content)

INTEGRATION
===========
In save_adls.py, find where the "Full Page Content" cell is populated.
Currently it does something like:
  cell.text = chunk_data.get('content', '')
  OR
  cell.add_paragraph(chunk_data.get('content', ''))

Replace those calls with:
  from co_content_renderer import render_content_into_cell
  render_content_into_cell(cell, chunk_data.get('content', ''))

That's it. The function takes a python-docx _Cell object and the markdown text,
and writes properly formatted Word content into the cell.
"""

import re
from typing import List, Tuple, Optional


# ══════════════════════════════════════════════════════════════════════════════
# COLOR DEFINITIONS
# ══════════════════════════════════════════════════════════════════════════════

# RGB tuples for Word color formatting
_COLOR_RED = (192, 0, 0)      # C00000 — standard "red" in Word
_COLOR_BLUE = (0, 112, 192)   # 0070C0 — standard "blue" in Word
_COLOR_GRAY = (128, 128, 128) # 808080 — used for image markers


# ══════════════════════════════════════════════════════════════════════════════
# INLINE FORMATTING PARSER
# ══════════════════════════════════════════════════════════════════════════════
# Parses a single line of text with inline markdown markers and yields
# (text, formatting) tuples. Formatting is a dict with keys:
#   bold (bool), italic (bool), color (None|'red'|'blue'), url (str)

_INLINE_PATTERNS = [
    # Color tags must be processed FIRST (they're the outermost wrapper)
    # so that bold/italic inside them is handled within their text range.
    # <RED>text</RED>
    (re.compile(r'<RED>(.+?)</RED>', re.DOTALL), {'color': 'red'}),
    # <BLUE>text</BLUE>
    (re.compile(r'<BLUE>(.+?)</BLUE>', re.DOTALL), {'color': 'blue'}),
    # ***bold italic***  (must come BEFORE ** and *)
    (re.compile(r'\*\*\*([^*]+?)\*\*\*'), {'bold': True, 'italic': True}),
    # **bold**
    (re.compile(r'\*\*([^*]+?)\*\*'), {'bold': True}),
    # *italic* (require boundary to avoid false matches inside **bold**)
    (re.compile(r'(?<![\w*])\*([^*\n]+?)\*(?![\w*])'), {'italic': True}),
    # [text](url) — hyperlinks
    (re.compile(r'\[([^\]]+?)\]\(([^)]+?)\)'), {'_link': True}),
]


def parse_inline_runs(text: str) -> List[Tuple[str, dict]]:
    """
    Parse a string of text with markdown markers into a list of
    (text, formatting) tuples.

    Handles nesting: <RED>**bold red**</RED> produces a single run with
    bold=True, color='red'.

    Examples:
      "plain text"         → [("plain text", {})]
      "**bold**"           → [("bold", {"bold": True})]
      "a **b** c"          → [("a ", {}), ("b", {"bold": True}), (" c", {})]
      "<RED>*red italic*</RED>" → [("red italic", {"italic": True, "color": "red"})]
    """
    if not text:
        return []

    # Defensive normalization: collapse 4+ consecutive asterisks into 2.
    # This fixes leftover artifacts like '**Lean ****Development**' which
    # arise when the source had two adjacent bold runs not merged upstream.
    text = re.sub(r'\*{4,}', '**', text)

    # Strategy: recursively peel off the outermost wrapper.
    # We process color tags first (outermost), then bold, italic, links.

    runs = [(text, {})]

    for pattern, formatting in _INLINE_PATTERNS:
        new_runs = []
        for run_text, run_fmt in runs:
            if not run_text:
                continue
            # Find all matches in this run
            last_end = 0
            for m in pattern.finditer(run_text):
                # Text before the match keeps current formatting
                if m.start() > last_end:
                    new_runs.append((run_text[last_end:m.start()], dict(run_fmt)))

                # The matched text gets new formatting layered on
                inner_fmt = dict(run_fmt)
                if '_link' in formatting:
                    inner_fmt['url'] = m.group(2)
                    inner_text = m.group(1)
                else:
                    inner_fmt.update(formatting)
                    inner_text = m.group(1)

                new_runs.append((inner_text, inner_fmt))
                last_end = m.end()

            # Any remaining text after the last match
            if last_end < len(run_text):
                new_runs.append((run_text[last_end:], dict(run_fmt)))

        runs = new_runs

    # Final cleanup: strip orphan markers that survived parsing.
    # These come from malformed source content (e.g., stray `**` with no pair).
    cleaned = []
    for t, f in runs:
        if not t:
            continue
        # Strip stray asterisks from text content (anything not a marker)
        t = re.sub(r'\*+', '', t) if '*' in t else t
        # Strip stray angle bracket tags
        t = t.replace('<RED>', '').replace('</RED>', '').replace('<BLUE>', '').replace('</BLUE>', '')
        if t:
            cleaned.append((t, f))
    return cleaned


# ══════════════════════════════════════════════════════════════════════════════
# CONTENT FILTERING
# ══════════════════════════════════════════════════════════════════════════════

_FILTER_PATTERNS = [
    re.compile(r'^Section Break$', re.I),
    re.compile(r'^Sample Footer Text$', re.I),
    re.compile(r'^AI-generated content may be incorrect\.?$', re.I),
    re.compile(r'^\d{1,2}/\d{1,2}/\d{2,4}$'),
    re.compile(r'^\d{1,2}-\d{1,2}-\d{2,4}$'),
    re.compile(r'^Click to add (text|notes|content|title)$', re.I),
    re.compile(r'^CONFIDENTIAL\b.*INTERNAL USE\b', re.I),
    re.compile(r'^FOR INTERNAL USE ONLY$', re.I),
    re.compile(r'^PROPRIETARY( AND CONFIDENTIAL)?$', re.I),
    re.compile(r'^Page\s*\d+\s*(of\s*\d+)?$', re.I),
    re.compile(r'^\[IMAGE-ONLY SLIDE.*\]$', re.I),
    re.compile(r'^\[CHART.*\]$', re.I),
    re.compile(r'^\[MISSING_IMAGE:.*\]$', re.I),
]


def is_filter_line(line: str) -> bool:
    """Returns True if a line should be dropped (footer/artifact)."""
    if not line:
        return True
    stripped = line.strip()
    if not stripped:
        return False  # Keep blank lines for spacing
    for pat in _FILTER_PATTERNS:
        if pat.match(stripped):
            return True
    return False


# ══════════════════════════════════════════════════════════════════════════════
# TABLE BLOCK PARSER
# ══════════════════════════════════════════════════════════════════════════════
# A [TABLE] block in our markdown looks like:
#   [TABLE]
#   | Header 1 | Header 2 |
#   | --- | --- |
#   | Row 1 Col 1 | Row 1 Col 2 |
#   | Row 2 Col 1 | Row 2 Col 2 |

def parse_table_block(lines: List[str], start_idx: int) -> Tuple[Optional[List[List[str]]], int]:
    """
    Parse a [TABLE] markdown block starting at lines[start_idx].
    Returns (rows_data, end_idx_exclusive) where rows_data is a list of cell-text rows.
    Returns (None, start_idx) if no valid table found.
    """
    if start_idx >= len(lines):
        return None, start_idx

    # Skip the [TABLE] marker line
    i = start_idx
    if lines[i].strip() == '[TABLE]':
        i += 1
    else:
        return None, start_idx

    rows = []
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Skip the separator row (| --- | --- |)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        # Parse pipe-separated cells
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
        i += 1

    if not rows:
        return None, start_idx

    return rows, i


# ══════════════════════════════════════════════════════════════════════════════
# WORD FORMATTING APPLICATION
# ══════════════════════════════════════════════════════════════════════════════

def apply_runs_to_paragraph(paragraph, runs: List[Tuple[str, dict]]):
    """
    Add runs to a python-docx paragraph with formatting applied.
    `paragraph` is a docx.text.paragraph.Paragraph object.
    `runs` is a list of (text, formatting_dict) tuples.
    """
    try:
        from docx.shared import RGBColor
    except ImportError:
        # python-docx not available; fall back to plain text
        paragraph.add_run(''.join(t for t, _ in runs))
        return

    for text, fmt in runs:
        if not text:
            continue
        run = paragraph.add_run(text)
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        color = fmt.get('color')
        if color == 'red':
            run.font.color.rgb = RGBColor(*_COLOR_RED)
        elif color == 'blue':
            run.font.color.rgb = RGBColor(*_COLOR_BLUE)
        # Hyperlinks would need more work — for now just style as blue underline
        if fmt.get('url'):
            run.underline = True
            if not color:
                run.font.color.rgb = RGBColor(*_COLOR_BLUE)


def render_content_into_cell(cell, content: str, images: list = None):
    """
    Main entry point. Render markdown content into a python-docx _Cell.

    Steps:
      1. Clear any existing content in the cell
      2. Split into lines and strip footer artifacts
      3. Detect [TABLE] blocks → create nested Word tables
      4. Detect [IMAGE: id] markers → embed image at this position
      5. Plain lines → add as paragraphs with inline formatting

    Args:
        cell: python-docx Cell object
        content: markdown content string
        images: optional list of dicts with image data to resolve [IMAGE: id]
                markers. Each dict should have 'id' and one of:
                  - 'data' (base64 string) + 'format' (extension), OR
                  - 'image_key' (path) + 'data' (raw bytes), OR
                  - 'content' (base64 string)
                If a marker has no matching image, it's silently dropped.
    """
    # Clear existing cell content
    # python-docx doesn't have a clean clear() — we have to work with paragraphs
    # The cell already has one empty paragraph by default; we'll reuse it.

    if not content:
        return

    # [PATCH] Build image lookup map for [IMAGE: id] resolution.
    # Supports multiple key formats produced by different extractors.
    image_lookup = {}
    if images:
        for img in images:
            if not isinstance(img, dict):
                continue
            # Try every possible key the extractors might use
            for key_field in ('id', 'image_key', 'image_name'):
                k = img.get(key_field)
                if k:
                    image_lookup[str(k).strip().lower()] = img

    lines = content.split('\n')
    # Filter out footer artifacts
    lines = [l for l in lines if not is_filter_line(l)]
    # Collapse 3+ blank lines into 2
    cleaned = []
    blank_run = 0
    for line in lines:
        if not line.strip():
            blank_run += 1
            if blank_run <= 1:
                cleaned.append(line)
        else:
            blank_run = 0
            cleaned.append(line)
    lines = cleaned

    # Process line by line
    first_para_used = False  # Cell starts with one empty paragraph
    i = 0
    while i < len(lines):
        line = lines[i]

        # ── [TABLE] block ──
        if line.strip() == '[TABLE]':
            rows_data, new_i = parse_table_block(lines, i)
            if rows_data:
                _add_table_to_cell(cell, rows_data)
                i = new_i
                continue
            i += 1
            continue

        # ── [IMAGE: id] marker ──
        # Embeds an image at this exact position in the cell. The id is matched
        # against the cell's images list. If not found, the line is dropped.
        # Fixes reviewer comments about images being placed in wrong positions
        # or missing entirely.
        img_match = re.match(r'^\s*\[IMAGE:\s*([^\]]+?)\s*\]\s*$', line, re.IGNORECASE)
        if img_match:
            img_id = img_match.group(1).strip().lower()
            img_obj = image_lookup.get(img_id)
            if img_obj is not None:
                _embed_image_into_cell(cell, img_obj)
                first_para_used = True
            i += 1
            continue

        # ── Blank line ──
        if not line.strip():
            if first_para_used:
                cell.add_paragraph('')
            i += 1
            continue

        # ── Regular line with inline formatting ──
        runs = parse_inline_runs(line)
        if first_para_used:
            para = cell.add_paragraph()
        else:
            # Use the existing first paragraph
            para = cell.paragraphs[0]
            first_para_used = True
        apply_runs_to_paragraph(para, runs)

        i += 1


def _add_table_to_cell(cell, rows_data: List[List[str]]):
    """Add a Word table inside a cell. Each row is a list of cell text strings."""
    if not rows_data:
        return

    n_cols = max(len(r) for r in rows_data)

    # Create a paragraph to anchor the table after
    cell.add_paragraph()
    tbl = cell.add_table(rows=len(rows_data), cols=n_cols)

    try:
        tbl.style = 'Table Grid'
    except KeyError:
        pass  # Style not available

    for ri, row_data in enumerate(rows_data):
        for ci in range(n_cols):
            text = row_data[ci] if ci < len(row_data) else ''
            inner_cell = tbl.rows[ri].cells[ci]
            # Recursively render the cell text (it may have **bold** etc.)
            inner_cell.text = ''  # Clear default
            inner_para = inner_cell.paragraphs[0]
            runs = parse_inline_runs(text)
            apply_runs_to_paragraph(inner_para, runs)
            # Bold the header row
            if ri == 0:
                for run in inner_para.runs:
                    run.bold = True


def _embed_image_into_cell(cell, img_obj: dict, max_width_inches: float = 3.0):
    """
    Embed an image into a Word table cell at the current position.

    Accepts image dicts in any of the formats produced by the extractors:
      - {'id': ..., 'data': base64, 'format': 'png'}        (PPTX native)
      - {'image_key': path, 'image_name': name, ...}         (DOCX with file path)
      - {'content': base64, ...}                              (search_service resolved)
      - {'data': raw_bytes}                                   (raw bytes)

    Silently no-ops on any failure so a bad image doesn't break the whole CO.
    """
    try:
        import io
        import base64
        from docx.shared import Inches

        # Resolve image bytes from whatever format we got
        img_bytes = None

        # Format 1: base64 in 'data' or 'content'
        b64 = img_obj.get('data') or img_obj.get('content') or img_obj.get('base64')
        if isinstance(b64, str) and b64:
            try:
                # Strip data URI prefix if present
                if 'base64,' in b64:
                    b64 = b64.split('base64,', 1)[1]
                img_bytes = base64.b64decode(b64)
            except Exception:
                pass
        elif isinstance(b64, (bytes, bytearray)):
            img_bytes = bytes(b64)

        # Format 2: file path in image_key
        if img_bytes is None:
            path = img_obj.get('image_key') or img_obj.get('path')
            if path:
                try:
                    import os
                    if os.path.exists(path):
                        with open(path, 'rb') as f:
                            img_bytes = f.read()
                except Exception:
                    pass

        if not img_bytes:
            return  # nothing to embed

        # Add image to a new paragraph in the cell
        para = cell.add_paragraph()
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(max_width_inches))
    except Exception as e:
        # Failure to embed an image must never break the whole CO render
        try:
            import logging
            logging.getLogger(__name__).debug(f"_embed_image_into_cell failed: {e}")
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# CONVENIENCE: STRIP MARKERS FOR PLAIN-TEXT FALLBACK
# ══════════════════════════════════════════════════════════════════════════════

def strip_all_markers(text: str) -> str:
    """
    Strip all formatting markers from text, returning plain text.
    Used as a fallback when Word rendering isn't available.
    """
    if not text:
        return ""

    # Remove [TABLE] markers and table separator rows (keep cell text)
    text = re.sub(r'^\[TABLE\]\n?', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\|[\s\-:|]+\|$\n?', '', text, flags=re.MULTILINE)
    # Convert table rows to plain text: "| a | b |" → "a   b"
    text = re.sub(r'^\|\s*(.+?)\s*\|$',
                  lambda m: '   '.join(c.strip() for c in m.group(1).split('|')),
                  text, flags=re.MULTILINE)

    # Remove inline markers
    text = re.sub(r'\*\*\*([^*]+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*([^*]+?)\*\*', r'\1', text)
    text = re.sub(r'(?<![\w*])\*([^*\n]+?)\*(?![\w*])', r'\1', text)
    text = re.sub(r'<RED>(.+?)</RED>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<BLUE>(.+?)</BLUE>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\[([^\]]+?)\]\([^)]+?\)', r'\1', text)
    text = re.sub(r'\[(?:IMAGE-ONLY SLIDE|CHART|MISSING_IMAGE)[^\]]*\]', '', text)

    # Filter footer artifacts line-by-line
    out = [l for l in text.split('\n') if not is_filter_line(l)]
    return '\n'.join(out).strip()
