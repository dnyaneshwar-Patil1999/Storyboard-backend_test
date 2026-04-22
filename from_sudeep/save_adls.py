import io
import os
import base64
import re
import tempfile
import binascii
import json
import logging
from typing import Union
from datetime import datetime
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.document import Document as DocumentType
from docx.table import _Cell
from azure.storage.filedatalake import DataLakeServiceClient
from azure.core.exceptions import AzureError
from azure.storage.blob import BlobServiceClient
import os
from azure.core.exceptions import ServiceResponseError
import io, time, logging

load_dotenv()

AZURE_STORAGE_ACCOUNT_NAME = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
AZURE_STORAGE_ACCOUNT_KEY = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
AZURE_STORAGE_FILESYSTEM_NAME = os.getenv("AZURE_STORAGE_FILESYSTEM_NAME")

logger = logging.getLogger(__name__)

# ── [PATCH] Import markdown-to-Word renderer ──────────────────────────────────
# Renders extractor output markdown (**bold**, *italic*, <RED>, <BLUE>,
# [TABLE], [link](url)) into proper Word formatting inside cells.
# Fixes the "stars not from source" and "italic and red missing" review comments.
# If the module isn't present, we fall back to the legacy renderer.
try:
    from co_content_renderer import render_content_into_cell
    _CO_RENDERER_AVAILABLE = True
except ImportError:
    _CO_RENDERER_AVAILABLE = False
    logger.warning("co_content_renderer not available, using legacy renderer")
# ──────────────────────────────────────────────────────────────────────────────

def get_client_short_name(client_name: str) -> str:
    """
    Returns a client short name based on the following rules:
    - If single word and length >= 3: first 3 letters, capitalized (e.g., 'Aptara' -> 'APT')
    - If single word and length < 3: first 2 letters, capitalized (e.g., 'new' -> 'NE')
    - If two words: first 3 letters of each word, capitalized and concatenated (e.g., 'American Express' -> 'AMEEXP')
    """
    if not client_name:
        return ""
    words = client_name.split()
    if len(words) == 1:
        word = words[0]
        if len(word) >= 3:
            return word[:3].upper()
        else:
            return word[:2].upper()
    elif len(words) == 2:
        return (words[0][:3] + words[1][:3]).capitalize()
    else:
        # For more than two words, use first letter of each word, up to 6 chars
        return ''.join(w[0] for w in words).upper()[:6]
        
# ========== TABLE BORDER STYLING ==========
def set_table_borders(table):
    """Apply borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
 
    # Check if tblBorders element exists, create if not
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
 
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)

def add_markdown_table_to_doc(container, table_text):
    """Convert markdown table into DOCX table with correct header-to-data mapping."""
    lines = [l.strip() for l in table_text.strip().split("\n") if l.strip()]
    if not lines:
        return

    # Remove Markdown separator row (---|---|---)
    clean_lines = []
    for l in lines:
        if not re.match(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$', l):
            clean_lines.append(l)

    if not clean_lines:
        return

    # Extract header
    header = [h.strip() for h in clean_lines[0].strip('|').split('|')]
    num_cols = len(header)

    # Extract rows
    data_rows = []
    for row in clean_lines[1:]:
        cells = [c.strip() for c in row.strip('|').split('|')]
        if len(cells) < num_cols:
            cells += [""] * (num_cols - len(cells))
        elif len(cells) > num_cols:
            cells = cells[:num_cols]
        data_rows.append(cells)

    # Create DOCX table in container
    if isinstance(container, DocumentType):
        table = container.add_table(rows=len(data_rows) + 1, cols=num_cols)
    elif isinstance(container, _Cell):
        table = container.add_table(rows=len(data_rows) + 1, cols=num_cols)
    else:
        raise TypeError("Unsupported container type")

    table.style = "Table Grid"

# OLD CODE FOR FILLING TABLE - COMMENTED OUT TO AVOID ISSUES WITH FORMATTING AND IMAGES IN CELLS
    # # Fill header
    # for j, col in enumerate(header):
    #     table.cell(0, j).text = col
    #     table.cell(0, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # # Fill data rows
    # for i, row in enumerate(data_rows, start=1):
    #     for j, col in enumerate(row):
    #         table.cell(i, j).text = col

    # set_table_borders(table)

# NEW CODE FOR FILLING TABLE - HANDLES FORMATTING AND IMAGES IN CELLS

    # Fill header
    for j, col in enumerate(header):
        cell = table.cell(0, j)
        cell.paragraphs[0].clear()
        add_formatted_text(cell.paragraphs[0], col)

    # Fill data rows
    for i, row in enumerate(data_rows, start=1):
        for j, col in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].clear()
            add_formatted_text(cell.paragraphs[0], col)

def add_formatted_text(paragraph, text):
    """Adds bold and italic formatting to an existing paragraph."""
    if not text:
        return
   
    # First, check if the entire text is wrapped in ** (for PowerPoint style headings)
    if text.startswith('**') and text.endswith('**'):
        # Count how many ** pairs we have
        bold_pairs = text.count('**')
        if bold_pairs == 2:  # Only if there's exactly one pair at start and end
            run = paragraph.add_run(text[2:-2])
            run.bold = True
            return
   
    # For mixed formatting within text (Markdown style)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def add_image_content(container, image_content):
    """Add image content to container"""
    if not image_content:
        return
    
    # Extract base64 data
    if 'base64,' in image_content:
        base64_data = image_content.split('base64,')[1]
    else:
        base64_data = image_content
    
    try:
        image_bytes = base64.b64decode(base64_data)
        image_stream = io.BytesIO(image_bytes)
        
        if isinstance(container, DocumentType):
            container.add_picture(image_stream, width=Inches(4))
            container.add_paragraph()  # Add spacing
        elif isinstance(container, _Cell):
            p = container.add_paragraph()
            run = p.add_run()
            run.add_picture(image_stream, width=Inches(3))  # Smaller for table cells
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Warning: Could not add image: {e}")

def add_text_content(container, text_content):
    """Add text content to container, handling tables, headings, lists, and regular text for both formats"""
    if not text_content:
        return
   
    lines = text_content.split('\n')
    is_in_table = False
    table_buffer = []
    current_paragraph = None
    previous_line = None  # Track previous line to detect headings after ---
   
    def add_paragraph_to_container(c, text="", style=None):
        if isinstance(c, (DocumentType, _Cell)):
            return c.add_paragraph(text, style=style)
        else:
            raise TypeError("Unsupported container type")
   
    for i, line in enumerate(lines):
        stripped_line = line.strip()
       
        # Store current line as previous for next iteration
        current_line = stripped_line
        if i > 0:
            previous_line = lines[i-1].strip()
       
        if not stripped_line:
            # Empty line - reset current paragraph to force new paragraph for next content
            current_paragraph = None
            continue
           
        # Handle PowerPoint separator lines (---) - treat as section breaks
        if stripped_line.startswith('---') and len(stripped_line) >= 3:
            # Check if there's content after the separator that could be a heading
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # If next line is a bold heading, use that as the section title
                if next_line.startswith('**') and next_line.endswith('**') and next_line.count('**') == 2:
                    # The next line will be handled as a heading, so just skip this separator
                    continue
                else:
                    # No clear heading after separator, create a generic section break
                    if isinstance(container, DocumentType):
                        container.add_heading("Section Break", level=2)
                    else:
                        p = container.add_paragraph()
                        run = p.add_run("Section Break")
                        run.bold = True
                    current_paragraph = None
            else:
                # This is the last line, just skip it
                continue
            continue
           
        # Handle tables (for first outline format)
        if '|' in stripped_line:
            # Check if this is a table separator line
            if re.match(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$', stripped_line):
                if is_in_table:
                    table_buffer.append(line)
                continue
               
            if not is_in_table:
                is_in_table = True
                table_buffer.clear()
           
            table_buffer.append(line)
            current_paragraph = None
            continue
        else:
            # If we were in a table, process it
            if is_in_table and table_buffer:
                add_markdown_table_to_doc(container, '\n'.join(table_buffer))
                is_in_table = False
                table_buffer.clear()
 
        # Handle regular text content (not in table)
        if not is_in_table:
            # Handle headings for both formats
            if stripped_line.startswith('#'):
                # Markdown headings (first outline)
                if stripped_line.startswith('###'):
                    level = 3
                elif stripped_line.startswith('##'):
                    level = 2
                else:
                    level = 1
               
                heading_text = stripped_line.replace('#', '').strip()
               
                if isinstance(container, DocumentType):
                    container.add_heading(heading_text, level=level)
                else:
                    p = container.add_paragraph()
                    run = p.add_run(heading_text)
                    run.bold = True
                    if level == 1:
                        run.font.size = Pt(14)
                    elif level == 2:
                        run.font.size = Pt(12)
                current_paragraph = None
           
            # Handle PowerPoint-style bold key terms (entire line in **)
            elif stripped_line.startswith('**') and stripped_line.endswith('**'):
                # Count the number of ** pairs to ensure it's exactly one pair
                bold_pairs = stripped_line.count('**')
                if bold_pairs == 2:  # Exactly one pair at start and end
                    # Check if this comes after a --- separator to determine if it's a section heading
                    is_section_heading = (i > 0 and lines[i-1].strip().startswith('---'))
                   
                    if is_section_heading:
                        # Treat as section heading
                        if isinstance(container, DocumentType):
                            container.add_heading(stripped_line[2:-2], level=1)
                        else:
                            p = container.add_paragraph()
                            run = p.add_run(stripped_line[2:-2])
                            run.bold = True
                            run.font.size = Pt(14)
                    else:
                        # Treat as regular bold paragraph
                        if current_paragraph is None:
                            current_paragraph = add_paragraph_to_container(container)
                        elif current_paragraph.text:
                            current_paragraph = add_paragraph_to_container(container)
                        run = current_paragraph.add_run(stripped_line[2:-2])
                        run.bold = True
                   
                    current_paragraph = None  # Reset for next content
                else:
                    # It's mixed formatting, use the regular function
                    if current_paragraph is None:
                        current_paragraph = add_paragraph_to_container(container)
                    elif current_paragraph.text:
                        current_paragraph.add_run().add_break()
                    add_formatted_text(current_paragraph, stripped_line)
           
        # OLD CODE FOR BULLET POINTS - COMMENTED OUT TO AVOID ISSUES WITH MIXED FORMATTING
            # # Handle bullet points for both formats
            # elif stripped_line.startswith('-') or stripped_line.startswith('*'):
            #     # Handle bullet points - ALWAYS create new bullet paragraph
            #     current_paragraph = add_paragraph_to_container(container, style='List Bullet')
               
            #     # Remove the bullet character and process the text with formatting
            #     bullet_content = stripped_line[1:].strip()
            #     add_formatted_text(current_paragraph, bullet_content)
            #     # Don't reset current_paragraph here to allow line breaks within same bullet if needed
        
        # NEW CODE FOR BULLET POINTS AND NUMBERED LISTS - CHECKS FOR MIXED FORMATTING

            elif re.match(r'^(\s*)([-*•◦▪])\s', line):  # also catch • ◦ ▪
                indent_count = len(line) - len(line.lstrip())
                level = indent_count // 2  # 2 spaces per level
                if level == 0:
                    style = 'List Bullet'
                elif level == 1:
                    style = 'List Bullet 2'
                else:
                    style = 'List Bullet 3'
                current_paragraph = add_paragraph_to_container(container, style=style)
                bullet_content = stripped_line.lstrip('-*•◦▪').strip()
                add_formatted_text(current_paragraph, bullet_content)

    #-------------------------------------------------------------------------------------
            elif re.match(r'^\d+\.', stripped_line):
                # Handle numbered lists - ALWAYS create new numbered paragraph
                current_paragraph = add_paragraph_to_container(container, style='List Number')
                numbered_content = re.sub(r'^\d+\.\s*', '', stripped_line)
                add_formatted_text(current_paragraph, numbered_content)
                # Don't reset current_paragraph here to allow line breaks within same numbered item if needed
           
            else:
                # Regular text - handle mixed formatting
                if current_paragraph is None:
                    current_paragraph = add_paragraph_to_container(container)
                elif current_paragraph.text and current_paragraph.style.name not in ['List Bullet', 'List Number']:
                    # Only add line break if not in a list (lists handle their own formatting)
                    current_paragraph.add_run().add_break()
                add_formatted_text(current_paragraph, line.strip())
   
    # Process any remaining table content
    if is_in_table and table_buffer:
        add_markdown_table_to_doc(container, '\n'.join(table_buffer))

def add_text_or_table_or_image(container, content):
    """
    Process content while maintaining the exact order of text, tables, and images as they appear.
    """
    if not content:
        return
    
    # First, extract all image patterns and their positions
    image_patterns = []

    for match in re.finditer(r'!\[.*?\]\((data:image/.*?;base64,.*?)\)', content):
        image_patterns.append({
            'start': match.start(),
            'end': match.end(),
            'pattern': match.group(1),
            'type': 'markdown'
        })
    
    image_patterns.sort(key=lambda x: x['start'])
    
    # Split the content into segments: text and images in order
    segments = []
    last_pos = 0
    
    for img in image_patterns:
        # Add text before the image
        if img['start'] > last_pos:
            text_segment = content[last_pos:img['start']].strip()
            if text_segment:
                segments.append({'type': 'text', 'content': text_segment})
        
        # Add the image
        segments.append({'type': 'image', 'content': img['pattern']})
        last_pos = img['end']
    
    # Add remaining text after the last image
    if last_pos < len(content):
        text_segment = content[last_pos:].strip()
        if text_segment:
            segments.append({'type': 'text', 'content': text_segment})
    
    # If no images were found, treat the whole content as text
    if not segments:
        segments.append({'type': 'text', 'content': content})
    
    # Process each segment in order
    for segment in segments:
        if segment['type'] == 'text':
            # Process text content (which may include tables)
            add_text_content(container, segment['content'])
        elif segment['type'] == 'image':
            # Process image content
            add_image_content(container, segment['content'])


# def upload_to_azure_storage(stream, filename: str, client: str = None, project: str = None, module: str = None):
#     """
#     Upload document stream to Azure Data Lake Storage (ADLS Gen2) with proper path structure.
    
#     Args:
#         stream: The file content stream
#         filename: The filename to use
#         client: Client name for path structure
#         project: Project name for path structure
#         module: Module name for path structure
        
#     Returns:
#         str: URL path to the uploaded file
#     """
#     try:
#         # Create DataLakeServiceClient with connection pooling
#         datalake_service_client = DataLakeServiceClient(
#             account_url=f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net",
#             credential=AZURE_STORAGE_ACCOUNT_KEY
#         )

#         # Get filesystem client
#         file_system_client = datalake_service_client.get_file_system_client(
#             AZURE_STORAGE_FILESYSTEM_NAME
#         )

#         # Construct the appropriate path
#         if client and project and (module and module != "project-level"):
#             directory_path = f"{client}/{project}/{module}/Outline"
#             file_path = f"{directory_path}/{filename}"
#             # Create directories if they don't exist
#             try:
#                 file_system_client.create_directory(directory_path)
#                 logger.info(f"Created or verified directory path: {directory_path}")
#             except Exception as e:
#                 # Directory might already exist - this is fine
#                 logger.debug(f"Directory path check: {e}")

#         elif client and project:
#             directory_path = f"{client}/{project}/Outline"
#             file_path = f"{directory_path}/{filename}"
#             try:
#                 file_system_client.create_directory(directory_path)
#                 logger.info(f"Created or verified directory path: {directory_path}")
#             except Exception as e:
#                 # Directory might already exist - this is fine
#                 logger.debug(f"Directory path check: {e}")
#         else:
#             # Fallback to simple filename if path components not provided
#             file_path = filename
#             logger.warning(f"Uploading to root path because client/project/module not provided: {filename}")

#         # Ensure stream is at beginning
#         stream.seek(0)
#         data = stream.read()

#         # Create file client and upload with retry policy
#         file_client = file_system_client.get_file_client(file_path)
#         file_client.upload_data(data, overwrite=True)

#         # Build and return the ADLS file path
#         full_path = f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net/{AZURE_STORAGE_FILESYSTEM_NAME}/{file_path}"
#         logger.info(f"Successfully uploaded to ADLS: {full_path}")
#         return full_path

#     except Exception as e:
#         logger.error(f"Error uploading to ADLS: {e}", exc_info=True)
#         raise


def upload_to_azure_storage(stream, filename, client, project, module):
    datalake_service_client = DataLakeServiceClient(
        account_url=f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net",
        credential=AZURE_STORAGE_ACCOUNT_KEY
    )

    file_system_client = datalake_service_client.get_file_system_client(AZURE_STORAGE_FILESYSTEM_NAME)

    # ✅ Build the directory path properly
    if not module or module == "project-level":
        directory_path = f"{client}/{project}/Outline"
    else:
        directory_path = f"{client}/{project}/{module}/Outline"

    # ✅ Ensure the directory exists (create if missing)
    try:
        dir_client = file_system_client.get_directory_client(directory_path)
        dir_client.get_properties()  # Try fetching metadata
    except Exception:
        print(f"Creating missing directory path: {directory_path}")
        dir_client = file_system_client.create_directory(directory_path)

    # ✅ Create file client
    file_path = f"{directory_path}/{filename}"
    file_client = file_system_client.get_file_client(file_path)

    # Always create/overwrite file before appending
    print(f"Creating file: {file_path}")
    file_client.create_file()  # this actually creates an empty file

    # ✅ Upload in chunks
    stream.seek(0)
    chunk_size = 4 * 1024 * 1024  # 4 MB
    offset = 0

    while True:
        chunk = stream.read(chunk_size)
        if not chunk:
            break
        file_client.append_data(data=chunk, offset=offset, length=len(chunk))
        offset += len(chunk)

    # ✅ Finalize upload
    file_client.flush_data(offset)

    print(f"✅ Successfully uploaded to ADLS: {file_path}")
    return file_path

def clean_text(text: str) -> str:
    """
    Remove NULL bytes and invalid control characters from text.
    Keeps only XML-compatible characters.
    """
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    # Remove ASCII control chars except \t, \n, \r
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    
def save_outline_to_adls(data: dict, filename: str = None):
    """
    Save outline to ADLS with proper path structure and naming convention.
    
    Args:
        data: The outline data dictionary
        filename: Optional custom filename. If not provided, will generate based on convention.
        
    Returns:
        str: URL path to the saved file
    """
    # Extract client/project/module for path structure
    client = clean_text(data.get("client", ""))
    project = clean_text(data.get("project", ""))
    module = clean_text(data.get("module", ""))
    
    # Generate standardized filename if not provided
    if not filename:
        if str(module) == "project-level":
            client_short_name = get_client_short_name(client)
            filename = f"{client_short_name}_{project}_CO.docx"
        else:
            client_short_name = get_client_short_name(client)
            module_name = module.replace(" ", "")
            filename = f"{client_short_name}_{module_name}_CO.docx"
    
    # Create document in memory
    doc = Document()
    
    # Add main title
    doc.add_heading("Outline Document", 0)
    
    # Calculate total duration first
    outlines = data.get("outlines", [])
    total_duration_minutes = 0
    total_duration_seconds = 0
    
    for outline in outlines:
        duration_str = clean_text(str(outline.get("Durations_min", "")))
        try:
            minutes_match = re.search(r'(\d+)\s*min', duration_str)
            seconds_match = re.search(r'(\d+)\s*sec', duration_str)
            
            minutes = int(minutes_match.group(1)) if minutes_match else 0
            seconds = int(seconds_match.group(1)) if seconds_match else 0
            
            total_duration_minutes += minutes
            total_duration_seconds += seconds
            
            if total_duration_seconds >= 60:
                total_duration_minutes += total_duration_seconds // 60
                total_duration_seconds = total_duration_seconds % 60
        except (ValueError, AttributeError):
            try:
                numbers = re.findall(r'\d+', duration_str)
                if numbers:
                    total_duration_minutes += int(numbers[0])
            except (ValueError, IndexError):
                pass

    # Format total duration
    total_duration_str = f"{total_duration_minutes} min"
    if total_duration_seconds > 0:
        total_duration_str += f" {total_duration_seconds} sec"
    
    # Add project information
    doc.add_heading("Project Information", level=1)
    doc.add_paragraph(f"Client: {client}")
    doc.add_paragraph(f"Project: {project}")
    doc.add_paragraph(f"Module: {module}")
    doc.add_paragraph(f"Duration: {total_duration_str}")
    doc.add_paragraph()
    
    # Add context_prompt section
    context_prompt = data.get("context_prompt")
    doc.add_heading("User Prompt", level=2)
    if context_prompt:
        add_text_or_table_or_image(doc, clean_text(context_prompt))
    else:
        doc.add_paragraph("No user prompt provided.")
    doc.add_paragraph()
   
    # Add topic section
    topic = data.get("topic")
    doc.add_heading("Topic", level=2)
    if topic:
        add_text_or_table_or_image(doc, clean_text(topic))
    else:
        doc.add_paragraph("No topic specified.")
    doc.add_paragraph()
    
    # Outline content heading
    doc.add_heading("Outline Content", level=1)
    doc.add_paragraph()

    # Table headers with new order
    headers = [
        "Chapter",
        "Topic", 
        "Subtopic",
        "Full Page Content",
        "Duration (Mins)",
        "File",
        "Source Page"
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = clean_text(header)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add outline items with new column order
    for outline in outlines:
        row_cells = table.add_row().cells

        row_cells[0].text = clean_text(outline.get("Chapter", ""))      # Chapter
        row_cells[1].text = clean_text(outline.get("Topic", ""))        # Topic
        row_cells[2].text = clean_text(outline.get("Subtopic", ""))     # Subtopic

        # Full Page Content
        # [PATCH] Use the new markdown-aware renderer when available.
        # The legacy path ADDS "Section Break" as a heading whenever it sees
        # a --- separator, which is the source of the "Section Break not from
        # source" review comments. The new renderer correctly filters these,
        # handles **bold**, *italic*, <RED>, <BLUE>, [TABLE], and [link](url).
        raw_content = outline.get("Full_Page_Content", "")
        # [PATCH] Resolve the images list from the chunk so the renderer can
        # embed images at the position of [IMAGE: id] markers in content.
        # Different upstream paths use different field names; try all of them.
        chunk_images = (
            outline.get("images")
            or outline.get("Images")
            or outline.get("source_images")
            or []
        )
        # Some upstream paths store images as a JSON-encoded string
        if isinstance(chunk_images, str):
            try:
                chunk_images = json.loads(chunk_images)
            except Exception:
                chunk_images = []

        if raw_content and raw_content.strip():
            if _CO_RENDERER_AVAILABLE:
                # Clear existing paragraphs in the cell
                if row_cells[3].paragraphs:
                    row_cells[3]._element.clear()
                    # Re-add an empty paragraph so the cell has one to work with
                    row_cells[3]._element.append(
                        row_cells[3].add_paragraph()._element
                    )
                # [PATCH] Pass images list so [IMAGE: id] markers get resolved
                # to actual embedded images at their source position.
                render_content_into_cell(row_cells[3], raw_content, images=chunk_images)
            else:
                # Legacy fallback
                content = clean_text(raw_content)
                if row_cells[3].paragraphs:
                    row_cells[3]._element.clear()
                add_text_or_table_or_image(row_cells[3], content)

        row_cells[4].text = clean_text(outline.get("Durations_min", "")) # Duration
        row_cells[5].text = clean_text(outline.get("File", ""))         # File
        row_cells[6].text = clean_text(outline.get("Source_Page", ""))  # Source Page

        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Prepare ADLS client and directory path
    datalake_service_client = DataLakeServiceClient(
        account_url=f"https://{AZURE_STORAGE_ACCOUNT_NAME}.dfs.core.windows.net",
        credential=AZURE_STORAGE_ACCOUNT_KEY
    )
    file_system_client = datalake_service_client.get_file_system_client(AZURE_STORAGE_FILESYSTEM_NAME)
    # directory_path = f"{client}/{project}/{module}/Outline/"

    if not module or module == "project-level":
        directory_path = f"{client}/{project}/Outline/"
    else:
        directory_path = f"{client}/{project}/{module}/Outline/"

    # Use versioned filename logic
    versioned_filename = get_versioned_filename(directory_path, filename, file_system_client)

    # Save document to memory
    stream = io.BytesIO()
    doc.save(stream)

    # Upload to Azure with versioned filename
    return upload_to_azure_storage(stream, versioned_filename, client, project, module)

# def add_html_table_to_doc(container, html_table):
#     """
#     Parse an HTML table and add it as a real docx table to the container.
#     """
#     soup = BeautifulSoup(html_table, "html.parser")
#     table_tag = soup.find("table")
#     if not table_tag:
#         return

#     rows = table_tag.find_all("tr")
#     if not rows:
#         return

#     # Determine number of columns
#     first_row = rows[0]
#     cols = first_row.find_all(["td", "th"])
#     num_cols = len(cols)

#     # Create table in docx
#     if isinstance(container, DocumentType):
#         docx_table = container.add_table(rows=0, cols=num_cols)
#     elif isinstance(container, _Cell):
#         docx_table = container.add_table(rows=0, cols=num_cols)
#     else:
#         raise TypeError("Unsupported container type")

#     docx_table.style = "Table Grid"

#     for row in rows:
#         docx_row = docx_table.add_row().cells
#         cells = row.find_all(["td", "th"])
#         for i, cell in enumerate(cells):
#             docx_row[i].text = cell.get_text(strip=True)

#     set_table_borders(docx_table)


# def add_html_table_to_doc(container, html_table):
#     """
#     Parse an HTML table and add it as a real docx table to the container.
#     """
#     soup = BeautifulSoup(html_table, "html.parser")
#     table_tag = soup.find("table")
#     if not table_tag:
#         return

#     rows = table_tag.find_all("tr")
#     if not rows:
#         return

#     # Determine number of columns from the first row
#     first_row = rows[0]
#     cols = first_row.find_all(["td", "th"])
#     num_cols = len(cols)

#     # Create table in docx
#     if isinstance(container, DocumentType):
#         docx_table = container.add_table(rows=0, cols=num_cols)
#     elif isinstance(container, _Cell):
#         docx_table = container.add_table(rows=0, cols=num_cols)
#     else:
#         raise TypeError("Unsupported container type")

#     docx_table.style = "Table Grid"

#     for row in rows:
#         docx_row = docx_table.add_row().cells
#         cells = row.find_all(["td", "th"])
#         # Pad or truncate cells to match num_cols
#         cell_texts = [cell.get_text(strip=True) for cell in cells]
#         if len(cell_texts) < num_cols:
#             cell_texts += [""] * (num_cols - len(cell_texts))
#         elif len(cell_texts) > num_cols:
#             cell_texts = cell_texts[:num_cols]
#         for i, text in enumerate(cell_texts):
#             docx_row[i].text = text

#     set_table_borders(docx_table)

def add_html_table_to_doc(container, html_table):
    """
    Parse an HTML table and add it as a real docx table to the container.
    """
    soup = BeautifulSoup(html_table, "html.parser")
    table_tag = soup.find("table")
    if not table_tag:
        return

    rows = table_tag.find_all("tr")
    if not rows:
        return

    # Determine number of columns from the first non-empty row
    for r in rows:
        cols = r.find_all(["td", "th"])
        if cols:
            num_cols = len(cols)
            break
    else:
        num_cols = 1

    # Create docx table
    if hasattr(container, "add_table"):
        docx_table = container.add_table(rows=0, cols=num_cols)
    elif hasattr(container, "tables"):
        docx_table = container.add_table(rows=0, cols=num_cols)
    else:
        raise TypeError("Unsupported container type for adding a table")

    docx_table.style = "Table Grid"

    # Loop through HTML rows safely
    for row in rows:
        cells = row.find_all(["td", "th"])
        docx_row = docx_table.add_row().cells

        cell_texts = [cell.get_text(strip=True) for cell in cells]

        # Normalize cell count to num_cols
        if len(cell_texts) < len(docx_row):
            cell_texts += [""] * (len(docx_row) - len(cell_texts))
        elif len(cell_texts) > len(docx_row):
            cell_texts = cell_texts[:len(docx_row)]

        # Safely fill cells
        for i in range(min(len(docx_row), len(cell_texts))):
            docx_row[i].text = cell_texts[i]

    set_table_borders(docx_table)
    
def add_text_or_table_or_image(container, content):
    """
    Process content while maintaining the exact order of text, tables, HTML tables, and images as they appear.
    """
    if not content:
        return

    # Detect and process HTML tables
    html_table_matches = list(re.finditer(r'(<table[\s\S]*?</table>)', content, re.IGNORECASE))
    last_pos = 0
    for match in html_table_matches:
        start, end = match.span()
        # Add text before the table
        if start > last_pos:
            text_segment = content[last_pos:start].strip()
            if text_segment:
                _process_text_and_images(container, text_segment)
        # Add the HTML table
        add_html_table_to_doc(container, match.group(1))
        last_pos = end

    # Add any remaining text after the last table
    if last_pos < len(content):
        text_segment = content[last_pos:].strip()
        if text_segment:
            _process_text_and_images(container, text_segment)


def _process_text_and_images(container, content):
    """
    Helper to process text and images in order, including markdown tables.
    Handles both markdown and HTML image formats.
    """
    # Combine both markdown and HTML image patterns
    image_patterns = []
   
    # Pattern 1: Markdown style - ![alt text](data:image/...)
    for match in re.finditer(r'!\[.*?\]\((data:image/.*?;base64,[^)]+)\)', content):
        image_patterns.append({
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1),  # base64 data only
            'type': 'markdown'
        })
   
    # Pattern 2: HTML style - <img src="data:image/..." />
    for match in re.finditer(r'<img\s+[^>]*src="(data:image/.*?;base64,[^"]+)"[^>]*>', content):
        image_patterns.append({
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1),  # base64 data only
            'type': 'html'
        })
   
    # Pattern 3: HTML style with single quotes - <img src='data:image/...' />
    for match in re.finditer(r"<img\s+[^>]*src='(data:image/.*?;base64,[^']+)'[^>]*>", content):
        image_patterns.append({
            'start': match.start(),
            'end': match.end(),
            'content': match.group(1),  # base64 data only
            'type': 'html'
        })
 
    # Sort by position in content
    image_patterns.sort(key=lambda x: x['start'])
 
    segments = []
    last_pos = 0
 
    # Split content into segments (text and images)
    for img in image_patterns:
        # Add text segment before this image
        if img['start'] > last_pos:
            text_segment = content[last_pos:img['start']].strip()
            if text_segment:
                segments.append({'type': 'text', 'content': text_segment})
       
        # Add image segment
        segments.append({'type': 'image', 'content': img['content']})
        last_pos = img['end']
 
    # Add remaining text after last image
    if last_pos < len(content):
        text_segment = content[last_pos:].strip()
        if text_segment:
            segments.append({'type': 'text', 'content': text_segment})
 
    # If no segments found, treat entire content as text
    if not segments:
        segments.append({'type': 'text', 'content': content})
 
    # Process each segment
    for segment in segments:
        if segment['type'] == 'text':
            # Clean HTML tags except for tables - using your original approach
            clean_text = BeautifulSoup(segment['content'], "html.parser").get_text()
            add_text_content(container, clean_text)
        elif segment['type'] == 'image':
            # Add image content
            add_image_content(container, segment['content'])

def get_versioned_filename(directory_path: str, filename: str, file_system_client) -> str:
    """
    Returns a versioned filename if the file already exists in ADLS.
    Example: If 'ÁME_Module1_CO.docx' exists, returns 'ÁME_Module1_CO(1).docx', etc.
    """
    base, ext = os.path.splitext(filename)
    version = 1
    candidate = filename
    
    # Remove trailing slash if it exists
    directory_path = directory_path.rstrip('/')
        
    while True:
        # Build the full path with a single slash
        file_path = f"{directory_path}/{candidate}"
        try:
            file_client = file_system_client.get_file_client(file_path)
            exists = file_client.exists()
        except Exception as e:
            logger.error(f"Error checking if file exists: {e}")
            return candidate  # Return original filename on error
        
        if not exists:
            return candidate
        
        candidate = f"{base}({version}){ext}"  # Version before extension
        version += 1
